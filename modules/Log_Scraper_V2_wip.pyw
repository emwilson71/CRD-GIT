#---------------------------------------------------------------------------------------------
# Log_Scraper_V2_83
# JSmyser
#---------------------------------------------------------------------------------------------

import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, BooleanVar, ttk
import sys
import os
import re
import json
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import time
import tkcalendar  # Added for calendar popups
import subprocess
import webbrowser
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import fnmatch

# Default URLs for FIXS and Intranet searches
DEFAULT_FIXS_BASE_URL = 'https://canonmedical.my.salesforce.com/_ui/search/ui/UnifiedSearchResults?searchType=2&sen=a02&sen=a0s&sen=500&sen=005&sen=a0Y&sen=ka&sen=00O&str="" ""'
DEFAULT_INTRANET_BASE_URL = 'https://intranet.cmsu.com/_layouts/15/osssearchresults.aspx?u=https%3A%2F%2Fintranet%2Ecmsu%2Ecom&k="" ""'
DEFAULT_7Z_PATH = r"C:\Program Files\7-Zip\7z.exe"

def get_config_path(filename):
    if getattr(sys, 'frozen', False):
        # If the application is run as a bundle (exe), use the directory of the executable
        return os.path.join(os.path.dirname(sys.executable), filename)
    else:
        # Otherwise, use the directory of the Python script
        return os.path.join(os.path.dirname(__file__), filename)
    

#### Here is the start of my TextFileSearch Class and init. ------------------------------------------------------ 

class TextFileSearchGUI:
    def __init__(self, master):
        self.master = master
        master.geometry("750x730+5+10")
        master.title("Log_Scraper_V2.83")
        master.grid_columnconfigure(4, weight=1)
        master.grid_rowconfigure(7, weight=1)
        master.bind("<Configure>", self.on_resize)
        
        self.search_var = tk.StringVar()
        self.initial_search_done = False
        self.last_browse_location = {"path": ""}
        self.searching = None
        self.preferences = {}
        self.config_path = get_config_path("log_scraper_settings.json")
        self.load_preferences()
        
        self.search_paths = self.preferences.get('search_paths', {})
        self.search_path_order = self.preferences.get('search_path_order', [])
        self.preferences.setdefault('recent_count', 5)
        self.recent_count_var = tk.IntVar(value=self.preferences.get('recent_count', 5))
        
        # Ensure Recent 1-N are in search_path_order
        recent_count = self.recent_count_var.get()
        for i in range(1, recent_count + 1):
            name = f"Recent {i}"
            if name not in self.search_path_order:
                self.search_path_order.append(name)
        
        self.manage_window = None
        self.fixs_setup_window = None
        self.intranet_setup_window = None
        self.windows_setup_window = None
        self.seven_zip_path = self.preferences.get('seven_zip_path', DEFAULT_7Z_PATH)
        
        self.save_prefs_var = tk.BooleanVar(value=self.preferences.get('save_entries', False))
        self.show_last_file_var = tk.BooleanVar(value=self.preferences.get('show_last_file', False))
        self.path_var = tk.StringVar(value=self.preferences.get('path', ''))
        self.extensions_var = tk.StringVar(value=self.preferences.get('extensions', ''))
        self.main_term_var = tk.StringVar(value=self.preferences.get('main_term', ''))
        self.following_terms_var = tk.StringVar(value=self.preferences.get('following_terms', ''))
        self.main_context = tk.StringVar(value=self.preferences.get('main_context', '0,0'))
        self.following_context = tk.StringVar(value=self.preferences.get('following_context', '0,0'))
        self.search_locations = self.preferences.get('search_locations', {})
        self.search_location_order = self.preferences.get('search_location_order', [])
        self.fixs_base_url = self.preferences.get('fixs_base_url', DEFAULT_FIXS_BASE_URL)
        self.intranet_base_url = self.preferences.get('intranet_base_url', DEFAULT_INTRANET_BASE_URL)
        if not hasattr(self, 'configure_graph'):  # Ensure it’s set even if load fails
            self.configure_graph = {}

    ## START OF ROW 0, Tools and Path Entry.
        self.options_var = tk.StringVar(value="               Tool Menu")
        self.options_dropdown = ttk.Combobox(master, textvariable=self.options_var, state="readonly", width=25)
        self.options_dropdown.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.options_dropdown.bind("<<ComboboxSelected>>", self.on_options_select)
        self.options_dropdown['postcommand'] = self.update_options_dropdown
        self.options_dropdown.configure(values=[])
        self.update_options_dropdown()

        self.select_folder_button = tk.Button(master, text="Select Folder", command=self.select_folder)
        self.select_folder_button.grid(row=0, column=1, padx=5, pady=5, sticky="e")
        self.select_folder_button.bind("<Double-Button-3>", self.toggle_advanced_extract)  # Double right-click binding
        self.advanced_extract = False
        self.select_file_button = tk.Button(master, text="Select File", command=self.select_file)
        self.select_file_button.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        
        path_frame = tk.Frame(master)
        path_frame.grid(row=0, column=3, columnspan=3, padx=5, pady=5, sticky="ew")
        self.path_drop_button = tk.Button(path_frame, text="▼", width=1, font=("TkDefaultFont", 7), command=lambda: [self.path_combo.focus_set(), self.path_combo.event_generate("<Alt-Down>")])
        self.path_drop_button.grid(row=0, column=0, sticky="w")
        self.path_combo = self.create_context_menu(ttk.Combobox(path_frame, textvariable=self.path_var, state="normal", width=38))
        self.path_combo.grid(row=0, column=1, sticky="ew")
        self.update_path_combo()
        self.path_combo.bind("<Double-Button-1>", lambda event: [self.path_combo.focus_set(), self.path_combo.event_generate("<Alt-Down>")])
        self.path_combo.bind('<<ComboboxSelected>>', self.on_combo_select)  # Add this
        self.path_combo.bind("<Return>", lambda event: self.start_search())  # Bind Enter to start_search
        path_frame.columnconfigure(1, weight=1)
    
    ## START OF ROW 1, Filter Entry.
        tk.Label(master, text="Folder/.Ext/File Filters :").grid(row=1, column=0, columnspan=2, sticky="e", padx=5, pady=5)
        self.extensions_entry = self.create_context_menu(tk.Entry(master, textvariable=self.extensions_var))
        self.extensions_entry.grid(row=1, column=2, columnspan=4, padx=5, pady=5, sticky="ew")
        self.extensions_entry.bind("<Return>", lambda event: self.start_search())  # Bind Enter to start_search

    ## START OF ROW 2, Main Term Entry.
        main_term_frame = tk.Frame(master)
        main_term_frame.grid(row=2, column=0, columnspan=6, sticky="ew", padx=(70,0), pady=5)
        main_term_frame.grid_columnconfigure(3, weight=1)
        tk.Label(main_term_frame, text="Main: [Before,After]").grid(row=0, column=0, sticky="e")
        self.main_context_entry = self.create_context_menu(tk.Entry(main_term_frame, textvariable=self.main_context, width=5))
        self.main_context_entry.grid(row=0, column=1, padx=5, sticky="w")
        self.main_context_entry.bind("<Return>", lambda event: self.start_search())  # Bind Enter to start_search
        tk.Label(main_term_frame, text="[Terms]").grid(row=0, column=2, sticky="e")
        self.main_term_entry = self.create_context_menu(tk.Entry(main_term_frame, textvariable=self.main_term_var))
        self.main_term_entry.grid(row=0, column=3, columnspan=3, padx=5, sticky="ew")
        self.main_term_entry.bind("<Return>", lambda event: self.start_search())  # Bind Enter to start_search

    ## START OF ROW 3, First Following Entry.
        first_match_frame = tk.Frame(master)
        first_match_frame.grid(row=3, column=0, columnspan=6, sticky="ew", padx=(20,0), pady=5)
        first_match_frame.grid_columnconfigure(3, weight=1)
        tk.Label(first_match_frame, text="First Following: [Before,After]").grid(row=0, column=0, sticky="e")
        self.following_context_entry = self.create_context_menu(tk.Entry(first_match_frame, textvariable=self.following_context, width=5))
        self.following_context_entry.grid(row=0, column=1, padx=5, sticky="w")
        self.following_context_entry.bind("<Return>", lambda event: self.start_search())  # Bind Enter to start_search
        tk.Label(first_match_frame, text="[Terms]").grid(row=0, column=2, sticky="e")
        self.following_terms_entry = self.create_context_menu(tk.Entry(first_match_frame, textvariable=self.following_terms_var))
        self.following_terms_entry.grid(row=0, column=3, columnspan=3, padx=5, sticky="ew")
        self.following_terms_entry.bind("<Return>", lambda event: self.start_search())  # Bind Enter to start_search
   
    ## START OF ROW 4, Date Filter Entry.
        self.date_filter_frame = tk.Frame(master)
        self.date_filter_frame.grid(row=4, column=1, columnspan=6, sticky="w", padx=5, pady=5)
        self.date_filter_var = tk.BooleanVar(value=self.preferences.get('date_filter_enabled', False))
        self.date_filter_checkbox = tk.Checkbutton(self.date_filter_frame, text="Filter by Date", variable=self.date_filter_var)
        self.date_filter_checkbox.grid(row=0, column=0, sticky="w")
        tk.Label(self.date_filter_frame, text="Start:").grid(row=0, column=1, sticky="e")
        self.start_date_entry = self.create_context_menu(tkcalendar.DateEntry(self.date_filter_frame, width=10, date_pattern='yyyy-mm-dd'))
        self.start_date_entry.grid(row=0, column=2, sticky="w")
        tk.Label(self.date_filter_frame, text="End:").grid(row=0, column=3, sticky="e")
        self.end_date_entry = self.create_context_menu(tkcalendar.DateEntry(self.date_filter_frame, width=10, date_pattern='yyyy-mm-dd'))
        self.end_date_entry.grid(row=0, column=4, sticky="w")
        tk.Label(self.date_filter_frame, text="Preset:").grid(row=0, column=5, sticky="e")
        self.quick_date_var = tk.StringVar(value=self.preferences.get('quick_date', ''))
        self.quick_date_combobox = ttk.Combobox(self.date_filter_frame, textvariable=self.quick_date_var, state="readonly", width=10)
        self.quick_date_combobox['values'] = ["1 day", "1 week", "1 month", "1 year", "5 years"]
        self.quick_date_combobox.grid(row=0, column=6, sticky="w")
        self.quick_date_combobox.bind("<<ComboboxSelected>>", self.on_quick_date_select)

        # Load raw values from preferences
        quick_date = self.preferences.get('quick_date', '')
        start_date = self.preferences.get('start_date', '')
        end_date = self.preferences.get('end_date', '')

        # If quick_date is set and not Custom, calculate dates
        if quick_date and quick_date != "Custom":
            today = datetime.today()
            if quick_date == "1 day":
                start_date = today # Same day for 1 day
            elif quick_date == "1 week":
                start_date = today - timedelta(days=6)
            elif quick_date == "1 month":
                start_date = today - relativedelta(months=1)  # 1 month
            elif quick_date == "1 year":
                start_date = today - relativedelta(years=1)  # 1 year
            elif quick_date == "5 years":
                start_date = today - relativedelta(years=5)  # 5 years
            self.start_date_entry.set_date(start_date.date())
            self.end_date_entry.set_date(today.date())
        else:
            # Use stored start_date and end_date (for Custom or no quick_date)
            if start_date:
                self.start_date_entry.set_date(start_date)
            else:
                self.start_date_entry.delete(0, 'end')
            if end_date:
                self.end_date_entry.set_date(end_date)
            else:
                self.end_date_entry.delete(0, 'end') 

        def focus_and_raise_calendar(event):
            cal = event.widget._top_cal
            cal.transient(self.master)
            cal.lift()
            cal.focus_set()
        self.start_date_entry.bind("<Button-1>", focus_and_raise_calendar, add="+")
        self.end_date_entry.bind("<Button-1>", focus_and_raise_calendar, add="+")
        self.start_date_entry.bind("<<DateEntrySelected>>", self.on_date_change)
        self.end_date_entry.bind("<<DateEntrySelected>>", self.on_date_change)
        self.quick_date_var.trace_add("write", self.on_quick_date_change)
        self.start_date_entry.bind("<Return>", lambda event: self.start_search())
        self.end_date_entry.bind("<Return>", lambda event: self.start_search())
    
    ## START OF ROW 5, Quick Search, Search, Export, and Clear Fields Control.
        search_button_frame = tk.Frame(master)
        search_button_frame.grid(row=5, column=0, columnspan=5, sticky="w", padx=20, pady=3)
        self.quick_search_var = tk.StringVar(value="")
        self.quick_search_dropdown = ttk.Combobox(search_button_frame, textvariable=self.quick_search_var, state="readonly", width=40)
        self.quick_search_dropdown.grid(row=0, column=0, columnspan=1, padx=10, pady=5, sticky="w")
        self.quick_search_dropdown.bind("<<ComboboxSelected>>", self.on_quick_search_select)
        self.update_quick_search_dropdown()

        self.search_button = tk.Button(search_button_frame, text="Search", command=self.start_search, width=15)
        self.search_button.grid(row=0, column=1, padx=5, sticky="w")
        self.export_button = tk.Button(search_button_frame, text="Export Results", command=self.export_results, state="disabled")
        self.export_button.grid(row=0, column=2, padx=8, sticky="w")
        self.clear_button = tk.Button(search_button_frame, text="Clear All Fields", command=self.clear_input_fields)
        self.clear_button.grid(row=0, column=3, padx=20, sticky="w")
      # TEMPORARY GRAPH BUTTON
        self.graph_button = tk.Button(search_button_frame, text="Graph Results", command=self.graph_results)
        self.graph_button.grid(row=0, column=4, padx=5, pady=5, sticky="ew")
        
    ## START OF ROW 6, self.output Cycle Files/Results and Search Control.
        files_results_frame = tk.Frame(master)
        files_results_frame.grid(row=6, column=0, columnspan=6, padx=15, pady=0, sticky="w")
        tk.Label(files_results_frame, text="Cycle Files").grid(row=0, column=0, padx=(5,0))
        self.next_file_button = tk.Button(files_results_frame, text="↑", command=lambda: self.cycle_output_files(-1))
        self.next_file_button.grid(row=0, column=1)
        self.prev_file_button = tk.Button(files_results_frame, text="↓", command=lambda: self.cycle_output_files(1))
        self.prev_file_button.grid(row=0, column=2)
        tk.Label(files_results_frame, text="Cycle Results").grid(row=0, column=3, padx=(10,0))
        self.prev_result_button = tk.Button(files_results_frame, text="←", command=lambda: self.cycle_output_results(-1))
        self.prev_result_button.grid(row=0, column=4)
        self.next_result_button = tk.Button(files_results_frame, text="→", command=lambda: self.cycle_output_results(1))
        self.next_result_button.grid(row=0, column=5)
        self.show_last_file_checkbox = tk.Checkbutton(files_results_frame, text="Auto-scroll to \nLatest File", variable=self.show_last_file_var, command=self.toggle_show_last_file)
        self.show_last_file_checkbox.grid(row=0, column=6, padx=10, sticky="e")
        self.output_search_frame = tk.Frame(files_results_frame)
        self.output_search_frame.grid(row=0, column=7, padx=5)
        tk.Label(self.output_search_frame, text="Search:").grid(row=0, column=0)
        self.output_search_var = tk.StringVar()
        self.output_search_entry = tk.Entry(self.output_search_frame, textvariable=self.output_search_var, width=20)
        self.output_search_entry.grid(row=0, column=1)
        self.output_search_entry.bind("<Return>", lambda event: self.search_output(self.output_search_var.get(), 1))
        self.output_search_entry.bind("<Shift-Return>", lambda event: self.search_output(self.output_search_var.get(), -1))
        tk.Button(self.output_search_frame, text="↑", command=lambda: self.search_output(self.output_search_var.get(), -1)).grid(row=0, column=2)
        tk.Button(self.output_search_frame, text="↓", command=lambda: self.search_output(self.output_search_var.get(), 1)).grid(row=0, column=3)
        self.output_result_label = tk.Label(files_results_frame, text="File 0 of 0, Result 0 of 0")
        self.output_result_label.grid(row=0, column=8, padx=5, sticky="w")
    ## Bind Arrow Keys to Cycle Files and Results
        self.master.bind("<Up>", lambda event: self.handle_arrow_keys(event, -1, "files"))
        self.master.bind("<Down>", lambda event: self.handle_arrow_keys(event, 1, "files"))
        self.master.bind("<Left>", lambda event: self.handle_arrow_keys(event, -1, "results"))
        self.master.bind("<Right>", lambda event: self.handle_arrow_keys(event, 1, "results"))
    #    master.bind("<Up>", lambda event: self.cycle_output_files(-1))
    #    master.bind("<Down>", lambda event: self.cycle_output_files(1))
    #    master.bind("<XButton1>", lambda event: self.cycle_output_results(-1)) #Supposed to be Mouse Back.
    #    master.bind("<XButton2>", lambda event: self.cycle_output_results(1))  #Supposed to be Mouse Back.

      
    ## START OF ROW 7, self.output.
        self.output = scrolledtext.ScrolledText(master, wrap=tk.WORD, width=70, height=20)
        self.output.tag_config("file_marker", font=("Arial", 14))
        self.output.tag_config("sel_file", foreground="red", font=("Arial", 14, "bold"))
        self.output.tag_config("sel_result", foreground="red", font=("Arial", 12, "bold"))
        self.output.tag_config("search", background="orange")
        self.output.tag_config("sel", background="#1e90ff", foreground="white")
        self.output.tag_raise("sel")
        self.output.grid(row=7, column=0, columnspan=6, padx=5, pady=5, sticky="nsew")
        self.output.bind("<Button-3>", self.show_context_menu)
        self.output.bind("<Double-1>", self.take_me_to_log_from_event)

    ## START OF ROW 8, Loading and Status Info.
        self.progress = ttk.Progressbar(master, orient="horizontal", length=180, mode="determinate")
        self.progress.grid(row=8, column=0, columnspan=3, padx=5, pady=5, sticky="ew")
        self.status_bar = tk.Label(master, text="Ready", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.grid(row=8, column=3, columnspan=3, padx=5, pady=5, sticky="ew")

    ## Final Initializing.        
        master.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.current_file = None
        self.last_temp_dir = None
        self.zip_cache = {}
        self.last_filter = ""
        self.search_complete = False
        self.log_viewer_states = {}
        self.clicked_lines = {}

        def handle_return(event):  ## This is for when there isn't a focus in an tk.entry. 
            focused_widget = self.master.focus_get()
            if focused_widget == self.output_search_entry:
                self.search_output(self.output_search_var.get(), 1)
            elif focused_widget in (self.path_combo, self.extensions_entry, self.main_context_entry, 
                                    self.main_term_entry, self.following_context_entry, self.following_terms_entry,
                                    self.start_date_entry, self.end_date_entry):
                self.start_search()
       # master.bind('<Return>', handle_return) # This has been disabled because I added bind <Return> to eack tk.entry individually. 

    def handle_arrow_keys(self, event, direction, action):
        # List of input widgets where arrow keys should not cycle files/results
        input_widgets = (
            self.path_combo,
            self.extensions_entry,
            self.main_context_entry,
            self.main_term_entry,
            self.following_context_entry,
            self.following_terms_entry,
            self.start_date_entry,
            self.end_date_entry,
            self.output_search_entry,
            self.quick_search_dropdown,
            self.options_dropdown,
            self.quick_date_combobox
        )
        focused_widget = self.master.focus_get()
        # Only cycle if no input widget is focused
        if focused_widget not in input_widgets:
            if action == "files":
                self.cycle_output_files(direction)
            elif action == "results":
                self.cycle_output_results(direction)
        return "break"  # Prevent default arrow key behavior in some cases




#### GRAPH RESULTS FROM MAIN. ------------------------------------------------------------------------
    def graph_results(self):
    # Check number of main terms
        main_terms = [term.strip() for term in self.main_term_var.get().split(',') if term.strip()]
        following_terms = [term.strip() for term in self.following_terms_var.get().split(',') if term.strip()]
        if following_terms and len(main_terms) > 1:
            tk.messagebox.showwarning("Multiple Main Terms", "Graph results only works with one main term when following terms are specified.")
            return
                
    # Hardcoded settings (for rule defaults only)
        default_settings = {
            # Setup Example. Make sure to use all lowercase Terms.
            "center frequency 1st value": {"unit": "MHz"},
            "rf level": {"unit": "dB"},
            "s.a.r. (head) 2nd value": {"unit": "W/Kg. (limit)"},
            "s.a.r. (partial) 2nd value": {"unit": "W/Kg. (limit)"},
            "s.a.r. (wb) 2nd value": {"unit": "W/Kg. (limit)"}
        }

        if not hasattr(self, 'configure_graph'):
            self.configure_graph = default_settings.copy()
        # Normalize self.configure_graph keys: strip := and whitespace, then lowercase
        self.configure_graph = {k.rstrip(':=').strip().lower(): v for k, v in self.configure_graph.items()}

    ## Create graph window
        graph_window = tk.Toplevel(self.master)
        graph_window.title("Graph Results")

    ## Position graph window
        width = 980
        height = 600
        x = self.master.winfo_x() + 50
        y = self.master.winfo_y() + 50
        graph_window.geometry(f"{int(width)}x{int(height)}+{int(x)}+{int(y)}")
        graph_window.update_idletasks()  # Ensure geometry is applied
        
    ## Initialize data structures
        all_data = {}
        examples = {}
        series_lines = {}
        series_vars = {}
        checked_terms = []  # Master list of active terms
        term_units = {}  # Term-to-unit mapping from configure_graph_data
        check_vars = {}
        series_colors = {}  # Persistent color mapping for each term
        file_filter_popup = None  # Track file filter popup
        config_popup = None  # Track the popup
        last_update = 0    
        annotation = None  # Tooltip annotation (defined once, updated in update_graph)
        last_window_size = None  # Track window size for resize debouncing

        ## Set vibrant 20-color cycle for series
        vibrant_colors = [
            '#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd',
            '#00b7eb', '#ff1493', '#32cd32', '#ffa500', '#6a5acd',
            '#ffd700', '#00ced1', '#c71585', '#228b22', '#ff4500',
            '#9932cc', '#ffff00', '#20b2aa', '#ff69b4', '#4682b4'
        ]
        plt.rcParams['axes.prop_cycle'] = plt.cycler(color=vibrant_colors)  # Use vibrant colors

    ## Graph Window Auto-resizing setup
        fig = plt.Figure(figsize=(8, 4))
        ax = fig.add_subplot(111)
        canvas = FigureCanvasTkAgg(fig, master=graph_window)
        canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        fig.subplots_adjust(left=0.1, right=0.90, top=0.9, bottom=0.125)  # Fixed margins

    ## Control frame (all buttons on top, custom order)
        control_frame = tk.Frame(graph_window)
        control_frame.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)
    ## Checkbox frame (series checkboxes, auto-adjust columns)
        checkbox_frame = tk.Frame(graph_window)
        checkbox_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=5)
        checkbox_frame.update()  # Ensure checkbox frame and all widgets are fully realized

    ## Get terms
        if following_terms:
            main_term = self.main_term_var.get().strip()
            all_terms = [main_term] + following_terms
        else:
            all_terms = main_terms

    ## Parse output for terms, examples, and date range
        output_text = self.output.get("1.0", tk.END).strip()
        lines = output_text.split('\n')
        dates = []
        files = set()
        date_pattern = r"\d{4}/\d{2}/\d{2} \d{2}:\d{2}:\d{2}"
        num_pattern = r"[-+]?\d*\.?\d+"
        for line in lines:
            if line.startswith(">>>> "):
                files.add(line.strip(">>>> ").strip(" <<<<"))
            date_match = re.search(date_pattern, line)
            if date_match:
                date_str = date_match.group()
                dates.append(datetime.strptime(date_str, "%Y/%m/%d %H:%M:%S"))

        ## Extract data for .acqsts and .sigusr1 files
        def extract_acqsts(lines, checked_terms):
            data = {}
            examples = {}
            term_units = {}
            # Include variants with and without : or =
            term_variants = {term: [term, term + " :", term.rstrip(":"),
                                    term + " =", term.rstrip("=")] for term in checked_terms}
            current_file = None
            for line in lines:
                if line.startswith(">>>> "):
                    current_file = line.strip(">>>> ").strip(" <<<<")
                    continue
                if not line.strip() or current_file is None:
                    continue
                parts = line.split(" ", 2)
                date_match = re.search(date_pattern, parts[1] + " " + parts[2])
                if not date_match:
                    continue
                date_str = date_match.group()
                date = datetime.strptime(date_str, "%Y/%m/%d %H:%M:%S")
                text = parts[2].strip()
                text_lower = text.lower()
                for term in checked_terms:
                    term_lower = term.lower()
                    for variant in term_variants.get(term.rstrip(':=').strip(), [term]):
                        variant_lower = variant.lower()
                        if variant_lower in text_lower:
                            text_after_term = text[text_lower.find(variant_lower) + len(variant_lower):].strip()
                            # Only add to examples for non-numeric or '=' terms
                            if '=' in term or not re.search(r"[-+]?\d*\.?\d+[eE]?[-+]?\d*", text_after_term):
                                if term not in examples:
                                    examples[term] = text
                                continue
                            nums = re.findall(r"[-+]?\d*\.?\d+[eE]?[-+]?\d*", text_after_term)
                            if not nums:
                                continue
                            last_idx = 0
                            # Normalize term by stripping : = and whitespace
                            base_term = term.strip().rstrip(':=').strip()
                            if len(nums) == 1:
                                # Single value: no suffix
                                term_with_value = base_term
                                if term_with_value not in data:
                                    data[term_with_value] = []
                                data[term_with_value].append((date, float(nums[0]), current_file))
                                if term_with_value not in examples:
                                    examples[term_with_value] = text
                                    num_idx = text_after_term.find(nums[0], last_idx)
                                    unit = ""
                                    if num_idx != -1:
                                        post_num = text_after_term[num_idx + len(nums[0]):].strip()
                                        unit_match = re.match(r"^\S+", post_num)
                                        unit = unit_match.group() if unit_match else ""
                                        last_idx = num_idx + len(nums[0])
                                    lookup_key = term_with_value.lower()
                                    unit = default_settings.get(lookup_key, {}).get("unit", unit)
                                    unit = self.configure_graph.get(lookup_key, {}).get("unit", unit)
                                    term_units[term_with_value] = unit
                            else:
                                # Multiple values: add 1st, 2nd, etc.
                                for i, value in enumerate(nums, start=1):
                                    value_label = "1st" if i == 1 else "2nd" if i == 2 else "3rd" if i == 3 else f"{i}th"
                                    term_with_value = f"{base_term} {value_label} Value"
                                    if term_with_value not in data:
                                        data[term_with_value] = []
                                    data[term_with_value].append((date, float(value), current_file))
                                    if term_with_value not in examples:
                                        examples[term_with_value] = text
                                        num_idx = text_after_term.find(value, last_idx)
                                        unit = ""
                                        if num_idx != -1:
                                            post_num = text_after_term[num_idx + len(value):].strip()
                                            unit_match = re.match(r"^\S+", post_num)
                                            unit = unit_match.group() if unit_match else ""
                                        last_idx = num_idx + len(value)
                                        lookup_key = term_with_value.lower()
                                        unit = default_settings.get(lookup_key, {}).get("unit", unit)
                                        unit = self.configure_graph.get(lookup_key, {}).get("unit", unit)
                                        term_units[term_with_value] = unit
                            break
            return data, examples, term_units, set()

        ## Extract data for GCoilTemp*.log files
        def extract_gcoiltemp(lines, checked_terms):
            data = {}
            examples = {}
            term_units = {}
            unchecked_terms = set()
            column_names = [
                "ASGC1", "ASGC2", "ASGC3", "ASGC4/RearRF", "Bore1", "Bore2",
                "WB1", "WB2", "GantryCab/PenPanel", "MagFoot", "Sequence", "Scan"
            ]
            graphable_terms = set(column_names[:-2])
            checked_terms = list(set(checked_terms + column_names))
            temp_data = {term: [] for term in graphable_terms}
            current_file = None
            for line in lines:
                if line.startswith(">>>> "):
                    current_file = line.strip(">>>> ").strip(" <<<<")
                    continue
                if not line.strip() or current_file is None:
                    continue
                parts = line.split(" ", 2)
                if len(parts) < 3:
                    continue
                date_match = re.search(r"\d{4}/\d{2}/\d{2} \d{2}:\d{2}:\d{2}", parts[1] + " " + parts[2])
                if not date_match:
                    continue
                date_str = date_match.group()
                try:
                    date = datetime.strptime(date_str, "%Y/%m/%d %H:%M:%S")
                except ValueError:
                    continue
                text = parts[2].strip()
                text_parts = text.split(" : ", 1)
                if len(text_parts) < 2:
                    continue
                data_text = text_parts[1].strip()
                values = data_text.split(" : ")
                if len(values) != len(column_names):
                    values = values + [""] * (len(column_names) - len(values)) if len(values) < len(column_names) else values[:len(column_names)]
                for i, (value, term) in enumerate(zip(values, column_names)):
                    if term not in checked_terms:
                        continue
                    if term not in examples:
                        examples[term] = f"{term} : {value}"
                    if term not in graphable_terms or not re.search(r"[-+]?\d*\.?\d+[eE]?[-+]?\d*", value):
                        continue
                    try:
                        float_value = float(value)
                    except ValueError:
                        continue
                    if term not in data:
                        data[term] = []
                    data[term].append((date, float_value, current_file))
                    if term in graphable_terms:
                        temp_data[term].append(float_value)
                    term_units[term] = "°C"
            for term in graphable_terms:
                if temp_data[term] and all(abs(v - (-273.0)) < 1e-10 for v in temp_data[term]):
                    unchecked_terms.add(term)
            return data, examples, term_units, unchecked_terms

        def extract_data(checked_terms):
            file_type_map = {
                r'\.acqsts$': extract_acqsts,
            #    r'\.sigusr1$': extract_acqsts, # removed because it doesnt' have dates. 
                r'gcoiltemp.*\.log$': extract_gcoiltemp
            }
            data = {}
            examples = {}
            term_units = {}
            unchecked_terms = set()
            current_file = None
            file_lines = {}
            unrecognized_files = []
            for line in lines:
                if line.startswith(">>>> "):
                    current_file = line.strip(">>>> ").strip(" <<<<")
                    if current_file not in file_lines:
                        file_lines[current_file] = []
                    file_lines[current_file].append(line)
                    continue
                if current_file:
                    file_lines[current_file].append(line)
            for file, file_specific_lines in file_lines.items():
                matched = False
                for pattern, extract_func in file_type_map.items():
                    if re.search(pattern, file.lower(), re.IGNORECASE):
                        file_data, file_examples, file_units, file_unchecked = extract_func(file_specific_lines, checked_terms)
                        for term in file_data:
                            if term not in data:
                                data[term] = []
                            data[term].extend(file_data[term])
                        examples.update(file_examples)
                        term_units.update(file_units)
                        unchecked_terms.update(file_unchecked)
                        matched = True
                        break
                if not matched:
                    unrecognized_files.append(file)
            if unrecognized_files:
                def show_popup():
                    popup = tk.Toplevel(graph_window)
                    popup.title("Unrecognized File Types")
                    popup.transient(graph_window)  # Tie to graph_window
                    popup.attributes('-topmost', True)  # Bring to front
                    popup.resizable(False, False)
                    file_types = {os.path.splitext(file)[1].lower() for file in unrecognized_files}
                    file_types_joined = '\n'.join(file_types)
                    message = (
                    f"The following file types were skipped due to format mismatch.\n"
                    f"    ( i.e. date missing, no values, not set up in database, etc. ):\n\n"
                    f"{file_types_joined}"
                    )
                    tk.Label(popup, text=message, justify=tk.LEFT).pack(padx=10, pady=10)
                    tk.Button(popup, text="OK", command=popup.destroy).pack(pady=5)
                    popup.update_idletasks()  # Update after widgets are added
                    width = max(popup.winfo_width(), 300)  # Ensure minimum width
                    height = popup.winfo_height() + 10  # Add padding
                    x = graph_window.winfo_x() + (graph_window.winfo_width() - width) // 2
                    y = graph_window.winfo_y() + (graph_window.winfo_height() - height) // 2
                    popup.geometry(f"{width}x{height}+{x}+{y}")
                    popup.grab_set()  # Make modal
                    popup.bind('<Button-1>', lambda event: popup.destroy())  # Close on click
                    after_id = graph_window.after(1500, popup.destroy)  # Close after 1.5 seconds
                    popup.bind('<Button-1>', lambda event: graph_window.after_cancel(after_id))  # Cancel timer on click
                graph_window.after(100, show_popup)  # Delay popup creation
            return data, examples, term_units, unchecked_terms

    ## Setup initial checked terms ----------------------
        def setup_checked_terms():
            nonlocal checked_terms, term_units, all_data, examples, check_vars, series_vars
            checked_terms = []
            term_units.clear()
            check_vars.clear()
            series_vars.clear()
            all_data, examples, term_units, unchecked_terms = extract_data(all_terms)
            for term in examples:
                # Uncheck terms in unchecked_terms
                check_vars[term] = tk.BooleanVar(value=(term in all_data and term not in unchecked_terms))
                if term in all_data and term not in unchecked_terms:
                    checked_terms.append(term)
                    series_vars[term] = tk.BooleanVar(value=True)
            update_checkbox_layout()
            update_graph()

    ## Configure graph data     # Opens with configure graph button. -------------------
        def configure_graph_data(unchecked_terms=None):
            nonlocal config_popup, checked_terms, term_units, all_data, check_vars
            if unchecked_terms is None:
                unchecked_terms = set()
            if config_popup is not None and config_popup.winfo_exists():
                config_popup.lift()
                return
            table_popup = tk.Toplevel(graph_window)
            config_popup = table_popup
            table_popup.title("Configure Data")
            table_popup.transient(graph_window)
            tk.Label(table_popup, text="Define series settings (click unit to edit, hit enter to set.):").pack(pady=5)

            main_frame = tk.Frame(table_popup)
            main_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)

            btn_frame_select = tk.Frame(main_frame)
            btn_frame_select.pack(fill=tk.X, pady=2)
            tk.Button(btn_frame_select, text="Select All", command=lambda: [check_vars[tree.item(item, "values")[1]].set(True) or tree.set(item, "Check", "☑") for item in tree.get_children()]).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame_select, text="Unselect All", command=lambda: [check_vars[tree.item(item, "values")[1]].set(False) or tree.set(item, "Check", "☐") for item in tree.get_children()]).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame_select, text="Reset Selection", command=lambda: [check_vars[tree.item(item, "values")[1]].set(term in all_data and term not in unchecked_terms) or tree.set(item, "Check", "☑" if (term in all_data and term not in unchecked_terms) else "☐") for item in tree.get_children() for term in [tree.item(item, "values")[1]]]).pack(side=tk.LEFT, padx=5)

            tree_frame = tk.Frame(main_frame)
            tree_frame.pack(fill=tk.BOTH, expand=True)
            tree = ttk.Treeview(tree_frame, columns=("Check", "Term", "Unit", "Example"), show="headings")
            tree.heading("Check", text="")
            tree.heading("Term", text="Search Term")
            tree.heading("Unit", text="Unit")
            tree.heading("Example", text="First Example")
            tree.column("Check", width=30, stretch=False, anchor="center")
            tree.column("Term", stretch=False, anchor="w")
            tree.column("Unit", stretch=False, anchor="w")
            tree.column("Example", stretch=True, anchor="w")
            y_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
            x_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
            tree.configure(yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set)
            tree.grid(row=0, column=0, sticky="nsew")
            y_scrollbar.grid(row=0, column=1, sticky="ns")
            x_scrollbar.grid(row=1, column=0, sticky="ew")
            tree_frame.grid_rowconfigure(0, weight=1)
            tree_frame.grid_columnconfigure(0, weight=1)

            pending_changes = {}  # Store edits until Apply
            sorted_terms = sorted(examples.keys(), key=str.lower)
            for term in sorted_terms:
                example = examples[term]
                base_term = term if not term.endswith(" Value") else " ".join(term.split(" ")[:-2])
                term_idx = example.lower().find(base_term.lower())
                if term_idx != -1:
                    end_idx = term_idx + len(base_term)
                    if end_idx < len(example) and example[end_idx] in ":=":
                        end_idx += 1
                        while end_idx < len(example) and example[end_idx].isspace():
                            end_idx += 1
                    original_term = example[term_idx:end_idx]
                    text_after_term = example[end_idx:].strip()
                    display_example = original_term + (" " + text_after_term if text_after_term else "")
                else:
                    original_term = term
                    text_after_term = ""
                    display_example = example
                unit = term_units.get(term, "")
                nums = re.findall(r"[-+]?\d*\.?\d+[eE]?[-+]?\d*", text_after_term)
                if term.endswith(" Value"):
                    value_idx = {"1st": 1, "2nd": 2, "3rd": 3}.get(term.split()[-2], int(term.split()[-2][:-2]))
                    if nums and value_idx <= len(nums):
                        value = nums[value_idx - 1]
                        num_idx = text_after_term.find(value)
                        if num_idx != -1:
                            text_after_term = text_after_term[:num_idx] + f"{{ {value} }}" + text_after_term[num_idx + len(value):]
                            display_example = original_term + " " + text_after_term
                elif nums:
                    value = nums[0]
                    num_idx = text_after_term.find(value)
                    if num_idx != -1:
                        text_after_term = text_after_term[:num_idx] + f"{{ {value} }}" + text_after_term[num_idx + len(value):]
                        display_example = original_term + " " + text_after_term
                # Use check_vars to set initial checkbox state
                is_checked = "☑" if check_vars.get(term, tk.BooleanVar(value=False)).get() else "☐"
                tree.insert("", "end", values=(is_checked, term, unit, display_example))

            def toggle_check(event):
                if tree.identify_region(event.x, event.y) == "cell" and tree.identify_column(event.x) == "#1":
                    item = tree.identify_row(event.y)
                    if item:
                        term = tree.item(item, "values")[1]
                        current = check_vars[term].get()
                        check_vars[term].set(not current)
                        tree.set(item, "Check", "☑" if not current else "☐")
                    return "break"

            def edit_cell(event):
                if tree.identify_region(event.x, event.y) != "cell":
                    return
                col = tree.identify_column(event.x)
                if col == "#1" or col == "#4":
                    return
                item = tree.identify_row(event.y)
                if not item:
                    return
                x, y, width, height = tree.bbox(item, col)
                entry = ttk.Entry(tree_frame)
                entry.place(x=x, y=y, width=width, height=height)
                old_value = tree.item(item, "values")[int(col[1]) - 1]
                entry.insert(0, old_value)
                entry.focus_set()

                def save_edit(*args):
                    new_value = entry.get()
                    term = tree.item(item, "values")[1]
                    tree.set(item, col, new_value)
                    entry.destroy()
                    if col == "#2":  # Term
                        old_term = term
                        if old_term in check_vars:
                            check_vars[new_value] = check_vars.pop(old_term)
                        if old_term.rstrip(':=').strip().lower() in self.configure_graph:
                            self.configure_graph[new_value.rstrip(':=').strip().lower()] = self.configure_graph.pop(old_term.rstrip(':=').strip().lower())
                        if old_term in default_settings:
                            default_settings[new_value] = default_settings.pop(old_term)
                        tree.set(item, col, new_value)
                    elif col == "#3":  # Unit
                        pending_changes[term] = {"unit": new_value}

                entry.bind("<Return>", save_edit)
                entry.bind("<FocusOut>", save_edit)

            tree.bind("<Button-1>", toggle_check)
            tree.bind("<Button-1>", edit_cell, add="+")

            tree.update_idletasks()
            for col in ("Term", "Unit"):
                max_width = len(tree.heading(col)["text"]) * 6 + 10
                for item in tree.get_children():
                    text = tree.item(item, "values")[["Check", "Term", "Unit", "Example"].index(col)]
                    width = len(text) * 6 + 20
                    max_width = max(max_width, width)
                tree.column(col, width=max_width, stretch=False)

            btn_frame = tk.Frame(table_popup)
            btn_frame.pack(pady=5)

            def apply_changes():
                nonlocal checked_terms, term_units, all_data
                for widget in tree_frame.winfo_children():
                    if isinstance(widget, ttk.Entry) and widget.focus_get() == widget:
                        item = tree.focus()
                        if item:
                            col = "#3"
                            term = tree.item(item, "values")[1]
                            new_value = widget.get()
                            tree.set(item, col, new_value)
                            if col == "#3":
                                pending_changes[term] = {"unit": new_value}
                            widget.destroy()
                        break
                checked_terms = []
                term_units.clear()
                for item in tree.get_children():
                    term = tree.item(item, "values")[1]
                    unit = tree.item(item, "values")[2]
                    if check_vars[term].get():
                        checked_terms.append(term)
                        term_units[term] = unit
                for term, settings in pending_changes.items():
                    self.configure_graph[term.rstrip(':=').strip().lower()] = settings
                pending_changes.clear()
                table_popup.destroy()
                all_data = extract_data(checked_terms)[0]
                series_lines.clear()
                series_vars.clear()
                for term in checked_terms:
                    if term in all_data:
                        series_vars[term] = tk.BooleanVar(value=True)
                update_checkbox_layout()
                update_graph()

            tk.Button(btn_frame, text="Apply", command=apply_changes).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="Cancel", command=table_popup.destroy).pack(side=tk.LEFT, padx=5)

            table_popup.update_idletasks()
            width = max(800, table_popup.winfo_width())
            height = max(600, table_popup.winfo_height())
            x = graph_window.winfo_x() + (graph_window.winfo_width() - width) // 2 + 100
            y = graph_window.winfo_y() + (graph_window.winfo_height() - height) // 2
            table_popup.geometry(f"{width}x{height}+{x}+{y}")
            
    ## Update checkbox layout ----------------------------
        def update_checkbox_layout(event=None):
                nonlocal last_update
                try:
                    # Lightweight debounce: skip if called too soon
                    if last_update and (time.time() - last_update < 0.1):
                            return
                    last_update = time.time()
                    if not all_data or not checked_terms:
                            return
                    sorted_terms = sorted(checked_terms, key=str.lower)
                    if not sorted_terms:  # Avoid empty list error
                            return

                    max_text_width = max(len(f"{term} ({term_units.get(term, self.configure_graph.get(term.rstrip(':=').strip().lower(), {'unit': ''})['unit'])})") for term in sorted_terms) * 6 + 5 #+ 30
                    window_width = max(200, graph_window.winfo_width())
                    max_columns = max(1, window_width // max_text_width)
                    for widget in checkbox_frame.winfo_children():
                            widget.destroy()
                    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
                    for i, term in enumerate(sorted_terms):
                            if term not in series_vars:
                                    series_vars[term] = tk.BooleanVar(value=True)
                            if term not in series_colors:  # Assign color once
                                    series_colors[term] = colors[i % len(colors)]
                            color = series_colors[term]  # Use persistent color
                            unit = term_units.get(term, self.configure_graph.get(term.rstrip(':=').strip().lower(), {'unit': ''})['unit'])
                            inner_frame = tk.Frame(checkbox_frame)
                            inner_frame.grid(row=i // max_columns, column=i % max_columns, padx=5, pady=2, sticky="w")
                            tk.Label(inner_frame, bg=color, width=2, height=1).pack(side=tk.LEFT, padx=(0, 5))
                            tk.Checkbutton(inner_frame, text=f"{term} ({unit})", variable=series_vars[term], command=lambda t=term: toggle_series(t)).pack(side=tk.LEFT)
                except Exception as e:
                    None

    ## Debounced resize handler
        def handle_resize(event):
                nonlocal last_window_size
                current_size = (graph_window.winfo_width(), graph_window.winfo_height())
                if current_size != last_window_size:  # Only update on size change
                        last_window_size = current_size
                        update_checkbox_layout()

        ## Filter by files # Opens with File Filter button.
        def filter_by_files(): 
            nonlocal all_data
            global file_filter_popup
            if 'file_filter_popup' in globals() and file_filter_popup is not None and file_filter_popup.winfo_exists():
                file_filter_popup.lift()
                return
            
            file_filter_popup = tk.Toplevel(graph_window)
            file_filter_popup.title("File Filter")
            file_filter_popup.transient(graph_window)
            file_filter_popup.resizable(True, True)
            
            main_frame = tk.Frame(file_filter_popup)
            main_frame.pack(padx=10, pady=5, fill=tk.BOTH, expand=True)
            main_frame.grid_rowconfigure(1, weight=1)  # Treeview expands
            main_frame.grid_rowconfigure(2, weight=0)  # Buttons fixed
            main_frame.grid_columnconfigure(0, weight=1)
            
            label = tk.Label(main_frame, text="Select files to include (multiple selections allowed):")
            label.grid(row=0, column=0, pady=5, sticky="ew")
            
            tree_frame = tk.Frame(main_frame)
            tree_frame.grid(row=1, column=0, sticky="nsew")
            tree = ttk.Treeview(tree_frame, columns=("File",), show="headings", selectmode="extended")
            tree.heading("File", text="File Name")
            tree.column("File", width=300, anchor="w")
            y_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
            tree.configure(yscrollcommand=y_scrollbar.set)
            tree.grid(row=0, column=0, sticky="nsew")
            y_scrollbar.grid(row=0, column=1, sticky="ns")
            tree_frame.grid_rowconfigure(0, weight=1)
            tree_frame.grid_columnconfigure(0, weight=1)
            
            if not hasattr(filter_by_files, 'original_file_list'):
                filter_by_files.original_file_list = sorted([f for f in files if any(f in point[2] for term in all_data for point in all_data[term])])
            file_list = filter_by_files.original_file_list
            if not hasattr(filter_by_files, 'selected_files'):
                filter_by_files.selected_files = {}
                for file in file_list:
                    filter_by_files.selected_files[file] = tk.BooleanVar(value=True)
            
            selected_files = filter_by_files.selected_files
            for file in file_list:
                if file not in selected_files:
                    selected_files[file] = tk.BooleanVar(value=True)
                tree.insert("", "end", values=(file,))
                if selected_files[file].get():
                    tree.selection_add(tree.get_children()[-1])
            
            def select_all():
                for file in file_list:
                    selected_files[file].set(True)
                tree.selection_set(tree.get_children())
            
            def clear_all():
                for file in file_list:
                    selected_files[file].set(False)
                tree.selection_remove(tree.get_children())
            
            def update_selection(event=None):
                selected_items = tree.selection()
                selected_values = [tree.item(item, "values")[0] for item in selected_items]
                for file in file_list:
                    selected_files[file].set(file in selected_values)
            
            def apply_changes():
                nonlocal all_data
                selected = [file for file, var in selected_files.items() if var.get() and file in file_list]
                if not hasattr(filter_by_files, 'original_all_data'):
                    filter_by_files.original_all_data = all_data.copy()
                filtered_data = {}
                for term in filter_by_files.original_all_data:
                    filtered_points = [(d, v, f) for d, v, f in filter_by_files.original_all_data[term] if f in selected]
                    if filtered_points:
                        filtered_data[term] = filtered_points
                if filtered_data:
                    all_data.clear()
                    all_data.update(filtered_data)
                else:
                    for term in filter_by_files.original_all_data:
                        filtered_points = [(d, v, f) for d, v, f in filter_by_files.original_all_data[term] if f in file_list]
                        if filtered_points:
                            filtered_data[term] = filtered_points
                    if filtered_data:
                        all_data.clear()
                        all_data.update(filtered_data)
                file_filter_popup.destroy()
                update_graph()
            
            def cancel():
                file_filter_popup.destroy()
            
            tree.bind("<<TreeviewSelect>>", update_selection)
            
            btn_frame = tk.Frame(main_frame)
            btn_frame.grid(row=2, column=0, pady=5, sticky="ew")
            tk.Button(btn_frame, text="Select All", command=select_all).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="Clear All", command=clear_all).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="Apply", command=apply_changes).pack(side=tk.RIGHT, padx=5)
            tk.Button(btn_frame, text="Cancel", command=cancel).pack(side=tk.RIGHT, padx=5)
            
            file_filter_popup.update_idletasks()
            # Simple height calculation
            label_height = label.winfo_reqheight()
            btn_height = btn_frame.winfo_reqheight()
            tree_row_height = 30
            padding = 30
            tree_height = len(file_list) * tree_row_height
            graph_height = graph_window.winfo_height()
            # Cap total height, prioritize buttons
            height = min(label_height + tree_height + btn_height + padding, graph_height)
            width = 400
            x = graph_window.winfo_x() + (graph_window.winfo_width() - width) // 2
            y = graph_window.winfo_y() + (graph_window.winfo_height() - height) // 2
            x = max(0, min(x, graph_window.winfo_screenwidth() - width))
            y = max(0, min(y, graph_window.winfo_screenheight() - height))
            file_filter_popup.geometry(f"{width}x{height}+{x}+{y}")
            # Set Treeview rows (scrollbar handles overflow)
            tree.configure(height=min(len(file_list), (graph_height - label_height - btn_height - padding) // tree_row_height))
            file_filter_popup.minsize(width, label_height + btn_height + 10)  # Tight minsize for buttons
            file_filter_popup.protocol("WM_DELETE_WINDOW", cancel)

    ## Toggle series visibility ----------------------------
        def toggle_series(term):
            visible = series_vars[term].get()
            if term in series_lines:  # Only toggle if line exists
                line = series_lines[term]
                line.set_visible(visible)
                ax.relim()
                ax.autoscale_view(scaley=True, scalex=False)
                update_stats()
                canvas.draw()
            update_graph()  # Always update to sync series_lines

    ## Update graph
        def update_graph():
            nonlocal annotation  # Access the outer scope's annotation
            ax.clear()
            stats_lines = {}
            series_lines.clear()  # Clear lines, but colors persist in series_colors
            start_date = datetime.strptime(start_date_entry.get(), "%Y-%m-%d")
            end_date = datetime.strptime(end_date_entry.get(), "%Y-%m-%d") + timedelta(days=1) - timedelta(seconds=1)
            visible_series = False
            for term in checked_terms:  # Only iterate over checked_terms
                if term in series_vars and series_vars[term].get() and term in all_data:
                    filtered = [(d, v, f) for d, v, f in all_data[term] if start_date <= d <= end_date]
                    if filtered:
                        # Sort by date to ensure smooth lines
                        filtered.sort(key=lambda x: x[0])
                        dates, values, _ = zip(*filtered)
                        color = series_colors.get(term, plt.rcParams['axes.prop_cycle'].by_key()['color'][0])  # Fallback to first color
                        line, = ax.plot(dates, values, marker='o', color=color, linestyle='-')
                        series_lines[term] = line
                        visible_series = True
                        if stats_var.get():
                            min_val, max_val, mean_val = min(values), max(values), sum(values) / len(values)
                            stats_lines[term] = [
                                ax.axhline(min_val, color=color, ls='--', alpha=0.5),
                                ax.axhline(max_val, color=color, ls='--', alpha=0.5),
                                ax.axhline(mean_val, color=color, ls='-', alpha=0.5)
                            ]
                            # Format annotations with up to 6 digits (5 significant, no trailing zeros)
                            def format_value(val):
                                if abs(val) >= 100:
                                    return f'{val:.2f}'.rstrip('0').rstrip('.')
                                return f'{val:.4f}'.rstrip('0').rstrip('.')[:8]
                            ax.annotate(f'{term}\nMin: \n{format_value(min_val)}', 
                                        xy=(1, min_val), xycoords=('axes fraction', 'data'),
                                        xytext=(5, 0), textcoords='offset points', color=color, va='center', ha='left')
                            ax.annotate(f'{term}\nMax: \n{format_value(max_val)}', 
                                        xy=(1, max_val), xycoords=('axes fraction', 'data'),
                                        xytext=(5, 0), textcoords='offset points', color=color, va='center', ha='left')
                            ax.annotate(f'{term}\nMean: \n{format_value(mean_val)}', 
                                        xy=(1, mean_val), xycoords=('axes fraction', 'data'),
                                        xytext=(5, 0), textcoords='offset points', color=color, va='center', ha='left')
            if visible_series:
                ax.relim()
                ax.autoscale_view(scaley=True, scalex=False)
                # Format y-axis with up to 6 digits (5 significant, no trailing zeros)
                ax.yaxis.set_major_formatter(plt.FuncFormatter(
                    lambda x, _: f'{x:.4f}'.rstrip('0').rstrip('.')[:8] if abs(x) < 100 else f'{x:.2f}'.rstrip('0').rstrip('.')
                ))
            ax.set_xlabel("Date/Time")
            ax.set_ylabel("Value")
            ax.set_title(", ".join(main_terms))  # Use all main terms for title
            plt.xticks(rotation=45)
            canvas.draw()

            # Re-apply hover tooltip after redraw
            annotation = ax.annotate("", xy=(0, 0), xytext=(10, 10), textcoords="offset points",
                                    bbox=dict(boxstyle="round,pad=0.5", fc="lightblue", alpha=0.8),
                                    arrowprops=dict(arrowstyle="->"))
            annotation.set_visible(False)
            canvas.mpl_connect("motion_notify_event", on_hover)  # Reconnect each time

    ## Update stats
        def update_stats():
            stats_lines = {}
            if stats_var.get():
                start_date = datetime.strptime(start_date_entry.get(), "%Y-%m-%d")
                end_date = datetime.strptime(end_date_entry.get(), "%Y-%m-%d") + timedelta(days=1) - timedelta(seconds=1)
                for term, points in all_data.items():
                    if term in series_vars and series_vars[term].get():
                        filtered = [(d, v, f) for d, v, f in points if start_date <= d <= end_date]
                        if filtered:
                            _, values, _ = zip(*filtered)
                            color = series_lines[term].get_color()
                            min_val, max_val, mean_val = min(values), max(values), sum(values) / len(values)
                            stats_lines[term] = [
                                ax.axhline(min_val, color=color, ls='--', alpha=0.5),
                                ax.axhline(max_val, color=color, ls='--', alpha=0.5),
                                ax.axhline(mean_val, color=color, ls='-', alpha=0.5)
                            ]
            return stats_lines

    ## On hover tooltip
        def on_hover(event):
            if event.inaxes != ax or not series_lines or not ax.get_lines():  # Skip if no lines or figure is empty
                annotation.set_visible(False)
                canvas.draw_idle()
                return
            for term, line in series_lines.items():
                if not line.get_visible():
                    continue
                xdata, ydata = line.get_xdata(), line.get_ydata()
                cont, ind = line.contains(event)
                if cont and 'ind' in ind:
                    idx = ind["ind"][0]
                    x, y = xdata[idx], ydata[idx]
                    # Get unit from term_units or configure_graph
                    unit = term_units.get(term, self.configure_graph.get(term.rstrip(':=').strip().lower(), {'unit': ''})['unit'])
                    annotation.xy = (x, y)
                    annotation.set_text(f"{term}\n{y} {unit}".strip())  # Include series name, value, and unit
                    annotation.set_visible(True)
                    canvas.draw_idle()
                    return
            annotation.set_visible(False)
            canvas.draw_idle()

    ## On quick date select --------------------------
        def on_quick_date_select(event):
            quick_date = quick_date_var.get()
            if quick_date and quick_date != "Range from End":
                end_date = end_date_entry.get_date()
                if quick_date == "1 day":
                    start_date = end_date  # Same day for 1 day
                elif quick_date == "1 week":
                    start_date = end_date - timedelta(days=6)  # 7 days inclusive
                elif quick_date == "1 month":
                    start_date = end_date - relativedelta(months=1)  # 1 month
                elif quick_date == "1 year":
                    start_date = end_date - relativedelta(years=1)  # 1 year
                elif quick_date == "5 years":
                    start_date = end_date - relativedelta(years=5)  # 5 years
                start_date_entry.set_date(start_date)
                update_graph()

        ## Setup control frame buttons and date filtering
        tk.Button(control_frame, text="Configure Data", command=configure_graph_data).pack(side=tk.LEFT, padx=5)
        tk.Button(control_frame, text="File Filter", command=filter_by_files).pack(side=tk.LEFT, padx=5)

        date_filter_frame = tk.Frame(control_frame)
        date_filter_frame.pack(side=tk.LEFT, padx=5)
        tk.Label(date_filter_frame, text="Start:").pack(side=tk.LEFT, padx=2)
        start_date_entry = tkcalendar.DateEntry(date_filter_frame, width=10, date_pattern='yyyy-mm-dd')
        start_date_entry.pack(side=tk.LEFT, padx=2)
        tk.Label(date_filter_frame, text="End:").pack(side=tk.LEFT, padx=2)
        end_date_entry = tkcalendar.DateEntry(date_filter_frame, width=10, date_pattern='yyyy-mm-dd')
        end_date_entry.pack(side=tk.LEFT, padx=2)
        quick_date_var = tk.StringVar(value="Range from End")
        quick_date_combobox = ttk.Combobox(date_filter_frame, textvariable=quick_date_var, state="readonly", width=15)
        quick_date_combobox['values'] = ["Range from End", "1 day", "1 week", "1 month", "1 year", "5 years"]
        quick_date_combobox.pack(side=tk.LEFT, padx=2)
        tk.Button(date_filter_frame, text="<", command=lambda: cycle_date(-1)).pack(side=tk.LEFT, padx=2)
        tk.Button(date_filter_frame, text=">", command=lambda: cycle_date(1)).pack(side=tk.LEFT, padx=2)
        graph_window.bind("<Left>", lambda event: cycle_date(-1))
        graph_window.bind("<Right>", lambda event: cycle_date(1))

    ## Cycle date range
        def cycle_date(direction):
                quick_date = quick_date_var.get()
                if quick_date == "Range from End":
                        return
                current_start = datetime.strptime(start_date_entry.get(), "%Y-%m-%d")
                current_end = datetime.strptime(end_date_entry.get(), "%Y-%m-%d")
                # Clamp to data range
                current_start = max(current_start, datetime.combine(min_date, datetime.min.time()))
                current_end = min(current_end, datetime.combine(max_date, datetime.min.time()))
                range_days = (current_end - current_start).days
                if quick_date == "1 day":
                        delta = timedelta(days=1)
                elif quick_date == "1 week":
                        delta = timedelta(days=7)
                elif quick_date == "1 month":
                        delta = relativedelta(months=1)
                elif quick_date == "1 year":
                        delta = relativedelta(years=1)
                elif quick_date == "5 years":
                        delta = relativedelta(years=5)
                else:
                        return
                remaining = timedelta(0)  # Default
                def show_popup(message):
                        graph_window.lift()  # Ensure graph_window is on top
                        popup = tk.Toplevel(graph_window)
                        popup.transient(graph_window)  # Tie to graph_window
                        popup.attributes('-topmost', True)  # Bring to front
                        popup.resizable(False, False)
                        tk.Label(popup, text=message, font = ("Arial", 11)).pack(padx=10, pady=5)
                        popup.update_idletasks()  # Update after widgets are added
                        # Center popup on graph_window
                        gw_x = graph_window.winfo_x()
                        gw_y = graph_window.winfo_y()
                        gw_width = graph_window.winfo_width()
                        gw_height = graph_window.winfo_height()
                        popup_width = 300
                        popup_height = 50
                        x = gw_x + (gw_width - popup_width) // 2
                        y = gw_y + gw_height - popup_height - 10
                        popup.geometry(f"{popup_width}x{popup_height}+{x}+{y}")
                        popup.lift()  # Ensure on top
                        popup.focus_set()  # Ensure focus
                        graph_window.update()  # Force Tkinter rendering
                        graph_window.after(2000, popup.destroy)  # Close after 2 seconds
                if direction > 0:  # Moving forward
                        new_end = current_end + delta
                        new_start = current_start + delta
                        skip_count = 0  # Track iterations to detect skipping
                        # Check for data in the new range
                        while new_end.date() <= max_date:
                                has_data = any(d for d in dates if new_start.date() <= d.date() <= new_end.date())
                                if has_data:
                                        if skip_count > 0:  # Only show popup if we skipped
                                                graph_window.after(200, lambda: show_popup(f"Skipping blank section to:\n {new_start.date()} to {new_end.date()}"))
                                        break
                                skip_count += 1
                                new_end = new_end + delta
                                new_start = new_start + delta
                        if new_end.date() > max_date:
                                new_end = datetime.combine(max_date, datetime.min.time())
                                new_start = new_end
                                if quick_date == "1 week":
                                        new_start = new_end - timedelta(days=6)  # Enforce 6-day range
                                        if new_start.date() < min_date:
                                                new_start = datetime.combine(min_date, datetime.min.time())
                                                new_end = new_start + timedelta(days=6)
                                                if new_end.date() > max_date:
                                                        new_end = datetime.combine(max_date, datetime.min.time())
                                                        new_start = new_end - timedelta(days=6)
                                elif quick_date not in ["1 day", "1 week"]:  # Non-day/week ranges
                                        if 0 < remaining.days < delta.days:  # Partial range
                                                new_start = datetime.combine(max_date - remaining, datetime.min.time())
                                        else:
                                                new_start = new_end - (current_end - current_start)  # Preserve range
                                remaining = max_date - current_end.date()
                                graph_window.after(200, lambda: show_popup("You have reached the end of the data set"))
                                # Update UI even if no data to prevent stalling
                                if new_start != current_start or new_end != current_end:
                                        start_date_entry.set_date(new_start)
                                        end_date_entry.set_date(new_end)
                                        update_graph()
                        else:
                                remaining = max(max_date - new_end.date(), timedelta(0))
                else:  # Moving backward
                        new_start = current_start - delta
                        new_end = current_end - delta
                        skip_count = 0  # Track iterations to detect skipping
                        # Check for data in the new range
                        while new_start.date() >= min_date:
                                has_data = any(d for d in dates if new_start.date() <= d.date() <= new_end.date())
                                if has_data:
                                        if skip_count > 0:  # Only show popup if we skipped
                                                graph_window.after(200, lambda: show_popup(f"Skipping blank section to:\n {new_start.date()} to {new_end.date()}"))
                                        break
                                skip_count += 1
                                new_start = new_start - delta
                                new_end = new_end - delta
                        if new_start.date() < min_date:
                                new_start = datetime.combine(min_date, datetime.min.time())
                                new_end = new_start
                                if quick_date == "1 week":
                                        new_end = new_start + timedelta(days=6)  # Enforce 6-day range
                                        if new_end.date() > max_date:
                                                new_end = datetime.combine(max_date, datetime.min.time())
                                                new_start = new_end - timedelta(days=6)
                                elif quick_date not in ["1 day", "1 week"]:  # Non-day/week ranges
                                        if 0 < remaining.days < delta.days:  # Partial range
                                                new_end = datetime.combine(min_date + remaining, datetime.min.time())
                                        else:
                                                new_end = min(new_start + (current_end - current_start), datetime.combine(max_date, datetime.min.time()))  # Preserve range
                                remaining = current_start.date() - min_date
                                graph_window.after(200, lambda: show_popup("You have reached the end of the data set"))
                                # Update UI even if no data to prevent stalling
                                if new_start != current_start or new_end != current_end:
                                        start_date_entry.set_date(new_start)
                                        end_date_entry.set_date(new_end)
                                        update_graph()
                        else:
                                remaining = timedelta(0)  # No boundary
                if (new_start.date() >= min_date and new_end.date() <= max_date and
                    new_start <= new_end and
                    (new_start != current_start or new_end != current_end)):
                        start_date_entry.set_date(new_start)
                        end_date_entry.set_date(new_end)
                        update_graph()



                        
        if dates:
            min_date = min(dates).date()
            max_date = max(dates).date()
        else:
            min_date = datetime.today().date() - timedelta(days=30)
            max_date = datetime.today().date()
        start_date_entry.set_date(min_date)
        end_date_entry.set_date(max_date)

        initial_start_date = min_date
        initial_end_date = max_date
        tk.Button(control_frame, text="Reset Dates", command=lambda: [start_date_entry.set_date(initial_start_date), end_date_entry.set_date(initial_end_date), quick_date_var.set("Range from End"), update_graph()]).pack(side=tk.LEFT, padx=5)
          
        stats_var = tk.BooleanVar(value=False)
        tk.Checkbutton(control_frame, text="Min/Max/Mean", variable=stats_var, command=lambda: update_graph()).pack(side=tk.LEFT, padx=5)

        tk.Button(control_frame, text="Select All", command=lambda: [var.set(True) for var in series_vars.values()] and update_graph()).pack(side=tk.LEFT, padx=5)
        tk.Button(control_frame, text="Clear All", command=lambda: [var.set(False) for var in series_vars.values()] and update_graph()).pack(side=tk.LEFT, padx=5)

    ## Bind events
        quick_date_combobox.bind("<<ComboboxSelected>>", on_quick_date_select)
        start_date_entry.bind("<<DateEntrySelected>>", lambda e: [quick_date_var.set("Range from End"), update_graph()])
        end_date_entry.bind("<<DateEntrySelected>>", lambda e: [quick_date_var.set("Range from End"), update_graph()])
        graph_window.bind("<Configure>", handle_resize)  # Bind debounced resize handler

    ## Initial setup and graph population
        setup_checked_terms()  # Run once to populate initial data
        graph_window.update_idletasks()  # Ensure window geometry is processed before layout
        update_graph()  # Initial graph draw














#### Tool Menu Control and Options. ------------------------------------------------------------------

    def update_options_dropdown(self):
        state_indicator = " [ON] (Click Off)" if self.save_prefs_var.get() else " [OFF] (Click On)"
        self.options_dropdown['values'] = [
            "Instructions",
            f"Store Settings {state_indicator}",
            "7-Zip Path Setup",
            "Manage Folder/File Paths",
            "Search FIXS Setup",
            "Search Intranet Setup",
            "Search Windows Setup"
        ]
  
    def on_options_select(self, event): 
        selected = self.options_dropdown.get()
        self.options_var.set("               Tool Menu")
        if selected == "Instructions":
            self.show_instructions()
        elif selected.startswith("Store Settings"):
            current_state = self.save_prefs_var.get()
            if not current_state and not os.path.exists(self.config_path):  # Turning on and JSON doesn't exist
                confirm_popup = tk.Toplevel(self.master)
                confirm_popup.transient(self.master)
                confirm_popup.title("Confirm Store Settings")
                tk.Label(confirm_popup, text="You are turning on Store Settings.\nThis will save your preferences to log_scraper_settings.json\nfor future use. Would you like to proceed?", justify=tk.CENTER).pack(padx=10, pady=10)
                
                def on_yes():
                    self.save_prefs_var.set(True)
                    self.update_options_dropdown()
                    self.preferences['save_entries'] = True
                    with open(self.config_path, 'w') as file:
                        json.dump(self.preferences, file, indent=4)
                    confirm_popup.destroy()
                    self.show_state_popup("is Turned On\nand will save settings\nto log_scraper_settings.json")
                
                def on_no():
                    confirm_popup.destroy()
                
                # Button frame for centered layout
                button_frame = tk.Frame(confirm_popup)
                button_frame.pack(pady=5)
                tk.Button(button_frame, text="Yes", command=on_yes).pack(side=tk.LEFT, padx=(0, 25))  # 25px spacing on right
                tk.Button(button_frame, text="No", command=on_no).pack(side=tk.LEFT)  # No extra padding on right
                
                # Center confirm popup
                confirm_popup.update_idletasks()
                width = confirm_popup.winfo_width()
                height = confirm_popup.winfo_height()
                x = self.master.winfo_x() + (self.master.winfo_width() - width) // 2
                y = self.master.winfo_y() + (self.master.winfo_height() - height) // 2
                confirm_popup.geometry(f"{width}x{height}+{x}+{y}")
            else:
                # Toggle as usual if JSON exists or turning off
                self.save_prefs_var.set(not current_state)
                self.update_options_dropdown()
                new_state_popup = "is Turned On\nand will save settings\nto log_scraper_settings.json" if self.save_prefs_var.get() else "is Turned Off\nand will NOT save settings\nto log_scraper_settings.json"
                self.preferences['save_entries'] = self.save_prefs_var.get()
                with open(self.config_path, 'w') as file:
                    json.dump(self.preferences, file, indent=4)
                self.show_state_popup(new_state_popup)
                
        elif selected == "Search FIXS Setup":
            if self.fixs_setup_window and self.fixs_setup_window.winfo_exists():
                self.fixs_setup_window.lift()
                return
            save_state_msg = "Store Settings is already ON." if self.save_prefs_var.get() else "Store Settings is OFF. Clicking save will turn it on and create a log_scraper_setting.json file."
            setup_window = tk.Toplevel(self.master)
            self.fixs_setup_window = setup_window
            setup_window.title("Search FIXS Setup")
            setup_window.transient(self.master)
            setup_window.protocol("WM_DELETE_WINDOW", lambda: self.close_window('fixs'))
            tk.Label(setup_window, text=f"Enter the base web address for FIXS search.\nUse \"\" \"\" where the highlighted text should go.\n{save_state_msg}", 
                    justify=tk.CENTER, wraplength=300).grid(row=0, column=0, columnspan=2, pady=10)
            url_var = tk.StringVar(value=self.fixs_base_url)
            url_entry = self.create_context_menu(tk.Entry(setup_window, textvariable=url_var))  # Add context menu
            url_entry.grid(row=1, column=0, columnspan=2, pady=5, padx=10, sticky="ew")
            def adjust_width(*args):
                content_length = len(url_var.get())
                screen_width = setup_window.winfo_screenwidth()
                max_width = (screen_width - 40) // 6
                new_width = min(content_length + 2, max_width) or 10
                url_entry.config(width=new_width)
            url_var.trace_add("write", adjust_width)
            setup_window.update_idletasks()
            adjust_width()
            def save_url():
                new_url = url_var.get().strip()
                if new_url and '"" ""' in new_url:
                    self.fixs_base_url = new_url
                    if not self.save_prefs_var.get():
                        self.save_prefs_var.set(True)
                        self.update_options_dropdown()
                    self.preferences['save_entries'] = self.save_prefs_var.get()
                    self.preferences['fixs_base_url'] = self.fixs_base_url
                    with open(self.config_path, 'w') as file:
                        json.dump(self.preferences, file, indent=4)
                    self.close_window('fixs')
                else:
                    messagebox.showerror("Error", "URL must contain \"\" \"\" for text insertion!", parent=setup_window)
            def set_default():
                url_var.set(self.fixs_base_url)
                if self.save_prefs_var.get():
                    self.preferences['fixs_base_url'] = self.fixs_base_url
                    self.save_preferences()
            button_frame = tk.Frame(setup_window)
            button_frame.grid(row=2, column=0, columnspan=2, pady=10)
            tk.Button(button_frame, text="Save", command=save_url).pack(side=tk.LEFT, padx=5)
            tk.Button(button_frame, text="Set Default", command=set_default).pack(side=tk.LEFT, padx=5)
            tk.Button(button_frame, text="Close", command=lambda: self.close_window('fixs')).pack(side=tk.LEFT, padx=20)
            setup_window.columnconfigure(0, weight=1)
            setup_window.columnconfigure(1, weight=1)
            setup_window.update_idletasks()
            width = setup_window.winfo_reqwidth()
            height = setup_window.winfo_reqheight()
            x = self.master.winfo_x() + (self.master.winfo_width() - width) // 2
            y = self.master.winfo_y() + 70
            screen_width = setup_window.winfo_screenwidth()
            if x < 0: x = 0
            elif x + width > screen_width: x = screen_width - width
            setup_window.geometry(f"{width}x{height}+{x}+{y}")
        
        elif selected == "Search Intranet Setup":
            if self.intranet_setup_window and self.intranet_setup_window.winfo_exists():
                self.intranet_setup_window.lift()
                return
            save_state_msg = "Store Settings is already ON." if self.save_prefs_var.get() else "Store Settings is OFF. Clicking save will turn it on and create a log_scraper_setting.json file."
            setup_window = tk.Toplevel(self.master)
            self.intranet_setup_window = setup_window
            setup_window.title("Search Intranet Setup")
            setup_window.transient(self.master)
            setup_window.protocol("WM_DELETE_WINDOW", lambda: self.close_window('intranet'))
            tk.Label(setup_window, text=f"Enter the base web address for Intranet search.\nUse \"\" \"\" where the highlighted text should go.\n{save_state_msg}", 
                    justify=tk.CENTER, wraplength=300).grid(row=0, column=0, columnspan=2, pady=10)
            url_var = tk.StringVar(value=self.intranet_base_url)
            url_entry = self.create_context_menu(tk.Entry(setup_window, textvariable=url_var))  # Add context menu
            url_entry.grid(row=1, column=0, columnspan=2, pady=5, padx=10, sticky="ew")
            def adjust_width(*args):
                content_length = len(url_var.get())
                screen_width = setup_window.winfo_screenwidth()
                max_width = (screen_width - 40) // 6
                new_width = min(content_length + 2, max_width) or 10
                url_entry.config(width=new_width)
            url_var.trace_add("write", adjust_width)
            setup_window.update_idletasks()
            adjust_width()
            def save_url():
                new_url = url_var.get().strip()
                if new_url and '"" ""' in new_url:
                    self.intranet_base_url = new_url
                    if not self.save_prefs_var.get():
                        self.save_prefs_var.set(True)
                        self.update_options_dropdown()
                    self.preferences['save_entries'] = self.save_prefs_var.get()
                    self.preferences['intranet_base_url'] = self.intranet_base_url
                    with open(self.config_path, 'w') as file:
                        json.dump(self.preferences, file, indent=4)
                    self.close_window('intranet')
                else:
                    messagebox.showerror("Error", "URL must contain \"\" \"\" for text insertion!", parent=setup_window)
            def set_default():
                url_var.set(self.intranet_base_url)
                if self.save_prefs_var.get():
                    self.preferences['intranet_base_url'] = self.intranet_base_url
                    self.save_preferences()
            button_frame = tk.Frame(setup_window)
            button_frame.grid(row=2, column=0, columnspan=2, pady=10)
            tk.Button(button_frame, text="Save", command=save_url).pack(side=tk.LEFT, padx=5)
            tk.Button(button_frame, text="Set Default", command=set_default).pack(side=tk.LEFT, padx=5)
            tk.Button(button_frame, text="Close", command=lambda: self.close_window('intranet')).pack(side=tk.LEFT, padx=20)
            setup_window.columnconfigure(0, weight=1)
            setup_window.columnconfigure(1, weight=1)
            setup_window.update_idletasks()
            width = setup_window.winfo_reqwidth()
            height = setup_window.winfo_reqheight()
            x = self.master.winfo_x() + (self.master.winfo_width() - width) // 2
            y = self.master.winfo_y() + 70
            screen_width = setup_window.winfo_screenwidth()
            if x < 0: x = 0
            elif x + width > screen_width: x = screen_width - width
            setup_window.geometry(f"{width}x{height}+{x}+{y}")
        
        elif selected == "Search Windows Setup":
            if self.windows_setup_window and self.windows_setup_window.winfo_exists():
                self.windows_setup_window.lift()
                return
            def move_up():
                selected = location_listbox.curselection()
                if selected:
                    index = selected[0]
                    if index > 0:
                        text = location_listbox.get(index)
                        location_listbox.delete(index)
                        location_listbox.insert(index - 1, text)
                        location_listbox.selection_set(index - 1)
                        self.search_location_order.insert(index - 1, self.search_location_order.pop(index))
                        if self.save_prefs_var.get():
                            self.save_preferences()
            def move_down():
                selected = location_listbox.curselection()
                if selected:
                    index = selected[0]
                    if index < location_listbox.size() - 1:
                        text = location_listbox.get(index)
                        location_listbox.delete(index)
                        location_listbox.insert(index + 1, text)
                        location_listbox.selection_set(index + 1)
                        self.search_location_order.insert(index + 1, self.search_location_order.pop(index))
                        if self.save_prefs_var.get():
                            self.save_preferences()
            def sort_az():
                self.search_location_order.sort()
                refresh_locations()
                if self.save_prefs_var.get():
                    self.save_preferences()
            def add_or_update_location():
                name = name_entry.get().strip()
                path = path_entry.get().strip()
                if not name or not path or not os.path.exists(path):
                    messagebox.showerror("Error", "Name and a valid Path are required.", parent=setup_window)
                    return
                if name in self.search_locations and name != name_entry.original_name:
                    messagebox.showerror("Error", f"Name '{name}' already exists.", parent=setup_window)
                    return
                if len(self.search_locations) >= 5 and name_entry.original_name not in self.search_locations:
                    messagebox.showerror("Error", "Maximum of 5 locations reached.", parent=setup_window)
                    return
                if name_entry.original_name and name_entry.original_name in self.search_location_order:
                    index = self.search_location_order.index(name_entry.original_name)
                    self.search_location_order[index] = name
                    if name_entry.original_name in self.search_locations:
                        del self.search_locations[name_entry.original_name]
                elif name not in self.search_location_order:
                    self.search_location_order.append(name)
                self.search_locations[name] = path
                refresh_locations()
                if self.save_prefs_var.get():
                    self.save_preferences()
                name_entry.original_name = name
            def delete_location():
                selected = location_listbox.curselection()
                if not selected:
                    messagebox.showerror("Error", "Please select a location to delete.", parent=setup_window)
                    return
                name = location_listbox.get(selected[0]).split(":")[0].strip()
                popup = tk.Toplevel(setup_window)
                popup.transient(setup_window)
                popup.title("Confirm Deletion")
                tk.Label(popup, text=f"Delete '{name}'?").pack(padx=20, pady=10)
                def confirm():
                    if name in self.search_locations:
                        del self.search_locations[name]
                        self.search_location_order.remove(name)
                    refresh_locations()
                    if self.save_prefs_var.get():
                        self.save_preferences()
                    popup.destroy()
                def cancel():
                    popup.destroy()
                button_frame = tk.Frame(popup)
                button_frame.pack(pady=5)
                tk.Button(button_frame, text="Yes", command=confirm).grid(row=0, column=0, padx=10)
                tk.Button(button_frame, text="No", command=cancel).grid(row=0, column=1, padx=10)
                popup.update_idletasks()
                width = popup.winfo_width()
                height = popup.winfo_height()
                x = setup_window.winfo_x() + (setup_window.winfo_width() - width) // 2
                y = setup_window.winfo_y() + (setup_window.winfo_height() - height) // 2
                popup.geometry(f"{width}x{height}+{x}+{y}")
                popup.grab_set()
            def clear_fields():
                name_entry.delete(0, tk.END)
                name_entry.original_name = ""
                path_entry.delete(0, tk.END)
            def refresh_locations():
                location_listbox.delete(0, tk.END)
                for name in self.search_location_order:
                    if name in self.search_locations:
                        location_listbox.insert(tk.END, f"{name}: {self.search_locations[name]}")
            def on_location_select(event):
                selected = location_listbox.curselection()
                if not selected:
                    return
                item = location_listbox.get(selected[0])
                name, path = item.split(":", 1)
                name_entry.delete(0, tk.END)
                name_entry.insert(0, name.strip())
                name_entry.original_name = name.strip()
                path_entry.delete(0, tk.END)
                path_entry.insert(0, path.strip())
            def toggle_save_settings():
                if save_var.get() and not self.save_prefs_var.get():
                    if messagebox.askyesno("Enable Store Settings", "Enabling Store Settings will save changes to log_scraper_settings.json. Proceed?", parent=setup_window):
                        self.save_prefs_var.set(True)
                        self.update_options_dropdown()
                        self.save_preferences()
                    else:
                        save_var.set(False)
                elif not save_var.get():
                    self.save_prefs_var.set(False)
                    self.update_options_dropdown()
            def browse_path():
                folder = filedialog.askdirectory(initialdir=path_entry.get() or "C:/")
                if folder:
                    name_entry.delete(0, tk.END)
                    name_entry.original_name = ""
                    path_entry.delete(0, tk.END)
                    path_entry.insert(0, folder)
            setup_window = tk.Toplevel(self.master)
            self.windows_setup_window = setup_window
            setup_window.title("Search Windows Setup")
            setup_window.transient(self.master)
            setup_window.protocol("WM_DELETE_WINDOW", lambda: self.close_window('windows'))
            x = self.master.winfo_x() + 50
            y = self.master.winfo_y() + 70
            screen_width = setup_window.winfo_screenwidth()
            if x < 0: x = 0
            elif x + 400 > screen_width: x = screen_width - 400
            setup_window.geometry(f"400x400+{x}+{y}")
            setup_window.grid_rowconfigure(1, weight=1)
            setup_window.grid_columnconfigure(1, weight=1)
            tk.Label(setup_window, text="Existing Search Locations:").grid(row=0, column=0, columnspan=2, sticky="w", padx=5, pady=5)
            listbox_frame = tk.Frame(setup_window)
            listbox_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)
            listbox_frame.grid_rowconfigure(0, weight=1)
            listbox_frame.grid_columnconfigure(0, weight=1)
            location_listbox = tk.Listbox(listbox_frame, width=80, height=5)
            location_listbox.grid(row=0, column=0, sticky="nsew")
            scrollbar = tk.Scrollbar(listbox_frame, orient="vertical", command=location_listbox.yview)
            scrollbar.grid(row=0, column=1, sticky="ns")
            location_listbox.config(yscrollcommand=scrollbar.set)
            location_listbox.bind("<<ListboxSelect>>", on_location_select)
            refresh_locations()
            sort_frame = tk.Frame(setup_window)
            sort_frame.grid(row=2, column=0, columnspan=2, sticky="w", padx=5, pady=5)
            tk.Label(sort_frame, text="Sort Locations:").grid(row=0, column=0, sticky="w", padx=(0, 10))
            tk.Button(sort_frame, text="▲", command=move_up, width=3).grid(row=0, column=1, padx=2)
            tk.Button(sort_frame, text="▼", command=move_down, width=3).grid(row=0, column=2, padx=2)
            tk.Button(sort_frame, text="A-Z", command=sort_az, width=3).grid(row=0, column=3, padx=2)
            tk.Label(setup_window, text=" ").grid(row=3, column=0, columnspan=2, pady=5)
            tk.Label(setup_window, text="Name:").grid(row=4, column=0, sticky="e", padx=5, pady=2)
            name_entry = self.create_context_menu(tk.Entry(setup_window))
            name_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=2)
            name_entry.original_name = ""
            tk.Label(setup_window, text="Path:").grid(row=5, column=0, sticky="e", padx=5, pady=2)
            path_entry = self.create_context_menu(tk.Entry(setup_window))
            path_entry.grid(row=5, column=1, sticky="ew", padx=5, pady=2)
            save_var = tk.BooleanVar(value=self.save_prefs_var.get())
            tk.Checkbutton(setup_window, text="Store Settings", variable=save_var, command=toggle_save_settings).grid(row=6, column=0, sticky="w", padx=5, pady=5)
            tk.Button(setup_window, text="Browse", width=10, command=browse_path).grid(row=6, column=1, padx=50, pady=2, sticky="w")
            button_frame = tk.Frame(setup_window)
            button_frame.grid(row=7, column=0, columnspan=2, pady=10, sticky="w")
            tk.Button(button_frame, text="Delete Selected", command=delete_location).grid(row=0, column=0, padx=5)
            tk.Button(button_frame, text="Clear Fields", command=clear_fields).grid(row=0, column=1, padx=5)
            tk.Button(button_frame, text="Add/Update Location", command=add_or_update_location).grid(row=0, column=2, padx=5)
            tk.Button(button_frame, text="Close", command=lambda: self.close_window('windows')).grid(row=0, column=3, padx=20)
            setup_window.bind('<Return>', lambda event: add_or_update_location())
         
        elif selected == "7-Zip Path Setup":
            if hasattr(self, 'seven_zip_setup_window') and self.seven_zip_setup_window and self.seven_zip_setup_window.winfo_exists():
                self.seven_zip_setup_window.lift()
                return
            save_state_msg = "Store Settings is already ON." if self.save_prefs_var.get() else "Store Settings is OFF. Clicking save will turn it on."
            setup_window = tk.Toplevel(self.master)
            self.seven_zip_setup_window = setup_window
            setup_window.title("7-Zip Path Setup")
            setup_window.transient(self.master)
            setup_window.protocol("WM_DELETE_WINDOW", lambda: self.close_window('seven_zip'))
            tk.Label(setup_window, text=f"Enter the path to 7z.exe.\n{save_state_msg}", 
                    justify=tk.CENTER, wraplength=300).grid(row=0, column=0, columnspan=2, pady=10)
            path_var = tk.StringVar(value=self.seven_zip_path)
            path_entry = self.create_context_menu(tk.Entry(setup_window, textvariable=path_var))  # Add context menu
            path_entry.grid(row=1, column=0, columnspan=2, pady=5, padx=10, sticky="ew")
            def adjust_width(*args):
                content_length = len(path_var.get())
                screen_width = setup_window.winfo_screenwidth()
                max_width = (screen_width - 40) // 6
                new_width = min(content_length + 2, max_width) or 10
                path_entry.config(width=new_width)
            path_var.trace_add("write", adjust_width)
            setup_window.update_idletasks()
            adjust_width()
            def save_path():
                new_path = path_var.get().strip()
                if new_path and os.path.exists(new_path) and new_path.endswith("7z.exe"):
                    self.seven_zip_path = new_path
                    if not self.save_prefs_var.get():
                        self.save_prefs_var.set(True)
                        self.update_options_dropdown()
                    self.save_preferences()
                    self.close_window('seven_zip')
                else:
                    messagebox.showerror("Error", "Path must point to a valid 7z.exe!", parent=setup_window)
            def set_default():
                path_var.set(DEFAULT_7Z_PATH)
                self.seven_zip_path = DEFAULT_7Z_PATH
                if self.save_prefs_var.get():
                    self.save_preferences()
            button_frame = tk.Frame(setup_window)
            button_frame.grid(row=2, column=0, columnspan=2, pady=10)
            tk.Button(button_frame, text="Save", command=save_path).pack(side=tk.LEFT, padx=5)
            tk.Button(button_frame, text="Set Default", command=set_default).pack(side=tk.LEFT, padx=5)
            tk.Button(button_frame, text="Close", command=lambda: self.close_window('seven_zip')).pack(side=tk.LEFT, padx=20)
            setup_window.columnconfigure(0, weight=1)
            setup_window.columnconfigure(1, weight=1)
            setup_window.update_idletasks()
            width = setup_window.winfo_reqwidth()
            height = setup_window.winfo_reqheight()
            x = self.master.winfo_x() + (self.master.winfo_width() - width) // 2
            y = self.master.winfo_y() + 70
            screen_width = setup_window.winfo_screenwidth()
            if x < 0: x = 0
            elif x + width > screen_width: x = screen_width - width
            setup_window.geometry(f"{width}x{height}+{x}+{y}")
  
        elif selected == "Manage Folder/File Paths":
            if hasattr(self, 'manage_paths_window') and self.manage_paths_window and self.manage_paths_window.winfo_exists():
                self.manage_paths_window.lift()
                return
            def move_up():
                selected = path_listbox.curselection()
                if selected:
                    index = selected[0]
                    if index > 2:  # Below separator
                        actual_index = index - 2
                        text = path_listbox.get(index)
                        path_listbox.delete(index)
                        path_listbox.insert(index - 1, text)
                        # Swap with item above
                        self.search_path_order[actual_index], self.search_path_order[actual_index - 1] = \
                            self.search_path_order[actual_index - 1], self.search_path_order[actual_index]
                        if self.save_prefs_var.get():
                            self.save_preferences()
                        self.refresh_paths()
                        # Reapply selection after refresh
                        path_listbox.selection_set(index - 1)
                        path_listbox.activate(index - 1)
                        path_listbox.focus_set()
                        self.update_path_combo()

            def move_down():
                selected = path_listbox.curselection()
                if selected:
                    index = selected[0]
                    if index > 1 and index < path_listbox.size() - 1:
                        actual_index = index - 2
                        text = path_listbox.get(index)
                        path_listbox.delete(index)
                        path_listbox.insert(index + 1, text)
                        # Swap with item below
                        self.search_path_order[actual_index], self.search_path_order[actual_index + 1] = \
                            self.search_path_order[actual_index + 1], self.search_path_order[actual_index]
                        if self.save_prefs_var.get():
                            self.save_preferences()
                        self.refresh_paths()
                        # Reapply selection after refresh
                        path_listbox.selection_set(index + 1)
                        path_listbox.activate(index + 1)
                        path_listbox.focus_set()
                        self.update_path_combo()
            def sort_az():
                # Sort the entire search_path_order: non-Recent A-Z first, then Recent
                self.search_path_order.sort(key=lambda x: (x.startswith("Recent "), x.lower() if not x.startswith("Recent ") else x))
                self.refresh_paths()
                if self.save_prefs_var.get():
                    self.save_preferences()
                self.update_path_combo()
            
            
            def add_or_update_path():
                name = name_entry.get().strip()
                path = path_entry.get().strip()
                if not path:
                    messagebox.showerror("Error", "Path is required.", parent=setup_window)
                    return
                if not os.path.exists(path):
                    messagebox.showerror("Error", "Path must exist.", parent=setup_window)
                    return
                if name and name in self.search_paths and name != name_entry.original_name:
                    messagebox.showerror("Error", f"Name '{name}' already exists.", parent=setup_window)
                    return
                effective_name = name if name else path
                if name_entry.original_name and name_entry.original_name in self.search_path_order:
                    index = self.search_path_order.index(name_entry.original_name)
                    self.search_path_order[index] = effective_name
                    if name_entry.original_name in self.search_paths:
                        del self.search_paths[name_entry.original_name]
                elif effective_name not in self.search_path_order:
                    self.search_path_order.append(effective_name)
                self.search_paths[effective_name] = path
                self.refresh_paths()
                self.update_path_combo()
                if self.save_prefs_var.get():
                    self.save_preferences()
                name_entry.original_name = effective_name
            def delete_path():
                selected = path_listbox.curselection()
                if not selected:
                    messagebox.showerror("Error", "Select a path to delete.", parent=setup_window)
                    return
                index = selected[0]
                item = path_listbox.get(index)
                if item.startswith("(Current)"):
                    messagebox.showinfo("Info", "Cannot delete the current path.", parent=setup_window)
                    return
                if item == "-" * 20:
                    return  # Skip separator
                name = item.split(": ", 1)[0].strip()
                popup = tk.Toplevel(setup_window)
                popup.transient(setup_window)
                popup.title("Confirm Deletion")
                tk.Label(popup, text=f"Delete '{name}'?").pack(padx=20, pady=10)
                def confirm():
                    if name in self.search_paths:
                        del self.search_paths[name]
                        self.search_path_order.remove(name)
                        recent_paths = [(n, self.search_paths[n]) for n in self.search_path_order if n.startswith("Recent ")]
                        for n, _ in recent_paths:
                            self.search_path_order.remove(n)
                            del self.search_paths[n]
                        for i, (_, path) in enumerate(recent_paths, 1):
                            new_name = f"Recent {i}"
                            # Insert below separator (index 2+)
                            self.search_path_order.insert(min(len(self.search_path_order), 2 + i - 1), new_name)
                            self.search_paths[new_name] = path
                    self.refresh_paths()
                    self.update_path_combo()
                    if self.save_prefs_var.get():
                        self.save_preferences()
                    popup.destroy()
                def cancel():
                    popup.destroy()
                button_frame = tk.Frame(popup)
                button_frame.pack(pady=5)
                tk.Button(button_frame, text="Yes", command=confirm).grid(row=0, column=0, padx=10)
                tk.Button(button_frame, text="No", command=cancel).grid(row=0, column=1, padx=10)
                popup.update_idletasks()
                width = popup.winfo_width()
                height = popup.winfo_height()
                x = setup_window.winfo_x() + (setup_window.winfo_width() - width) // 2
                y = setup_window.winfo_y() + (setup_window.winfo_height() - height) // 2
                popup.geometry(f"{width}x{height}+{x}+{y}")
                popup.grab_set()
            def clear_fields():
                name_entry.delete(0, tk.END)
                name_entry.original_name = ""
                path_entry.delete(0, tk.END)
            def on_path_select(event):
                selected = path_listbox.curselection()
                if not selected:
                    return
                index = selected[0]
                item = path_listbox.get(index)
                if item == "-" * 20:
                    clear_fields()
                    return
                name_part = item.split(": ", 1)[0].strip()
                path = item.split(": ", 1)[1] if ": " in item else ""
                if name_part == "(Current)" or name_part.startswith("Recent "):
                    name = ""
                    path = self.path_var.get().strip() if name_part == "(Current)" else self.search_paths.get(name_part, "")
                else:
                    name = name_part
                    path = self.search_paths.get(name_part, "")
                name_entry.delete(0, tk.END)
                name_entry.insert(0, name)
                name_entry.original_name = name
                path_entry.delete(0, tk.END)
                path_entry.insert(0, path)
            def toggle_save_settings():
                if save_var.get() and not self.save_prefs_var.get():
                    if messagebox.askyesno("Enable Store Settings", "Enabling Store Settings will save changes to log_scraper_settings.json. Proceed?", parent=setup_window):
                        self.save_prefs_var.set(True)
                        self.update_options_dropdown()
                        self.save_preferences()
                    else:
                        save_var.set(False)
                elif not save_var.get():
                    self.save_prefs_var.set(False)
                    self.update_options_dropdown()
            def update_recent_count(*args):
                new_count = recent_count_var.get()
                old_count = len([n for n in self.search_path_order if n.startswith("Recent ")])
                
                # Get current recent paths and their positions
                recent_items = [(i, n, self.search_paths.get(n)) for i, n in enumerate(self.search_path_order) if n.startswith("Recent ")]
                recent_paths = {n: p for _, n, p in recent_items if p}
                
                if new_count > old_count:
                    # Add new Recent slots at the end of the current Recent sequence
                    last_recent_index = max([i for i, n, _ in recent_items] + [-1]) if recent_items else len(self.search_path_order) - 1
                    for i in range(old_count + 1, new_count + 1):
                        new_name = f"Recent {i}"
                        if new_name not in self.search_path_order:
                            self.search_path_order.insert(last_recent_index + 1, new_name)
                            last_recent_index += 1
                elif new_count < old_count:
                    # Remove excess Recent slots from the end, preserving order
                    for i in range(old_count, new_count, -1):
                        name = f"Recent {i}"
                        if name in self.search_path_order:
                            self.search_path_order.remove(name)
                            if name in self.search_paths:
                                del self.search_paths[name]

                # Update paths for remaining Recent slots
                for i, name in enumerate([n for n in self.search_path_order if n.startswith("Recent ")]):
                    if i < len(recent_paths):
                        self.search_paths[name] = list(recent_paths.values())[i]
                    else:
                        self.search_paths[name] = None

                self.refresh_paths()
                self.update_path_combo()
                if self.save_prefs_var.get():
                    self.preferences['recent_count'] = new_count
                    self.save_preferences()

            def select_folder():
                folder = filedialog.askdirectory(initialdir=path_entry.get() or "C:/")
                if folder:
                    path_entry.delete(0, tk.END)
                    path_entry.insert(0, folder)
                    name_entry.delete(0, tk.END)
                    name_entry.original_name = ""

            def select_file():
                file = filedialog.askopenfilename(initialdir=path_entry.get() or "C:/")
                if file:
                    path_entry.delete(0, tk.END)
                    path_entry.insert(0, file)
                    name_entry.delete(0, tk.END)
                    name_entry.original_name = ""

            setup_window = tk.Toplevel(self.master)
            self.manage_paths_window = setup_window
            setup_window.title("Manage Folder/File Paths")
            setup_window.transient(self.master)
            setup_window.protocol("WM_DELETE_WINDOW", lambda: self.close_window('paths'))
            setup_window.grid_rowconfigure(1, weight=1)
            setup_window.grid_columnconfigure(1, weight=1)
            x = self.master.winfo_x() + 50
            y = self.master.winfo_y() + 70
            screen_width = setup_window.winfo_screenwidth()
            if x < 0: x = 0
            elif x + 400 > screen_width: x = screen_width - 400
            setup_window.geometry(f"600x450+{x}+{y}")
            tk.Label(setup_window, text="Existing Folder/File Search Paths:").grid(row=0, column=0, columnspan=2, sticky="w", padx=5, pady=5)
            listbox_frame = tk.Frame(setup_window)
            listbox_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)
            listbox_frame.grid_rowconfigure(0, weight=1)
            listbox_frame.grid_columnconfigure(0, weight=1)
            path_listbox = tk.Listbox(listbox_frame, width=80, height=6)
            path_listbox.grid(row=0, column=0, sticky="nsew")
            scrollbar = tk.Scrollbar(listbox_frame, orient="vertical", command=path_listbox.yview)
            scrollbar.grid(row=0, column=1, sticky="ns")
            path_listbox.config(yscrollcommand=scrollbar.set)
            path_listbox.bind("<<ListboxSelect>>", on_path_select)
            self.refresh_paths()
            sort_frame = tk.Frame(setup_window)
            sort_frame.grid(row=2, column=0, columnspan=2, sticky="w", padx=5, pady=5)
            tk.Label(sort_frame, text="Sort Paths:").grid(row=0, column=0, sticky="w", padx=(0, 10))
            tk.Button(sort_frame, text="▲", command=move_up, width=3).grid(row=0, column=1, padx=2)
            tk.Button(sort_frame, text="▼", command=move_down, width=3).grid(row=0, column=2, padx=2)
            tk.Button(sort_frame, text="A-Z", command=sort_az, width=3).grid(row=0, column=3, padx=2)
            tk.Label(sort_frame, text="Show Recent:").grid(row=0, column=4, sticky="w", padx=(10, 2))
            recent_count_var = tk.IntVar(value=self.preferences.get('recent_count', 10))
            recent_count_spinbox = tk.Spinbox(sort_frame, from_=0, to=10, textvariable=recent_count_var, width=5, command=update_recent_count)
            recent_count_spinbox.grid(row=0, column=5, sticky="w")
            recent_count_var.trace_add("write", lambda *args: update_recent_count())
            tk.Label(setup_window, text=" ").grid(row=3, column=0, columnspan=2, pady=5)
            tk.Label(setup_window, text="Name (optional):").grid(row=4, column=0, sticky="e", padx=5, pady=2)
            name_entry = self.create_context_menu(tk.Entry(setup_window))
            name_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=2)
            name_entry.original_name = ""
            tk.Label(setup_window, text="Path:").grid(row=5, column=0, sticky="e", padx=5, pady=2)
            path_entry = self.create_context_menu(tk.Entry(setup_window))
            path_entry.grid(row=5, column=1, sticky="ew", padx=5, pady=2)
            browse_frame = tk.Frame(setup_window)
            browse_frame.grid(row=6, column=1, sticky="w", padx=5, pady=2)
            tk.Button(browse_frame, text="Select Folder", command=select_folder).grid(row=0, column=0, padx=5)
            tk.Button(browse_frame, text="Select File", command=select_file).grid(row=0, column=1, padx=5)
            save_var = tk.BooleanVar(value=self.save_prefs_var.get())
            tk.Checkbutton(setup_window, text="Store Settings", variable=save_var, command=toggle_save_settings).grid(row=6, column=0, sticky="w", padx=5, pady=5)
            button_frame = tk.Frame(setup_window)
            button_frame.grid(row=7, column=0, columnspan=2, pady=10, sticky="w")
            tk.Button(button_frame, text="Delete Selected", command=delete_path).grid(row=0, column=0, padx=5)
            tk.Button(button_frame, text="Clear Fields", command=clear_fields).grid(row=0, column=1, padx=5)
            tk.Button(button_frame, text="Add/Update Path", command=add_or_update_path).grid(row=0, column=2, padx=5)
            tk.Button(button_frame, text="Close", command=lambda: self.close_window('paths')).grid(row=0, column=3, padx=20)
            setup_window.bind('<Return>', lambda event: add_or_update_path())

    def show_state_popup(self, message):
        popup = tk.Toplevel(self.master)
        popup.overrideredirect(True)
        tk.Label(popup, text=f"Store Settings {message}", justify=tk.CENTER).pack(padx=10, pady=5)
        popup.update_idletasks()
        width = popup.winfo_width()
        height = popup.winfo_height()
        x = self.master.winfo_x() + (self.master.winfo_width() - width) // 2
        y = self.master.winfo_y() + (self.master.winfo_height() - height) // 2
        popup.geometry(f"{width}x{height}+{x}+{y}")
        popup.after(3500, popup.destroy)   
    
    def close_window(self, window_type):
        """Helper method to close a setup window and clear its reference."""
        if window_type == 'fixs' and self.fixs_setup_window:
            self.fixs_setup_window.destroy()
            self.fixs_setup_window = None
        elif window_type == 'intranet' and self.intranet_setup_window:
            self.intranet_setup_window.destroy()
            self.intranet_setup_window = None
        elif window_type == 'windows' and self.windows_setup_window:
            self.windows_setup_window.destroy()
            self.windows_setup_window = None
        elif window_type == 'seven_zip' and self.seven_zip_setup_window:
            self.seven_zip_setup_window.destroy()
            self.seven_zip_setup_window = None
        elif window_type == 'paths' and self.manage_paths_window:
            self.manage_paths_window.destroy()
            self.manage_paths_window = None
      
    def show_instructions(self):
        title = self.master.title()
        instructions = (
            f"      {title}                                                                                           ",
            " Made By JSmyser\n\n\n",
            
            "           Instructions:\n\n",
            "   1. Select Folder or Select File to find a search location.\n\n",
            "   2. Enter one or all of the following to start search-\n\n",
            "       * Enter the Folder/Ext./File Filter you would like to filter, separated by commas.\n",
            "           • If you only search for a filter and no [Terms] it will return the entire file.\n",
            "                               •• Look below for Using Folder/Ext./File Switches ••\n\n",
            "       * Enter Main Term(s) to search, separated by commas.\n",
            "           • Prefix with '-' to exclude terms.\n\n",
            "       * Enter First Following term(s) to find the first match that follows each of the Main Term(s),\n",
            " separated by commas.\n",
            "           • Prefix with '-' to exclude terms.\n",
            "           • Prefix with '--' to define term which restarts search pairing. example --scan completed.\n\n",
            "   3. Set the number of lines before & after the Main Terms and First Following Terms you would like to find, \n separated by comma or period.\n\n",
            "   4. Click 'Search' or press Enter to begin the search, or use Quick Searches for predefined searches.\n\n",
            "   5. Results will appear in the output window. Use 'Export Results' to save findings.\n\n",
            
            "           Right click results:\n",
            "   Select Take Me To log to view entire log in Log Viewer. You can also doubleclick on the line.\n",
            "   Highlight an error and select 'Search FIXS'. This will open up a FIXS search.\n",
            "   Highlight an error and select 'Search Intranet'. This will open a search on the Intranet.\n",
            "   Highlight an error and select 'Search Windows'. This opens a sub-menu to search previously \n",
            "       saved Windows Folders on the local computer or you can browse for a new location.\n\n",
            
            "           Log Viewer:\n",
            "   Cycle Files: Cycles between Files that were found in search.\n",
            "   Cycle Results: Cycles through the Main terms within the file.\n",
            "   Search: Searches for the term entered in the search box within the file.\n\n",
            
            "           Log Viewer Highlighting Key:\n",
            "   Rt Click=Yellow. Main Terms=Blue. First Following=Green. \n",
            "   Before & After lines=Tan. Cycle Results=Red. Search Terms=Orange. \n\n\n",
                       
            "           Tool Menu:\n",
            "   Tool Menu - Store Setting creates a .json files to store all settings.\n", 
            "       • The .json needs to be in the same folder as Log_Scraper.\n",
            "   Tool Menu - Manage Folder/File Paths - Manages the folder/file path dropdown.\n",
            "       • Allows you to sort and store recent paths & your favorite search paths. (i.g., errors folder)\n",
            "   Tool Menu - Search Intranet/FIXS Setup - Allows changing of the URL if it changes in the future.\n",
            "   Tool Menu - Search Windows Setup - Manages the Rt Click search windows locations on laptop.\n",
            "\n\n",

            "           Other Notes:\n",
            "   Make sure you have 7-zip installed if you want to extract compressed files.\n",
            "   Quick Searches can be managed via the ********** Manage Quick Search ********** at the top of the dropdown.\n\n",
             
            "           Using Folder/.Ext/File Filter Switches:\n",
            "       Basic formula = (Folders OR Compressed Files) AND (Files OR Extensions)\n",
            " /folder/: Include only files within a folder + subfolders named 'folder' (e.g., /myfolder/).\n",
            " -/folder/: Exclude files within a folder + subfolders named 'folder' (e.g., -/myfolder/).\n",
            " .ext: Include only files with the specified extension (e.g., .txt, .sigusr1).\n",
            " -.ext: Exclude files with the specified extension (e.g., -.txt).\n",
            " filename: Include only the specific file (e.g., myfile.txt).\n",
            " -filename: Exclude the specific file (e.g., -myfile.txt).\n",
            " nested compressed filename: Include only the specific compressed file  (e.g., acqman.cab). \n",
            " -nested compressed filename: Exclude the specific compressed file  (e.g., -acqman.cab). \n",
            "   •• Nested compressed files filters are very useful when searching full savelogs. ••\n",
            " Wildcards: Each of these switches can use * for multiple characters & ? for a single character. \n\n",
        )
        
        # Create a Toplevel Instructions windowS
        instr_window = tk.Toplevel(self.master)
        instr_window.title("Instructions")
        instr_window.geometry("700x700+50+50")  # Set custom geometry here (width x height + x + y)
        instr_window.resizable(True, True)  # Set True for resizable. 
        instr_window.configure(bg="#f0f0f0")  # Light gray background
        
        # Add a scrolled text widget to display instructions
        instr_text = scrolledtext.ScrolledText(instr_window, wrap=tk.WORD, width=70, height=30, bg="#ffffff", fg="black", font=("Arial", 10), borderwidth=2, relief="groove")
        instr_text.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        # Configure a tag for the larger font
        instr_text.tag_configure("large_font", font=("Arial", 11, "bold"))

        # Define the line numbers to make large_font
        header_indices = []  # Example: [0, 2, 8, 15]

        # Insert instructions, applying large_font tag to header lines
        for i, line in enumerate(instructions):
            if i in header_indices or line.strip().endswith(":"):
                instr_text.insert(tk.END, line, "large_font")
            else:
                instr_text.insert(tk.END, line)

        instr_text.config(state=tk.DISABLED)  # Make it read-only


#### Export Results and Clear All Fields Buttons. ------------------------------------------------------

    def export_results(self):
        output_text = self.output.get("1.0", tk.END).strip()
        if not output_text:
            messagebox.showinfo("Export", "No results to export.", parent=self.master)
            return

        now = datetime.now().strftime("%Y%m%d%H%M%S%f")[:-3]
        initial_filename = f"{now}_LogScraper"

        # Set initial directory based on self.path_var
        initial_path = self.path_var.get().strip()
        if ": " in initial_path:
            initial_path = initial_path.split(": ", 1)[1].strip()  # Strip "Recent 1: " or "(Current): "
        if os.path.isfile(initial_path):  # If it's a file, use parent directory
            initial_path = os.path.dirname(initial_path)
        if not initial_path or not os.path.exists(initial_path):
            initial_path = "C:/"

        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            initialfile=initial_filename,
            initialdir=initial_path,  # Added
            filetypes=[("Text files", "*.txt"), ("DOCX files", "*.docx"), ("All files", "*.*")],
            title="Save Results As",
            parent=self.master
        )
        if not file_path:
            return

        progress_window = tk.Toplevel(self.master)
        progress_window.title("Exporting...")
        progress_window.geometry("300x120+500+300")
        progress_window.transient(self.master)
        progress_window.grab_set()
        tk.Label(progress_window, text="Exporting document...").pack(pady=10)
        progress = ttk.Progressbar(progress_window, mode='determinate', maximum=100)
        progress.pack(pady=10, padx=20, fill=tk.X)
        cancel_var = tk.BooleanVar(value=False)
        tk.Button(progress_window, text="Cancel", command=lambda: cancel_var.set(True)).pack(pady=5)
        progress_window.update()

        success = False
        if file_path.endswith('.docx'):
            success = self.export_to_docx(output_text, file_path, progress, cancel_var, progress_window)
        else:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    lines = output_text.splitlines()
                    total_lines = len(lines)
                    for i, line in enumerate(lines):
                        if cancel_var.get():
                            progress_window.destroy()
                            return
                        f.write(line + '\n')
                        progress['value'] = ((i + 1) / total_lines) * 100
                        progress_window.update()
                success = True
            except Exception as e:
                progress_window.destroy()
                messagebox.showerror("Error", f"Failed to export: {str(e)}", parent=self.master)

        progress_window.destroy()
        if success and not cancel_var.get():
            popup = tk.Toplevel(self.master)
            popup.title("Export")
            popup.geometry("300x120+500+300")
            popup.transient(self.master)
            popup.grab_set()
            tk.Label(popup, text="Export Successful!\nOpen file?", font=("Arial", "12")).pack(pady=10)
            btn_frame = tk.Frame(popup)
            btn_frame.pack(pady=10)
            tk.Button(btn_frame, text="Yes", command=lambda: [os.startfile(file_path), popup.destroy()]).pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="No", command=popup.destroy).pack(side=tk.LEFT, padx=5)
    
    
    def clear_input_fields(self):
        # Clear all Fields button. 
        self.extensions_var.set('')  # Clear extensions
        self.main_term_var.set('')    # Clear main term
        self.following_terms_var.set('')  # Clear first following terms
        self.main_context.set('0,0')  # Reset main context
        self.following_context.set('0,0')  # Reset following context
        self.start_date_entry.delete(0, 'end')  # Clear start date
        self.end_date_entry.delete(0, 'end')  # Clear end date
        self.quick_date_var.set('')  # Clear quick selection
        self.date_filter_var.set(False)  # Disable date filter
        self.output.delete('1.0', tk.END)
        self.export_button['state'] = 'disabled'  # Disable the export button if there's no content


#### Right Click Control for Main page Entry Boxes and Log_Viewer ------------------------------------------------------

    def create_context_menu(self, widget):
        context_menu = tk.Menu(widget, tearoff=0)
        context_menu.add_command(label="Cut", command=lambda: self.cut_to_clipboard(widget))
        context_menu.add_command(label="Copy", command=lambda: self.copy_to_clipboard(widget))
        context_menu.add_command(label="Paste", command=lambda: self.paste_from_clipboard(widget))
        context_menu.add_command(label="Delete", command=lambda: self.delete_selection(widget))
        context_menu.add_command(label="Select All", command=lambda: self.select_all(widget))
        context_menu.add_command(label="Clear Field", command=lambda: self.clear_field(widget))  # New option
        
        def popup(event):
            # Set focus to the widget before showing the context menu
            widget.focus_set()            
            context_menu.post(event.x_root, event.y_root)

        widget.bind("<Button-3>", popup)
        return widget

    def cut_to_clipboard(self, widget):
        try:
            selected_text = widget.selection_get()  # Get the selected text
            widget.delete(tk.SEL_FIRST, tk.SEL_LAST)  # Delete the selected text from the widget
            self.master.clipboard_clear()  # Clear the clipboard
            self.master.clipboard_append(selected_text)  # Append the selected text to the clipboard
        except tk.TclError:
            pass  # No selection to cut
 
    def copy_to_clipboard(self, widget):
        try:
            widget.event_generate("<<Copy>>")
            self.master.clipboard_clear()
            self.master.clipboard_append(widget.selection_get())
        except tk.TclError:
            pass

    def paste_from_clipboard(self, widget):
        try:
            widget.event_generate("<<Paste>>")
        except tk.TclError:
            pass

    def delete_selection(self, widget):
        try:
            widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
        except tk.TclError:
            pass  # No selection to delete

    def select_all(self, widget):
        if isinstance(widget, (tk.Text, scrolledtext.ScrolledText)):
            widget.tag_add("sel", "1.0", tk.END)
            widget.mark_set(tk.INSERT, "1.0")
            widget.see("1.0")
        elif isinstance(widget, tk.Entry):
            widget.select_range(0, tk.END)
            widget.icursor(tk.END)

    def clear_field(self, widget):
        try:
            widget.delete(0, tk.END)  # Clear the entire entry field
        except tk.TclError:
            pass  # Handle any potential errors gracefully


#### Load and save Preferences to log_scraper_settings.json -----------------------------------------------------------------

    def load_preferences(self):
        self.preferences = {}
        self.config_path = get_config_path("log_scraper_settings.json")
        try:
            with open(self.config_path, "r") as file:
                self.preferences = json.load(file)
                if not self.preferences.get('quick_search_order') and self.preferences.get('quick_searches'):
                    self.preferences['quick_search_order'] = list(self.preferences['quick_searches'].keys())
                if 'configure_graph' in self.preferences:
                    loaded_graph = self.preferences['configure_graph']
                    self.configure_graph = {}
                    for term, settings in loaded_graph.items():
                        clean_term = re.sub(r'[:;,\s]+$', '', term)
                        # Note: examples isn't available here, so we'll rely on graph_results to match later
                        self.configure_graph[clean_term] = settings
        except FileNotFoundError:
            self.preferences = {
                'save_entries': False,
                'quick_searches': {},
                'quick_search_order': [],
                'manage_quick_search_geometry': '600x500+50+50',
                'last_browse_location': '',
                'show_last_file': False,
                'date_filter_enabled': False,
                'start_date': '',
                'end_date': '',
                'quick_date': '',
                'path': '',
                'extensions': '',
                'main_term': '',
                'following_terms': '',
                'main_context': '0,0',
                'following_context': '0,0',
                'window_geometry': '750x730+5+10',
                'log_viewer_geometry': '700x700+60+60',
                'search_locations': {},
                'search_location_order': [],
                'fixs_base_url': DEFAULT_FIXS_BASE_URL,
                'intranet_base_url': DEFAULT_INTRANET_BASE_URL,
                'seven_zip_path': DEFAULT_7Z_PATH,
                'search_paths': {},
                'search_path_order': [],
                'recent_count': 5,
                'configure_graph': {}
            }
            self.configure_graph = {}
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load preferences: {str(e)}", parent=self.master)
            self.preferences = {
                'save_entries': False,
                'quick_searches': {},
                'quick_search_order': [],
                'manage_quick_search_geometry': '600x500+50+50',
                'last_browse_location': '',
                'show_last_file': False,
                'date_filter_enabled': False,
                'start_date': '',
                'end_date': '',
                'quick_date': '',
                'path': '',
                'extensions': '',
                'main_term': '',
                'following_terms': '',
                'main_context': '0,0',
                'following_context': '0,0',
                'window_geometry': '750x730+5+10',
                'log_viewer_geometry': '700x700+60+60',
                'search_locations': {},
                'search_location_order': [],
                'fixs_base_url': DEFAULT_FIXS_BASE_URL,
                'intranet_base_url': DEFAULT_INTRANET_BASE_URL,
                'seven_zip_path': DEFAULT_7Z_PATH,
                'search_paths': {},
                'search_path_order': [],
                'recent_count': 5,
                'configure_graph': {}
            }
            self.configure_graph = {}

        # Apply loaded settings that need immediate effect
        self.last_browse_location = {"path": self.preferences.get('last_browse_location', '')}
        if 'window_geometry' in self.preferences:
            self.master.geometry(self.preferences['window_geometry'])
        self.search_locations = self.preferences.get('search_locations', {})
        self.search_location_order = self.preferences.get('search_location_order', [])
        self.seven_zip_path = self.preferences.get('seven_zip_path', DEFAULT_7Z_PATH)
        self.search_paths = self.preferences.get('search_paths', {})
        self.search_path_order = self.preferences.get('search_path_order', [])
        self.fixs_base_url = self.preferences.get('fixs_base_url', DEFAULT_FIXS_BASE_URL)
        self.intranet_base_url = self.preferences.get('intranet_base_url', DEFAULT_INTRANET_BASE_URL)

    def on_resize(self, event):
        if event.widget == self.master:
            self.preferences['window_geometry'] = f"{event.width}x{event.height}+{event.x}+{event.y}"
            if self.save_prefs_var.get():
                self.save_preferences()

    def save_preferences(self):
        if not self.save_prefs_var.get():
            return
        try:
            self.preferences['path'] = self.path_var.get()
            self.preferences['extensions'] = self.extensions_var.get()
            self.preferences['main_term'] = self.main_term_var.get()
            self.preferences['following_terms'] = self.following_terms_var.get()
            self.preferences['main_context'] = self.main_context.get()
            self.preferences['following_context'] = self.following_context.get()
            self.preferences['save_entries'] = self.save_prefs_var.get()
            self.preferences['window_geometry'] = self.master.geometry()
            self.preferences['log_viewer_geometry'] = self.preferences.get('log_viewer_geometry', '700x700+60+60')
            self.preferences['last_browse_location'] = self.last_browse_location["path"]
            self.preferences['show_last_file'] = self.show_last_file_var.get()
            self.preferences['date_filter_enabled'] = self.date_filter_var.get()
            self.preferences['start_date'] = self.start_date_entry.get() if self.start_date_entry.get() else ''
            self.preferences['end_date'] = self.end_date_entry.get() if self.end_date_entry.get() else ''
            self.preferences['quick_date'] = self.quick_date_var.get()
            self.preferences['search_locations'] = self.search_locations
            self.preferences['search_location_order'] = self.search_location_order
            self.preferences['seven_zip_path'] = self.seven_zip_path
            self.preferences['search_paths'] = self.search_paths
            self.preferences['search_path_order'] = self.search_path_order
            self.preferences['fixs_base_url'] = self.fixs_base_url
            self.preferences['intranet_base_url'] = self.intranet_base_url
            self.preferences['recent_count'] = self.preferences.get('recent_count', 5)

            # Save configure_graph
            cleaned_graph = {}
            for term, settings in self.configure_graph.items():
                cleaned_term = re.sub(r'[:;,\s]+$', '', term)
                cleaned_graph[cleaned_term] = settings
            self.preferences['configure_graph'] = cleaned_graph

            # Save quick search order from listbox if manage window exists
            if hasattr(self, 'quick_search_listbox') and self.manage_window and self.manage_window.winfo_exists():
                self.preferences['quick_search_order'] = []
                for i in range(self.quick_search_listbox.size()):
                    item = self.quick_search_listbox.get(i)
                    name = item if not item.startswith("(Auto) ") else item[7:]
                    if name not in self.preferences['quick_search_order']:
                        self.preferences['quick_search_order'].append(name)

            with open(self.config_path, "w") as file:
                json.dump(self.preferences, file, indent=4)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save preferences: {str(e)}", parent=self.master)
            


#### Select Folder/File and Path Dropdown --------------------------------------------------

    def select_folder(self):
        """Open folder picker at parent directory if path_var is a file."""
        initial_path = self.path_var.get().strip()
        if ": " in initial_path:
            initial_path = initial_path.split(": ", 1)[1].strip()  # Strip "Recent 1: " or "(Current): "     
        if os.path.isfile(initial_path):  # If it's a file (compressed or not), use parent directory
            containing_folder = os.path.dirname(initial_path)
            initial_path = os.path.dirname(containing_folder)
        elif os.path.isdir(initial_path):
            initial_path = os.path.dirname(initial_path)        
        if not initial_path or not os.path.exists(initial_path):
            initial_path = "C:/"      
        folder_selected = filedialog.askdirectory(initialdir=initial_path)
        if folder_selected:
            self.path_var.set(folder_selected)
            self.refresh_paths()
            self.update_path_combo()

    def toggle_advanced_extract(self, event):
        """Toggle advanced extraction with a checkbox in a confirmation popup."""
        popup = tk.Toplevel(self.master)
        popup.transient(self.master)
        popup.title("Advanced Extraction Setting")
        message = (
            "                    Advanced Compressed File Extractor:\n\n"
            "This extracts all compressed files within folders/subfolders.\n"
            "  Doing this may require significant time, hard drive space,\n"
            "                      and could risk crashing the system.\n\n"
            "                              Proceed with Caution!"
        )
        tk.Label(popup, text=message, justify="left").pack(padx=20, pady=10)

        # Checkbox inside popup
        advanced_var = tk.BooleanVar(value=self.advanced_extract)  # Reflect current state
        checkbox = tk.Checkbutton(
            popup, 
            text="Enable Advanced Extraction", 
            variable=advanced_var
        )
        checkbox.pack(pady=5)

        def confirm():
            self.advanced_extract = advanced_var.get()  # Set based on checkbox state
            popup.destroy()
        
        def cancel():
            popup.destroy()  # No change to self.advanced_extract
        
        button_frame = tk.Frame(popup)
        button_frame.pack(pady=5)
        tk.Button(button_frame, text="OK", command=confirm).grid(row=0, column=0, padx=10)
        tk.Button(button_frame, text="Cancel", command=cancel).grid(row=0, column=1, padx=10)
        
        popup.update_idletasks()
        width = popup.winfo_width()
        height = popup.winfo_height()
        top = self.master.winfo_toplevel()
        top.update()
        x = top.winfo_x() + (top.winfo_width() - width) // 2
        y = top.winfo_y() + (top.winfo_height() - height) // 2
        popup.geometry(f"{width}x{height}+{x}+{y}")
        popup.grab_set()

    def select_file(self):
        """Open file picker at parent directory if path_var is a file."""
        initial_path = self.path_var.get().strip()
        if ": " in initial_path:
            initial_path = initial_path.split(": ", 1)[1].strip()  # Strip "Recent 1: " or "(Current): "
        if os.path.isfile(initial_path):  # If it's a file (compressed or not), use parent directory
            initial_path = os.path.dirname(initial_path)
        if not initial_path or not os.path.exists(initial_path):
            initial_path = "C:/"
        file_selected = filedialog.askopenfilename(initialdir=initial_path)
        if file_selected:
            self.path_var.set(file_selected)
            self.refresh_paths()
            self.update_path_combo()

    def on_combo_select(self, event):
        """Sync self.path_var with combobox selection, stripping labels."""
        selected = self.path_combo.get().strip()
        if selected and ": " in selected:
            path = selected.split(": ", 1)[1].strip()
            self.path_var.set(path)       

    def get_path_combo_values(self):
        """Generate combobox values without display prefix tweaks."""
        values = []
        current_path = self.path_var.get().strip()
        if current_path:
            values.append(f"(Current): {current_path}")
        for name in self.search_path_order:
            if name in self.search_paths and self.search_paths[name]:
                display = name if name != self.search_paths[name] else self.search_paths[name]
                prefix = f"(Recent {name.split()[1]}) " if name.startswith("Recent ") else ""
                values.append(f"{prefix}{display}")
        return values

        # Refreshes the listbox in Manage Path Window
    def refresh_paths(self):
        """Refresh Manage Paths listbox."""
        if hasattr(self, 'manage_paths_window') and self.manage_paths_window:
            try:
                path_listbox = self.manage_paths_window.children['!frame'].children['!listbox']
                path_listbox.delete(0, tk.END)
                current_path = self.path_var.get().strip()
                if current_path:
                    path_listbox.insert(tk.END, f"(Current): {current_path}")
                    path_listbox.insert(tk.END, "-" * 20)
                for name in self.search_path_order:
                    if name in self.search_paths and self.search_paths[name]:
                        path_listbox.insert(tk.END, f"{name}: {self.search_paths[name]}")
                    elif name.startswith("Recent "):
                        path_listbox.insert(tk.END, f"{name}:")
            except tk.TclError:
                pass

    def update_path_combo(self): # Refreshes the combobox Path on the Main page.
        """Refresh main combobox."""
        values = []
        current_path = self.path_var.get().strip()
        if current_path:
            values.append(f"(Current): {current_path}")
        for name in self.search_path_order:
            if name in self.search_paths and self.search_paths[name]:
                values.append(f"{name}: {self.search_paths[name]}")
            elif name.startswith("Recent "):
                values.append(f"{name}:")
        self.path_combo['values'] = values



#### Quick Search Dropdown and Manage Quick Search. -----------------------------------------------

    def update_quick_search_dropdown(self):
        quick_searches = []
        for name in self.preferences.get('quick_search_order', []):
            if name in self.preferences.get('quick_searches', {}):
                if self.preferences['quick_searches'][name].get('auto_start', False):
                    quick_searches.append(f"(Auto) {name}")
                else:
                    quick_searches.append(name)
        
        self.quick_search_dropdown['values'] = ["********** Manage Quick Search **********", "-" * 45] + quick_searches
        self.quick_search_var.set('               Quick Search:')  # Display "Quick Search:" when idle
        self.quick_search_dropdown.config(height=20)  # Doubled from default (typically ~10) to 20

    def on_quick_search_select(self, event):
        selected_search = self.quick_search_var.get()
        if selected_search == "********** Manage Quick Search **********":
            self.open_manage_quick_search_window()
            self.quick_search_var.set('               Quick Search:')  # Reset to idle state after action
        elif selected_search == "-" * 45:
            self.quick_search_var.set('               Quick Search:')  # Reset to idle state
        elif selected_search:
            # Check if the selected search includes "(Auto)"
            name = selected_search
            if name.startswith("(Auto) "):
                name = name[7:]  # Remove the "(Auto) " prefix before lookup
            
            if name in self.preferences['quick_searches']:
                search_data = self.preferences['quick_searches'][name]
                self.extensions_var.set(search_data['extensions'])
                self.main_term_var.set(search_data['main_term'])
                self.following_terms_var.set(search_data['following_terms'])
                self.main_context.set(search_data['main_context'])
                self.following_context.set(search_data['following_context'])
                if selected_search.startswith("(Auto) "):  # If auto start was selected, initiate search
                    self.start_search()  # Method to start searching
                self.quick_search_var.set('               Quick Search:')  # Reset to idle state after action
            else:
                messagebox.showerror("Error", "Selected quick search not found.", parent=self.master)
                self.quick_search_var.set('               Quick Search:')  # Reset to idle state

    def open_manage_quick_search_window(self, event=None):  # Added optional event parameter
        # Check if the window is already open
        if self.manage_window is not None and self.manage_window.winfo_exists():
            self.manage_window.lift()  # Bring existing window to the front
            return

        def move_up():
            selected = quick_search_listbox.curselection()
            if selected:
                index = selected[0]
                if index > 0:
                    text = quick_search_listbox.get(index)
                    quick_search_listbox.delete(index)
                    quick_search_listbox.insert(index - 1, text)
                    quick_search_listbox.selection_set(index - 1)

                    # Update the order in preferences
                    self.preferences['quick_search_order'].insert(index - 1, self.preferences['quick_search_order'].pop(index))
                    self.save_preferences()
                    self.update_quick_search_dropdown()

        def move_down():
            selected = quick_search_listbox.curselection()
            if selected:
                index = selected[0]
                if index < quick_search_listbox.size() - 1:
                    text = quick_search_listbox.get(index)
                    quick_search_listbox.delete(index)
                    quick_search_listbox.insert(index + 1, text)
                    quick_search_listbox.selection_set(index + 1)

                    # Update the order in preferences
                    self.preferences['quick_search_order'].insert(index + 1, self.preferences['quick_search_order'].pop(index))
                    self.save_preferences()
                    self.update_quick_search_dropdown()

        def sort_az():
            # Sort quick searches alphabetically
            self.preferences['quick_search_order'].sort()
            self.save_preferences()
            refresh_quick_searches()
            self.update_quick_search_dropdown()

        def add_or_update_quick_search():
            if not self.save_prefs_var.get():
                messagebox.showwarning(
                    "Warning",
                    "Store Settings is not enabled. Quick search changes won't be saved. Enable it via the Options dropdown.",
                    parent=manage_window
                )
                return

            name = name_entry.get().strip()
            extensions = extensions_entry.get().strip()
            main_term = main_term_entry.get().strip()
            following_terms = following_terms_entry.get().strip()
            main_context = main_context_entry.get().strip()
            following_context = following_context_entry.get().strip()
            auto_start = auto_start_var.get()  # Get the state of the checkbox

            # Validation based on Auto Start

            if not (name and main_context and following_context):
                messagebox.showerror("Error", "Name, Main Term (before,after), and First Following (before,after) are required.", parent=manage_window)
                return

            if not main_context:
                main_context = "0,0"  # Default value
            if not following_context:
                following_context = "0,0"  # Default value
            
            # If the name has changed and the original name exists, update its position
            if name_entry.original_name and name_entry.original_name != name:
                if name_entry.original_name in self.preferences['quick_searches']:
                    # Find the original position in quick_search_order
                    try:
                        index = self.preferences['quick_search_order'].index(name_entry.original_name)
                        # Replace the old name with the new name at the same position
                        self.preferences['quick_search_order'][index] = name
                    except ValueError:
                        # If not found (shouldn't happen), append as fallback
                        self.preferences['quick_search_order'].append(name)
                    # Remove the old entry from quick_searches
                    del self.preferences['quick_searches'][name_entry.original_name]
            elif not name_entry.original_name:
                # If it's a new entry, append it to the end
                self.preferences['quick_search_order'].append(name)

            # Update or add the quick search with the new settings
            self.preferences['quick_searches'][name] = {
                'extensions': extensions,
                'main_term': main_term,
                'following_terms': following_terms,
                'main_context': main_context,
                'following_context': following_context,
                'auto_start': auto_start  # Include the auto_start flag
            }
            
            self.save_preferences()  # Save to file (only if save_prefs_var is True)
            refresh_quick_searches()
            self.update_quick_search_dropdown()  # Update dropdown in main window
            name_entry.original_name = name  # Track the current name for future updates

        def copy_quick_search():
            current_name = name_entry.get()
            if not current_name:
                messagebox.showwarning("Warning", "Please select or enter a name for the search to copy.", parent=manage_window)
                return

            # Gather all the data from the current fields
            new_name = f"{current_name} (Copy)"
            new_extensions = extensions_entry.get()
            new_main_term = main_term_entry.get()
            new_following_terms = following_terms_entry.get()
            new_main_context = main_context_entry.get()
            new_following_context = following_context_entry.get()
            new_auto_start = auto_start_var.get()

            # Check if this new name already exists to avoid overwriting
            if new_name in self.preferences['quick_searches']:
                messagebox.showwarning("Warning", f"A quick search with the name '{new_name}' already exists.", parent=manage_window)
                return

            # Add the new search to preferences
            self.preferences['quick_searches'][new_name] = {
                'extensions': new_extensions,
                'main_term': new_main_term,
                'following_terms': new_following_terms,
                'main_context': new_main_context,
                'following_context': new_following_context,
                'auto_start': new_auto_start
            }

            # Add the new search to the order list
            self.preferences['quick_search_order'].append(new_name)

            # Refresh the listbox and dropdown to show the new quick search
            self.save_preferences()  # Save to file
            refresh_quick_searches()              
            self.update_quick_search_dropdown()  

            # Select the newly added search in the listbox
            for index in range(quick_search_listbox.size()):
                item = quick_search_listbox.get(index)
                if item == f"(Auto) {new_name}" if new_auto_start else new_name:
                    quick_search_listbox.selection_clear(0, tk.END)
                    quick_search_listbox.selection_set(index)
                    quick_search_listbox.see(index)  # Scroll to make the item visible
                    quick_search_listbox.event_generate("<<ListboxSelect>>")  # Trigger the selection event
                    break

        def delete_quick_search():
            selected = quick_search_listbox.curselection()
            if not selected:
                messagebox.showerror("Error", "Please select a quick search to delete.", parent=manage_window)
                return

            name = quick_search_listbox.get(selected[0])
            actual_name = name if not name.startswith("(Auto) ") else name[7:]  # Remove "(Auto) " if present

            if messagebox.askyesno("Confirm Deletion", f"Are you sure you want to delete '{name}'?", parent=manage_window):
                if actual_name in self.preferences['quick_searches']:
                    del self.preferences['quick_searches'][actual_name]
                    self.preferences['quick_search_order'].remove(actual_name)
                    self.save_preferences()  # Save to file
                    refresh_quick_searches()
                    self.update_quick_search_dropdown()  # Update dropdown in main window
                else:
                    messagebox.showerror("Error", f"Quick search '{name}' not found.", parent=manage_window)

            # Refresh the listbox and dropdown to show the new quick search
            self.save_preferences()  # Save to file
            refresh_quick_searches()              
            self.update_quick_search_dropdown()  

        def import_from_main():
            # Import current main window search settings into entry fields
            name_entry.delete(0, tk.END)  # Clear the name field
            name_entry.original_name = ""  # Reset tracking
            extensions_entry.delete(0, tk.END)
            extensions_entry.insert(0, self.extensions_var.get().strip())
            main_term_entry.delete(0, tk.END)
            main_term_entry.insert(0, self.main_term_var.get().strip())
            following_terms_entry.delete(0, tk.END)
            following_terms_entry.insert(0, self.following_terms_var.get().strip())
            main_context_entry.delete(0, tk.END)
            main_context_entry.insert(0, self.main_context.get().strip() or "0,0")
            following_context_entry.delete(0, tk.END)
            following_context_entry.insert(0, self.following_context.get().strip() or "0,0")
            auto_start_var.set(False)  # Default to unchecked

        def clear_fields():
            # Clear only the before/after entry fields to "0,0"
            main_context_entry.delete(0, tk.END)
            main_context_entry.insert(0, "0,0")
            following_context_entry.delete(0, tk.END)
            following_context_entry.insert(0, "0,0")

        def refresh_quick_searches():
            quick_search_listbox.delete(0, tk.END)
            for name in self.preferences.get('quick_search_order', []):
                if name in self.preferences.get('quick_searches', {}):
                    display_name = f"(Auto) {name}" if self.preferences['quick_searches'][name].get('auto_start', False) else name
                    quick_search_listbox.insert(tk.END, display_name)

        def on_quick_search_select(event):
            selected = quick_search_listbox.curselection()
            if not selected:
                return
            name = quick_search_listbox.get(selected[0])
            if name.startswith("(Auto) "):
                name = name[7:]  # Remove "(Auto) " prefix before lookup
            if name in self.preferences['quick_searches']:
                search_data = self.preferences['quick_searches'][name]
                name_entry.delete(0, tk.END)
                name_entry.insert(0, name)
                name_entry.original_name = name  # Track the original name for update checks
                extensions_entry.delete(0, tk.END)
                extensions_entry.insert(0, search_data['extensions'])
                main_term_entry.delete(0, tk.END)
                main_term_entry.insert(0, search_data['main_term'])
                following_terms_entry.delete(0, tk.END)
                following_terms_entry.insert(0, search_data['following_terms'])
                main_context_entry.delete(0, tk.END)
                main_context_entry.insert(0, search_data['main_context'])
                following_context_entry.delete(0, tk.END)
                following_context_entry.insert(0, search_data['following_context'])
                auto_start_var.set(search_data.get('auto_start', False))  # Set the checkbox state
            else:
                # Reset all fields if the search isn't found
                name_entry.delete(0, tk.END)
                extensions_entry.delete(0, tk.END)
                main_term_entry.delete(0, tk.END)
                following_terms_entry.delete(0, tk.END)
                main_context_entry.delete(0, tk.END)
                main_context_entry.insert(0, "0,0")
                following_context_entry.delete(0, tk.END)
                following_context_entry.insert(0, "0,0")
                auto_start_var.set(False)  # Uncheck if not found

        # Ensure 'quick_search_order' exists in preferences
        if 'quick_search_order' not in self.preferences:
            self.preferences['quick_search_order'] = list(self.preferences['quick_searches'].keys())

        # Create the management window
        manage_window = tk.Toplevel(self.master)
        self.manage_window = manage_window  # Store the reference
        manage_window.title("Manage Quick Searches")
        manage_window.geometry(self.preferences.get('manage_quick_search_geometry', "510x475+50+50"))
        manage_window.transient(self.master)  # Keep on top of Log Scraper only

        # Configure grid weights for resizing
        manage_window.grid_rowconfigure(1, weight=1)  # Listbox row expands vertically
        manage_window.grid_rowconfigure(3, weight=0)  # Blank space row
        manage_window.grid_columnconfigure(1, weight=1)  # Column 1 expands horizontally for stretchable entries

        # List of quick searches
        tk.Label(manage_window, text="Existing Quick Searches:").grid(row=0, column=0, columnspan=2, sticky="w", padx=5, pady=5)
        
        # Frame for listbox and scrollbar
        listbox_frame = tk.Frame(manage_window)
        listbox_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)
        listbox_frame.grid_rowconfigure(0, weight=1)
        listbox_frame.grid_columnconfigure(0, weight=1)

        quick_search_listbox = tk.Listbox(listbox_frame, width=80, height=10)
        quick_search_listbox.grid(row=0, column=0, sticky="nsew")
        
        # Scrollbar
        scrollbar = tk.Scrollbar(listbox_frame, orient="vertical", command=quick_search_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        
        quick_search_listbox.config(yscrollcommand=scrollbar.set)
        quick_search_listbox.bind("<<ListboxSelect>>", on_quick_search_select)  # Bind selection event
        refresh_quick_searches()

        # Sort Quick Searches label and buttons frame
        sort_frame = tk.Frame(manage_window)
        sort_frame.grid(row=2, column=0, columnspan=2, sticky="w", padx=5, pady=5)
        tk.Label(sort_frame, text="Sort Quick Searches:").grid(row=0, column=0, sticky="w", padx=(0, 0))
        up_button = tk.Button(sort_frame, text="▲", command=move_up, width=3)
        up_button.grid(row=0, column=1, padx=2)
        down_button = tk.Button(sort_frame, text="▼", command=move_down, width=3)
        down_button.grid(row=0, column=2, padx=2)
        az_button = tk.Button(sort_frame, text="A-Z", command=sort_az, width=3)
        az_button.grid(row=0, column=3, padx=2)
        copy_button = tk.Button(sort_frame, text="Copy Quick Search", command=copy_quick_search)
        copy_button.grid(row=0, column=4, padx=(40, 5))
        import_button = tk.Button(sort_frame, text="Import from Main", command=import_from_main)
        import_button.grid(row=0, column=5, padx=5)

        # Blank space (replacing separator)
        tk.Label(manage_window, text=" ").grid(row=3, column=0, columnspan=2, pady=5)

        # Name row with Auto Start checkbox
        tk.Label(manage_window, text="Name:").grid(row=4, column=0, sticky="e", padx=5, pady=2)
        name_entry = self.create_context_menu(tk.Entry(manage_window))
        name_entry.grid(row=4, column=1, sticky="ew", padx=5, pady=2)
        name_entry.original_name = ""  # Track the original name for update checks
        auto_start_var = tk.BooleanVar()
        auto_start_checkbox = tk.Checkbutton(manage_window, text="Auto Start", variable=auto_start_var)
        auto_start_checkbox.grid(row=4, column=0, sticky="w", padx=100, pady=2)

        # File Extensions
        tk.Label(manage_window, text="Folder/Ext./File Filter (comma-separated, optional):").grid(row=5, column=0, sticky="e", padx=5, pady=2)
        extensions_entry = self.create_context_menu(tk.Entry(manage_window))
        extensions_entry.grid(row=5, column=1, sticky="ew", padx=5, pady=2)

        # Main Term
        tk.Label(manage_window, text="Main Term (comma-separated):").grid(row=6, column=0, sticky="e", padx=5, pady=2)
        main_term_entry = self.create_context_menu(tk.Entry(manage_window))
        main_term_entry.grid(row=6, column=1, sticky="ew", padx=5, pady=2)

        # First Following
        tk.Label(manage_window, text="First Following (comma-separated, optional):").grid(row=7, column=0, sticky="e", padx=5, pady=2)
        following_terms_entry = self.create_context_menu(tk.Entry(manage_window))
        following_terms_entry.grid(row=7, column=1, sticky="ew", padx=5, pady=2)

        # Main Term (before,after)
        tk.Label(manage_window, text="Main Term (before,after):").grid(row=8, column=0, sticky="e", padx=5, pady=2)
        main_context_entry = self.create_context_menu(tk.Entry(manage_window, width=10))
        main_context_entry.insert(0, "0,0")  # Default value
        main_context_entry.grid(row=8, column=1, sticky="w", padx=5, pady=2)

        # First Following (before,after)
        tk.Label(manage_window, text="First Following (before,after):").grid(row=9, column=0, sticky="e", padx=5, pady=2)
        following_context_entry = self.create_context_menu(tk.Entry(manage_window, width=10))
        following_context_entry.insert(0, "0,0")  # Default value
        following_context_entry.grid(row=9, column=1, sticky="w", padx=5, pady=2)

        # Buttons frame (tied to left in column 0)
        button_frame = tk.Frame(manage_window)
        button_frame.grid(row=10, column=0, columnspan=5, sticky="w", padx=45, pady=10)  # Adjusted padx to 35
        
        delete_button = tk.Button(button_frame, text="Delete Selected", command=delete_quick_search)
        delete_button.grid(row=0, column=0, padx=10)
        clear_button = tk.Button(button_frame, text="Clear Fields", command=clear_fields)
        clear_button.grid(row=0, column=1, padx=5)
        add_update_button = tk.Button(button_frame, text="Add/Update Quick Search", command=add_or_update_quick_search)
        add_update_button.grid(row=0, column=2, padx=5)
        close_button = tk.Button(button_frame, text="Close", command=lambda: self.on_manage_window_closing(manage_window))
        close_button.grid(row=0, column=3, padx=(20, 5))

        # Bind Enter key to Add/Update Quick Search
        manage_window.bind('<Return>', lambda event: add_or_update_quick_search())

        # Handle window closing
        manage_window.protocol("WM_DELETE_WINDOW", lambda: self.on_manage_window_closing(manage_window))

        # Ensure 'quick_search_order' exists in preferences
        if 'quick_search_order' not in self.preferences:
            self.preferences['quick_search_order'] = list(self.preferences['quick_searches'].keys())

    def on_manage_window_closing(self, window):
        # Save the geometry of the Manage Quick Searches window
        self.preferences['manage_quick_search_geometry'] = window.geometry()
        if self.save_prefs_var.get():
            self.save_preferences()
        self.manage_window = None  # Clear the reference
        window.destroy()




#### self.output Cycle Files, Cycle Results, Auto Scroll, Search, and Display. 

    def cycle_output_files(self, direction):
        text = self.output.get("1.0", tk.END)
        lines = text.splitlines()
        current_pos = self.output.index(tk.INSERT).split('.')[0]  # Cursor line

        # Find all file markers
        file_positions = [i + 1 for i, line in enumerate(lines) if line.strip().startswith(">>>> ") and line.strip().endswith(" <<<<")]
        if not file_positions:
            return

        # Find current file position relative to cursor
        current_line = int(current_pos)
        if direction > 0:  # Down
            new_file_idx = next((i for i, pos in enumerate(file_positions) if pos > current_line), 0)
        else:  # Up
            new_file_idx = max([i for i, pos in enumerate(file_positions) if pos < current_line], default=-1)
            if new_file_idx == -1:
                new_file_idx = len(file_positions) - 1
        new_line = file_positions[new_file_idx]

        # Clear previous selection
        self.output.tag_remove("sel_file", "1.0", tk.END)
        self.output.tag_remove("sel_result", "1.0", tk.END)
        # Highlight new file with sel_file tag
        self.output.tag_add("sel_file", f"{new_line}.0", f"{new_line + 1}.0")
        # Move cursor and scroll (2 lines from top)
        self.output.mark_set(tk.INSERT, f"{new_line}.0")
        self.output.yview(f"{new_line - 2}.0" if new_line > 2 else "1.0")
        self.update_output_result_label()
               
    def cycle_output_results(self, direction):
        text = self.output.get("1.0", tk.END)
        lines = text.splitlines()
        current_pos = self.output.index(tk.INSERT).split('.')[0]  # Cursor line

        # Find all .mm lines
        mm_positions = [i + 1 for i, line in enumerate(lines) if ".mm" in line]
        if not mm_positions:
            return

        # Find current result position relative to cursor
        current_line = int(current_pos)
        if direction > 0:  # Right
            new_mm_idx = next((i for i, pos in enumerate(mm_positions) if pos > current_line), 0)
        else:  # Left
            new_mm_idx = max([i for i, pos in enumerate(mm_positions) if pos < current_line], default=-1)
            if new_mm_idx == -1:
                new_mm_idx = len(mm_positions) - 1
        new_line = mm_positions[new_mm_idx]

        # Calculate group start using first number in main_context
        try:
            before_context = int(self.main_context.get().split(',')[0].strip())
        except (ValueError, IndexError):
            before_context = 0  # Default to 0 if invalid
        group_start = max(1, new_line - before_context)  # Don’t go below line 1

        # Clear previous selection
        self.output.tag_remove("sel_file", "1.0", tk.END)
        self.output.tag_remove("sel_result", "1.0", tk.END)
        # Highlight new .mm line with sel_result tag
        self.output.tag_add("sel_result", f"{new_line}.0", f"{new_line + 1}.0")
        # Move cursor and scroll (group start 2 lines from top)
        self.output.mark_set(tk.INSERT, f"{new_line}.0")
        self.output.yview(f"{group_start - 2}.0" if group_start > 2 else "1.0")
        self.update_output_result_label()

    def toggle_show_last_file(self):
        """Scroll to the last file marker in the output when the checkbox is checked."""
        if self.show_last_file_var.get():  # If checkbox is checked
            text = self.output.get("1.0", tk.END)
            lines = text.splitlines()
            # Find all file markers
            file_positions = [i + 1 for i, line in enumerate(lines) if line.strip().startswith(">>>> ") and line.strip().endswith(" <<<<")]
            if file_positions:
                last_file_line = file_positions[-1]  # Last file marker
                # Clear previous selection
                self.output.tag_remove("sel_file", "1.0", tk.END)
                self.output.tag_remove("sel_result", "1.0", tk.END)
                # Highlight the last file
                self.output.tag_add("sel_file", f"{last_file_line}.0", f"{last_file_line + 1}.0")
                # Scroll to the last file (2 lines from top if possible)
                self.output.yview(f"{last_file_line - 2}.0" if last_file_line > 2 else "1.0")
                self.output.mark_set(tk.INSERT, f"{last_file_line}.0")
                self.update_output_result_label()

    def search_output(self, term, direction):
        if not term:
            return
        # Get current cursor position
        current_index = self.output.index(tk.INSERT)
        current_line, current_col = map(int, current_index.split('.'))
        lines = self.output.get("1.0", tk.END).split('\n')
        
        # Adjust starting position based on direction
        if direction == 1:  # Down
            start_line = current_line
            start_col = current_col + 1 if current_col > 0 else 0  # Move past current position
            end_line = len(lines) + 1
            step = 1
        else:  # Up
            start_line = current_line
            start_col = current_col - 1  # Move before current position
            end_line = 0
            step = -1
        
        found = False
        wrapped = False
        
        # First search from current position to end/beginning
        for i in range(start_line, end_line, step):
            line_text = lines[i - 1].lower()
            term_lower = term.lower()
            
            if direction == 1:
                start_pos = line_text.find(term_lower, start_col if i == start_line else 0)
            else:
                start_pos = line_text.rfind(term_lower, 0, start_col + 1 if i == start_line else None)
            
            if start_pos != -1:
                # Move cursor and highlight
                self.output.mark_set(tk.INSERT, f"{i}.{start_pos}")
                self.output.see(tk.INSERT)
                self.output.tag_remove("search", "1.0", tk.END)
                self.output.tag_add("search", f"{i}.{start_pos}", f"{i}.{start_pos + len(term)}")
                found = True
                break
            start_col = 0  # Reset column for subsequent lines
        
        # Wrap around if not found
        if not found:
            wrapped = True
            if direction == 1:
                start_line = 1
                end_line = current_line
                step = 1
                start_col = 0
            else:
                start_line = len(lines)
                end_line = current_line
                step = -1
                start_col = float('inf')
            
            for i in range(start_line, end_line, step):
                line_text = lines[i - 1].lower()
                term_lower = term.lower()
                
                if direction == 1:
                    start_pos = line_text.find(term_lower)
                else:
                    start_pos = line_text.rfind(term_lower)
                
                if start_pos != -1:
                    self.output.mark_set(tk.INSERT, f"{i}.{start_pos}")
                    self.output.see(tk.INSERT)
                    self.output.tag_remove("search", "1.0", tk.END)
                    self.output.tag_add("search", f"{i}.{start_pos}", f"{i}.{start_pos + len(term)}")
                    found = True
                    break
        
        # Update result label after search
        self.update_output_result_label()
        
        if not found:
            popup = tk.Toplevel(self.master)
            popup.transient(self.master)
            popup.overrideredirect(True)  # No window decorations
            popup.geometry(f"300x50+{self.master.winfo_x() + 200}+{self.master.winfo_y() + 200}")  # Center-ish
            tk.Label(popup, text="No matches found", bg="lightblue", fg="black").pack(expand=True, fill="both")
            popup.lift()
            popup.after(2000, popup.destroy)  # Close after 2 seconds
        elif wrapped:
            popup = tk.Toplevel(self.master)
            popup.transient(self.master)
            popup.overrideredirect(True)  # No window decorations
            popup.geometry(f"300x50+{self.master.winfo_x() + 200}+{self.master.winfo_y() + 200}")  # Center-ish
            tk.Label(popup, text="Search wrapped around to the beginning/end.", bg="lightblue", fg="black").pack(expand=True, fill="both")
            popup.lift()
            popup.after(2000, popup.destroy)  # Close after 2 seconds
 

    def update_output_result_label(self):
        text = self.output.get("1.0", tk.END)
        lines = text.splitlines()
        current_pos = self.output.index(tk.INSERT).split('.')[0]  # Cursor line
        
        try:
            # Find all file markers
            file_positions = [i + 1 for i, line in enumerate(lines) if line.strip().startswith(">>>> ") and line.strip().endswith(" <<<<")]
            total_files = len(file_positions) if file_positions else 1
            
            # Find current file index (last file marker before or at cursor)
            current_line = int(current_pos)
            file_idx_list = [i for i, pos in enumerate(file_positions) if pos <= current_line]
            current_file_idx = file_idx_list[-1] + 1 if file_idx_list else 1
            current_file_line = file_positions[current_file_idx - 1] if file_idx_list else 1
            
            # Find all .mm results
            mm_positions = [i + 1 for i, line in enumerate(lines) if ".mm" in line]
            total_mm = self.result_count if hasattr(self, 'result_count') else len(mm_positions)
            
            # Define current file's end (next file marker or end of text)
            next_file_idx = current_file_idx if current_file_idx < len(file_positions) else len(file_positions)
            file_end_line = file_positions[next_file_idx] - 1 if next_file_idx < len(file_positions) else len(lines) + 1
            
            # Find current .mm result within this file (closest to cursor)
            file_mm_positions = [pos for pos in mm_positions if current_file_line < pos <= file_end_line]
            if file_mm_positions:
                mm_idx_list = [i for i, pos in enumerate(mm_positions) if pos in file_mm_positions and pos <= current_line]
                current_mm_idx = mm_idx_list[-1] + 1 if mm_idx_list else min((i + 1 for i, pos in enumerate(mm_positions) if pos in file_mm_positions), default=0)
            else:
                current_mm_idx = 0  # No .mm in this file
            
            # Update label
            self.output_result_label.config(text=f"File {current_file_idx} of {total_files}, Result {current_mm_idx} of {total_mm}")
            self.output.update_idletasks()  # Force UI refresh
        except Exception:
            self.output_result_label.config(text="File 0 of 0, Result 0 of 0")

    


#### Right click menu for self.output ------------------------------------------------------
    
    def show_context_menu(self, event):  
        try:
            file_name = None
            result_line_number = None

            self.output.focus_set()
            index = self.output.index(f"@{event.x},{event.y}")
            line_number = int(index.split('.')[0])

            start_index = self.output.search(">>>> ", f"{line_number}.0", "1.0", regexp=False, backwards=True)
            if start_index:
                line_num = int(start_index.split('.')[0])
                full_line = self.output.get(f"{line_num}.0", f"{line_num}.end").rstrip('\n')
                if full_line.startswith(">>>> "):
                    file_name = full_line[5:-5].strip()           


            line_text = self.output.get(f"{line_number}.0", f"{line_number}.end")
            match = re.match(r'(\d+)\.([mmcfl]{2})\s', line_text)
            if match:
                result_line_number = int(match.group(1))

            sel_text = self.output.get(tk.SEL_FIRST, tk.SEL_LAST) if self.output.tag_ranges(tk.SEL) else None

            context_menu = tk.Menu(self.master, tearoff=0)
            if file_name and result_line_number is not None:
                context_menu.add_command(label="Take Me To Log", command=lambda: self.take_me_to_log(file_name, result_line_number))
            else:
                context_menu.add_command(label="Take Me To Log", state=tk.DISABLED)
            context_menu.add_separator()

            if sel_text:
                context_menu.add_command(label="Search FIXS", command=lambda: self.search_fixs(sel_text))
            else:
                context_menu.add_command(label="Search FIXS", state=tk.DISABLED)

            if sel_text:
                context_menu.add_command(label="Search Intranet", command=lambda: self.search_intranet(sel_text))
            else:
                context_menu.add_command(label="Search Intranet", state=tk.DISABLED)

            if sel_text:
                search_submenu = tk.Menu(context_menu, tearoff=0)
                for name in self.search_location_order:
                    path = self.search_locations.get(name, '')
                    if os.path.exists(path.replace('/', '\\')):
                        search_submenu.add_command(label=f"{name}", 
                                                  command=lambda p=path: self.search_windows_location(sel_text, p))
                search_submenu.add_command(label="Browse", command=lambda: self.browse_search_windows(sel_text, "output"))
                context_menu.add_cascade(label="Search Windows", menu=search_submenu)
            else:
                context_menu.add_command(label="Search Windows", state=tk.DISABLED)

            context_menu.add_separator()
            context_menu.add_command(label="Copy", command=lambda: self.copy_to_clipboard(self.output))
            context_menu.add_command(label="Select All", command=lambda: self.select_all(self.output))
            context_menu.post(event.x_root, event.y_root)
        except tk.TclError:
            context_menu = tk.Menu(self.master, tearoff=0)
            context_menu.add_command(label="Take Me To Log", state=tk.DISABLED)
            context_menu.add_separator()
            context_menu.add_command(label="Search FIXS", state=tk.DISABLED)
            context_menu.add_command(label="Search Intranet", state=tk.DISABLED)
            context_menu.add_command(label="Search Windows", state=tk.DISABLED)
            context_menu.add_separator()
            context_menu.add_command(label="Copy", command=lambda: self.copy_to_clipboard(self.output))
            context_menu.add_command(label="Select All", command=lambda: self.select_all(self.output))
            context_menu.post(event.x_root, event.y_root)
        except ValueError:
            pass
    
    # def take_me_to_log and def take_me_to_log_from_event both moved to Log_Viewer section. 
                 
    def search_fixs(self, search_text):
        url = self.fixs_base_url.replace('"" ""', search_text)
        webbrowser.open(url)

    def search_intranet(self, search_text):
        url = self.intranet_base_url.replace('"" ""', search_text)
        webbrowser.open(url)

    def search_windows_location(self, search_text, path):
        path_normalized = path.replace('/', '\\')
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]  # Milliseconds, 3 digits
        display_name = f"{path_normalized} {timestamp}"
        cmd = f'explorer "search-ms:displayname={display_name}&crumb=location:{path_normalized}&query={search_text}"'
        subprocess.Popen(cmd, shell=True)

    def browse_search_windows(self, search_text, caller="output"):
        browse_window = tk.Toplevel(self.master)
        browse_window.title("Browse Search Location")
        browse_window.transient(self.master)
        browse_window.grab_set()

        tk.Label(browse_window, text="Select a folder to search your highlighted terms.", 
                 justify=tk.CENTER, wraplength=300).pack(pady=10)
        folder_var = tk.StringVar(value=self.last_browse_location.get("path", ""))
        tk.Entry(browse_window, textvariable=folder_var).pack(pady=5, fill=tk.X, padx=10)
        
        button_frame = tk.Frame(browse_window)
        button_frame.pack(pady=10)
        tk.Button(button_frame, text="Browse", 
                 command=lambda: [folder_var.set(filedialog.askdirectory(initialdir=folder_var.get() or "C:/")), 
                                  self.last_browse_location.__setitem__("path", folder_var.get()) if folder_var.get() else None]).pack(side=tk.LEFT)
        tk.Button(button_frame, text="Search Now", command=lambda: search_now()).pack(side=tk.LEFT, padx=(10, 40))
        tk.Button(button_frame, text="Close", command=browse_window.destroy).pack(side=tk.LEFT)

        def search_now():
            folder = folder_var.get()
            if not folder or not os.path.exists(folder):
                messagebox.showerror("Error", "Please select a valid folder!", parent=browse_window)
                return
            folder_normalized = folder.replace('/', '\\')
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]  # Milliseconds, 3 digits
            display_name = f"{folder_normalized} {timestamp}"
            cmd = f'explorer "search-ms:displayname={display_name}&crumb=location:{folder_normalized}&query={search_text}"'
            subprocess.Popen(cmd, shell=True)
            browse_window.destroy()

        def open_setup():
            browse_window.destroy()
            self.options_dropdown.set("Search Windows Setup")
            self.on_options_select(tk.Event())

        tk.Button(browse_window, text="Open Search Windows Setup", 
                 command=open_setup).pack(pady=5)

        browse_window.update_idletasks()
        width = browse_window.winfo_reqwidth()
        height = browse_window.winfo_reqheight()
        if caller == "output":
            x = self.master.winfo_x() + (self.master.winfo_width() - width) // 2
            y = self.master.winfo_y() + (self.master.winfo_height() - height) // 2
        else:  # log viewer
            log_window = [w for w in self.log_viewer_states if self.log_viewer_states[w]['log_text'] == caller][0]
            x = log_window.winfo_x() + (log_window.winfo_width() - width) // 2
            y = log_window.winfo_y() + (log_window.winfo_height() - height) // 2
        browse_window.geometry(f"{width}x{height}+{x}+{y}")
        browse_window.grab_release()




    def on_quick_date_select(self, event):
        selection = self.quick_date_var.get()
        if selection and selection != "Custom":
            today = datetime.today()
            if selection == "1 day":
                start_date = today # Same day for 1 day
            elif selection == "1 week":
                start_date = today - timedelta(days=6)
            elif selection == "1 month":
                start_date = today - relativedelta(months=1)  # 1 month
            elif selection == "1 year":
                start_date = today - relativedelta(years=1)  # 1 year
            elif selection == "5 years":
                start_date = today - relativedelta(years=5)  # 5 years
            self.start_date_entry.set_date(start_date.date())
            self.end_date_entry.set_date(today.date())
        elif not selection:
            self.start_date_entry.delete(0, 'end')
            self.end_date_entry.delete(0, 'end')

    def on_date_change(self, event): # Quick Change display's Custom if custom date is selected. 
        self.quick_date_var.set("Custom")
        self.date_filter_var.set(True)

    def on_quick_date_change(self, *args): # Date Filter Checkbox
        if self.quick_date_var.get():
            self.date_filter_var.set(True)


#### Main Log Search Section -----------------------------------------------------------------
#*#
    def start_search(self):
        # Debug: Log input path and extensions
        path = self.path_var.get().strip()
#        extensions_input_raw = self.extensions_var.get().strip()
#        print(f"start_search: Input path='{path}', extensions_var='{extensions_input_raw}'")
        
        if self.searching:
            popup = tk.Toplevel(self.master)
            popup.transient(self.master)
            popup.overrideredirect(True)
            message = (
                "A search is already running. Please wait for it to complete or\n"
                "abort the current search to start a new one."
            )
            tk.Label(popup, text=message, justify="left", bg="yellow", fg="black").pack(padx=10, pady=10)
            popup.update_idletasks()
            width = popup.winfo_width()
            height = popup.winfo_height()
            top = self.master.winfo_toplevel()
            top.update()
            top_x = top.winfo_x()
            top_y = top.winfo_y()
            top_width = top.winfo_width()
            top_height = top.winfo_height()
            if top_x <= 0 and top_y <= 0 and top_width < 100 and top_height < 100:
                screen_width = top.winfo_screenwidth()
                screen_height = top.winfo_screenheight()
                x = (screen_width - width) // 2
                y = (screen_height - height) // 2
            else:
                x = top_x + (top_width - width) // 2
                y = top_y + (top_height - height) // 2
            x = max(0, x)
            y = max(0, y)
            popup.geometry(f"{width}x{height}+{x}+{y}")
            popup.update()
            popup.lift()
            popup.after(4000, popup.destroy)
            return
        
        if ": " in path:
            path = path.split(": ", 1)[1].strip()
        if not path or not os.path.exists(path):
            messagebox.showerror("Error", "Please select a valid folder or file.", parent=self.master)
            return

        # Store current search state
        self.current_search_path = path
        self.current_search_filter = self.extensions_var.get().replace(' ', '')

        old_current = getattr(self, 'last_searched_path', None)
        self.path_var.set(path)

        recent_count = self.recent_count_var.get()
        recent_slots = [f"Recent {i}" for i in range(1, recent_count + 1)]
        user_shortcuts = [n for n in self.search_path_order if not n.startswith("Recent ")]

        recent_paths = [self.search_paths.get(slot) for slot in recent_slots if self.search_paths.get(slot)]
        if path in recent_paths:
            recent_paths.remove(path)
        
        if old_current and old_current != path and os.path.exists(old_current):
            existing_paths = [self.search_paths.get(n) for n in user_shortcuts] + recent_paths
            if old_current not in existing_paths:
                recent_paths.insert(0, old_current)

        for slot in recent_slots:
            self.search_paths[slot] = None
        for i, p in enumerate(recent_paths[:recent_count], 1):
            self.search_paths[f"Recent {i}"] = p

        self.refresh_paths()
        self.update_path_combo()
        if self.save_prefs_var.get():
            self.save_preferences()
        
        self.last_searched_path = path

        is_shortcut = path in [self.search_paths.get(n) for n in user_shortcuts]

        if os.path.isdir(path):
            if not hasattr(self, 'last_prompted_path') or self.last_prompted_path != path:
                has_compressed = False
                for root, _, files in os.walk(path):
                    for file in files:
                        ext = os.path.splitext(file)[1].lower()[1:]
                        if ext in ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2']:
                            has_compressed = True
                            break
                    if has_compressed:
                        break
                if has_compressed:
                    if self.advanced_extract:
                        self.extract_compressed = True
                    else:
                        self.extract_compressed = False
                        popup = tk.Toplevel(self.master)
                        popup.transient(self.master)
                        popup.overrideredirect(True)
                        message = (
                            "Searching by folder will not extract compressed files due to risk of\n"
                            "extracting too much and crashing the system. To extract compressed\n"
                            "files, select them individually from Select File."
                        )
                        tk.Label(popup, text=message, justify="left", bg="yellow", fg="black").pack(padx=10, pady=10)
                        popup.update_idletasks()
                        width = popup.winfo_width()
                        height = popup.winfo_height()
                        top = self.master.winfo_toplevel()
                        top.update()
                        top_x = top.winfo_x()
                        top_y = top.winfo_y()
                        top_width = top.winfo_width()
                        top_height = top.winfo_height()
                        if top_x <= 0 and top_y <= 0 and top_width < 100 and top_height < 100:
                            screen_width = top.winfo_screenwidth()
                            screen_height = top.winfo_screenheight()
                            x = (screen_width - width) // 2
                            y = (screen_height - height) // 2
                        else:
                            x = top_x + (top_width - width) // 2
                            y = top_y + (top_height - height) // 2
                        x = max(0, x)
                        y = max(0, y)
                        popup.geometry(f"{width}x{height}+{x}+{y}")
                        popup.update()
                        popup.lift()
                        popup.after(4000, popup.destroy)
                else:
                    self.extract_compressed = False
                self.last_prompted_path = path
            else:
                self.extract_compressed = False
        else:
            self.extract_compressed = True

        self.refresh_paths()
        self.update_path_combo()
        if self.save_prefs_var.get():
            self.save_preferences()
        
        self.last_searched_path = path

        if os.path.isfile(path):
            files = [path.replace('\\', '/')]
            path_type = "file"
            is_compressed = os.path.splitext(path)[1].lower()[1:] in ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2']
        else:
            path_type = "directory"
            files = [os.path.join(root, file).replace('\\', '/') for root, _, files in os.walk(path) for file in files]
            is_compressed = False

        current_filter = self.current_search_filter  # Use stored filter
        path_changed = not hasattr(self, 'last_path') or self.last_path != self.current_search_path
        filter_changed = current_filter != self.last_filter

        if path_changed or filter_changed:
            self.zip_cache.clear()
            if self.last_temp_dir and os.path.exists(self.last_temp_dir):
                shutil.rmtree(self.last_temp_dir, ignore_errors=True)
            self.last_temp_dir = None
            self.last_filter = current_filter

        self.last_path = self.current_search_path  # Use stored path

        # Parse filter inputs
        extensions_input = [item.strip() for item in self.extensions_var.get().split(',') if item.strip()]
        compressed_exts = ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2']
#        print(f"start_search: extensions_input={extensions_input}")
        self.extensions = {ext.lstrip('.').lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and ext.startswith('.') and '*' not in ext and '?' not in ext and '[' not in ext}
        self.exclude_extensions = {ext[1:].lstrip('.').lower() for ext in extensions_input if ext.startswith('-.') and '*' not in ext and '?' not in ext and '[' not in ext}
        self.extension_patterns = {ext.lstrip('.').lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and ext.startswith('.') and ('*' in ext or '?' in ext or '[' in ext)}
        self.exclude_extension_patterns = {ext[1:].lstrip('.').lower() for ext in extensions_input if ext.startswith('-.') and ('*' in ext or '?' in ext or '[' in ext)}
        self.include_folders = {ext.strip('/').lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and ext.startswith('/') and ext.endswith('/') and '*' not in ext and '?' not in ext and '[' not in ext}
        self.exclude_folders = {ext[1:].strip('/').lower() for ext in extensions_input if ext.startswith('-/') and ext.endswith('/') and '*' not in ext and '?' not in ext and '[' not in ext}
        self.include_folder_patterns = {ext.strip('/').lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and ext.startswith('/') and ext.endswith('/') and ('*' in ext or '?' in ext or '[' in ext)}
        self.exclude_folder_patterns = {ext[1:].strip('/').lower() for ext in extensions_input if ext.startswith('-/') and ext.endswith('/') and ('*' in ext or '?' in ext or '[' in ext)}
        self.include_files = {ext.lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and not ext.startswith(('.', '/')) and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] not in compressed_exts and '*' not in ext and '?' not in ext and '[' not in ext}
        self.exclude_files = {ext[1:].lower() for ext in extensions_input if ext.startswith('-') and not ext.startswith(('-/', '-.', ':')) and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] not in compressed_exts and '*' not in ext and '?' not in ext and '[' not in ext}
        self.include_file_patterns = {ext.lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and not ext.startswith(('.', '/')) and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] not in compressed_exts and ('*' in ext or '?' in ext or '[' in ext)}
        self.exclude_file_patterns = {ext[1:].lower() for ext in extensions_input if ext.startswith('-') and not ext.startswith(('-/', '-.', ':')) and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] not in compressed_exts and ('*' in ext or '?' in ext or '[' in ext)}
        self.include_compressed_files = {ext.lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and not ext.startswith('/') and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] in compressed_exts and '*' not in ext and '?' not in ext and '[' not in ext}
        self.include_compressed_file_patterns = {ext.lower() for ext in extensions_input if ext and not ext.startswith(('-', ':')) and not ext.startswith('/') and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] in compressed_exts and ('*' in ext or '?' in ext or '[' in ext)}
        self.exclude_compressed_files = {ext[1:].lower() for ext in extensions_input if ext.startswith('-') and not ext.startswith(('-/', '-.', ':')) and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] in compressed_exts and '*' not in ext and '?' not in ext and '[' not in ext}
        self.exclude_compressed_file_patterns = {ext[1:].lower() for ext in extensions_input if ext.startswith('-') and not ext.startswith(('-/', '-.', ':')) and '/' not in ext and os.path.splitext(ext)[1].lower()[1:] in compressed_exts and ('*' in ext or '?' in ext or '[' in ext)}
        self.stop_folders = {ext.strip('/').lower() for ext in extensions_input if ext and not ext.startswith('-') and ext.startswith(':') and ext.endswith('/:') and '*' not in ext and '?' not in ext and '[' not in ext}
        self.stop_folder_patterns = {ext.strip('/').lower() for ext in extensions_input if ext and not ext.startswith('-') and ext.startswith(':') and ext.endswith('/:') and ('*' in ext or '?' in ext or '[' in ext)}
        self.stop_compressed_files = {ext[1:-1].lower() for ext in extensions_input if ext and not ext.startswith('-') and ext.startswith(':') and ext.endswith(':') and os.path.splitext(ext[1:-1])[1].lower()[1:] in compressed_exts and '*' not in ext[1:-1] and '?' not in ext[1:-1] and '[' not in ext[1:-1]}
        self.stop_compressed_file_patterns = {ext[1:-1].lower() for ext in extensions_input if ext and not ext.startswith('-') and ext.startswith(':') and ext.endswith(':') and os.path.splitext(ext[1:-1])[1].lower()[1:] in compressed_exts and ('*' in ext[1:-1] or '?' in ext[1:-1] or '[' in ext[1:-1])}
#        print(f"start_search: Parsed filter sets: extensions={self.extensions}, exclude_extensions={self.exclude_extensions}, extension_patterns={self.extension_patterns}, exclude_extension_patterns={self.exclude_extension_patterns}, include_folders={self.include_folders}, exclude_folders={self.exclude_folders}, include_folder_patterns={self.include_folder_patterns}, exclude_folder_patterns={self.exclude_folder_patterns}, include_files={self.include_files}, exclude_files={self.exclude_files}, include_file_patterns={self.include_file_patterns}, exclude_file_patterns={self.exclude_file_patterns}, include_compressed_files={self.include_compressed_files}, include_compressed_file_patterns={self.include_compressed_file_patterns}, exclude_compressed_files={self.exclude_compressed_files}, exclude_compressed_file_patterns={self.exclude_compressed_file_patterns}, stop_folders={self.stop_folders}, stop_folder_patterns={self.stop_folder_patterns}, stop_compressed_files={self.stop_compressed_files}, stop_compressed_file_patterns={self.stop_compressed_file_patterns}")

        if not any([self.extensions, self.extension_patterns, self.include_folders, self.include_folder_patterns, self.include_files, self.include_file_patterns, self.include_compressed_files, self.include_compressed_file_patterns, self.exclude_extensions, self.exclude_extension_patterns, self.exclude_folders, self.exclude_folder_patterns, self.exclude_files, self.exclude_file_patterns, self.exclude_compressed_files, self.exclude_compressed_file_patterns, self.stop_folders, self.stop_folder_patterns, self.stop_compressed_files, self.stop_compressed_file_patterns]):
            self.extensions = None
            self.extension_patterns = set()
            self.include_folders = set()
            self.include_folder_patterns = set()
            self.include_files = set()
            self.include_file_patterns = set()
            self.include_compressed_files = set()
            self.include_compressed_file_patterns = set()
            self.exclude_extensions = set()
            self.exclude_extension_patterns = set()
            self.exclude_folders = set()
            self.exclude_folder_patterns = set()
            self.exclude_files = set()
            self.exclude_file_patterns = set()
            self.exclude_compressed_files = set()
            self.exclude_compressed_file_patterns = set()
            self.stop_folders = set()
            self.stop_folder_patterns = set()
            self.stop_compressed_files = set()
            self.stop_compressed_file_patterns = set()
#            print(f"start_search: No filter sets specified, reset to defaults")

        filtered_files = self.filter_files(files, is_compressed=is_compressed, extensions=self.extensions, exclude_extensions=self.exclude_extensions, extension_patterns=self.extension_patterns, exclude_extension_patterns=self.exclude_extension_patterns, include_folders=self.include_folders, exclude_folders=self.exclude_folders, include_folder_patterns=self.include_folder_patterns, exclude_folder_patterns=self.exclude_folder_patterns, include_files=self.include_files, exclude_files=self.exclude_files, include_file_patterns=self.include_file_patterns, exclude_file_patterns=self.exclude_file_patterns, include_compressed_files=self.include_compressed_files, include_compressed_file_patterns=self.include_compressed_file_patterns, exclude_compressed_files=self.exclude_compressed_files, exclude_compressed_file_patterns=self.exclude_compressed_file_patterns, stop_folders=self.stop_folders, stop_folder_patterns=self.stop_folder_patterns, stop_compressed_files=self.stop_compressed_files, stop_compressed_file_patterns=self.stop_compressed_file_patterns)

        if self.date_filter_var.get():
            start_date_str = self.start_date_entry.get()
            end_date_str = self.end_date_entry.get()
            if not start_date_str or not end_date_str:
                messagebox.showerror("Error", "Please select both start and end dates when date filter is enabled.", parent=self.master)
                return
            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d') + timedelta(days=1) - timedelta(seconds=1)
                filtered_files = [f for f in filtered_files if not os.path.splitext(f)[1].lower()[1:] in compressed_exts or self.is_within_date_range(self.get_file_date(f))]
            except ValueError:
                messagebox.showerror("Error", "Invalid date format.", parent=self.master)
                return

        main_terms = {term.strip().lower() for term in self.main_term_var.get().split(',') if term.strip() and not term.startswith('-')}
        exclude_main_terms = {term.strip().lower()[1:] for term in self.main_term_var.get().split(',') if term.strip().startswith('-')}
        following_terms = []
        stop_after = []
        raw_following_terms = self.following_terms_var.get().split(',')
        for term in raw_following_terms:
            term = term.strip()
            if term.startswith('--'):
                stop_after.append(term[2:].lower())
            elif term and not term.startswith('-'):
                following_terms.append(term.lower())
        exclude_following_terms = {term.strip().lower()[1:] for term in raw_following_terms if term.strip().startswith('-') and not term.strip().startswith('--')}

        use_filter_search = not main_terms and not following_terms and self.extensions_var.get().strip()

        if not use_filter_search:
            main_context = self.main_context.get().strip() or '0,0'
            try:
                main_before, main_after = map(int, re.split(r'[.,]', main_context))
            except ValueError:
                messagebox.showerror("Error", "Main Term (before,after) must be two integers separated by a comma or period.", parent=self.master)
                return

            following_context = self.following_context.get().strip() or '0,0'
            try:
                following_before, following_after = map(int, re.split(r'[.,]', following_context))
            except ValueError:
                messagebox.showerror("Error", "First Following (before,after) must be two integers separated by a comma or period.", parent=self.master)
                return
        else:
            main_before = main_after = following_before = following_after = 0

        self.output.delete('1.0', tk.END)
        self.export_button['state'] = 'normal'
        self.result_count = 0
        self.result_files = set()
        has_results = False
        error_files = []

        self.progress['value'] = 0
        self.progress['maximum'] = len(filtered_files)
        self.status_bar.config(text="Preparing search...")
        self.master.update_idletasks()

        self.search_files_queue = filtered_files  # Store files queue
        self.search_button.config(text="Abort Search", command=self.abort_search)
        self.searching = True

        self.master.after(0, self.do_search, self.search_files_queue, path_type, self.extensions, self.exclude_extensions, self.exclude_folders, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, 0, use_filter_search)


######## Bookmark Filter Files Text Entry Box
    ## Filter_Files is a hierarchical search with two levels. 
    # (Folders OR Compressed Files) AND (Files OR Extensions)
    def filter_files(self, files, is_compressed, extensions, exclude_extensions, extension_patterns, exclude_extension_patterns, include_folders, exclude_folders, include_folder_patterns, exclude_folder_patterns, include_files, exclude_files, include_file_patterns, exclude_file_patterns, include_compressed_files, include_compressed_file_patterns, exclude_compressed_files, exclude_compressed_file_patterns, stop_folders, stop_folder_patterns, stop_compressed_files, stop_compressed_file_patterns):
        filtered = []
        compressed_exts = ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2']
        
        for file in files:
            file_name = os.path.basename(file).lower()
            file_ext = os.path.splitext(file_name)[1].lower()[1:] if os.path.splitext(file_name)[1] else ''
            file_dir = os.path.normpath(os.path.dirname(file)).lower()

            # Allow compressed files for extraction unless excluded
            if is_compressed and file_ext in compressed_exts:
                if include_compressed_files and file_name in include_compressed_files:
#                    print(f"[41] filter_files: {file} included, matched include_compressed_files={file_name}")
                    filtered.append(file)
                    continue
                if include_compressed_file_patterns:
                    for pattern in include_compressed_file_patterns:
                        if fnmatch.fnmatch(file_name, pattern.lower()):
#                            print(f"[43] filter_files: {file} included, matched include_compressed_file_patterns={pattern}")
                            filtered.append(file)
                            break
                    else:
#                        print(f"[44] filter_files: {file} included as compressed file for extraction")
                        filtered.append(file)
                    continue
                if exclude_compressed_files and file_name in exclude_compressed_files:
#                    print(f"[42] filter_files: {file} skipped, in exclude_compressed_files={exclude_compressed_files}")
                    continue
                if exclude_compressed_file_patterns:
                    for pattern in exclude_compressed_file_patterns:
                        if fnmatch.fnmatch(file_name, pattern.lower()):
#                            print(f"[45] filter_files: {file} skipped, matched exclude_compressed_file_patterns={pattern}")
                            break
                    else:
#                        print(f"[46] filter_files: {file} included as compressed file for extraction")
                        filtered.append(file)
                    continue
#                print(f"[47] filter_files: {file} included as compressed file for extraction")
                filtered.append(file)
                continue

            # Folder filters
            folder_match = True
            if include_folders or include_folder_patterns:
                folder_match = False
                for folder in include_folders:
                    folder = folder.lower().strip(os.sep)
                    if folder in [part.lower() for part in file_dir.split(os.sep)]:
                        folder_match = True
#                        print(f"[11] filter_files: {file} matched include_folders={folder}")
                        break
                if not folder_match:
                    for pattern in include_folder_patterns:
                        for dir_part in file_dir.split(os.sep):
                            if fnmatch.fnmatch(dir_part.lower(), pattern.lower()):
                                folder_match = True
#                                print(f"[12] filter_files: {file} matched include_folder_patterns={pattern}")
                                break
                        if folder_match:
                            break
                if not folder_match:
#                    print(f"[13] filter_files: {file} skipped, no include folder match")
                    continue
            
            if exclude_folders:
                if any(folder.lower().strip(os.sep) in [part.lower() for part in file_dir.split(os.sep)] for folder in exclude_folders):
#                    print(f"[14] filter_files: {file} skipped, matched exclude_folders={exclude_folders}")
                    continue
            if exclude_folder_patterns:
                if any(fnmatch.fnmatch(dir_part.lower(), pattern.lower()) for pattern in exclude_folder_patterns for dir_part in file_dir.split(os.sep)):
#                    print(f"[15] filter_files: {file} skipped, matched exclude_folder_patterns={exclude_folder_patterns}")
                    continue
            
            # File and extension filters with OR logic
            file_or_ext_match = False
            # File filters
            if include_files and file_name in include_files:
                file_or_ext_match = True
#                print(f"[21] filter_files: {file} matched include_files={file_name}")
            elif include_file_patterns:
                for pattern in include_file_patterns:
                    if fnmatch.fnmatch(file_name, pattern.lower()):
                        file_or_ext_match = True
#                        print(f"[23] filter_files: {file} matched include_file_patterns={pattern}")
                        break
            # Extension filters
            if not file_or_ext_match and (extensions or extension_patterns):
                if extensions and file_ext and file_ext in extensions:
                    file_or_ext_match = True
#                    print(f"[31] filter_files: {file} matched extensions={file_ext}")
                elif extension_patterns and file_ext:
                    for pattern in extension_patterns:
                        if fnmatch.fnmatch(file_ext, pattern.lower()):
                            file_or_ext_match = True
#                            print(f"[32] filter_files: {file} matched extension_patterns={pattern}")
                            break
            # If no file or extension filters are specified, allow all files; otherwise, require a match
            if not include_files and not include_file_patterns and not extensions and not extension_patterns:
                file_or_ext_match = True
            if not file_or_ext_match:
#                print(f"[33] filter_files: {file} skipped, no match in include_files={include_files}, include_file_patterns={include_file_patterns}, extensions={extensions}, or extension_patterns={extension_patterns}")
                continue
            
            # Apply exclusion filters
            if exclude_files and file_name in exclude_files:
#                print(f"[22] filter_files: {file} skipped, in exclude_files={exclude_files}")
                continue
            if exclude_file_patterns:
                if any(fnmatch.fnmatch(file_name, pattern.lower()) for pattern in exclude_file_patterns):
#                    print(f"[25] filter_files: {file} skipped, matched exclude_file_patterns={exclude_file_patterns}")
                    continue
            if exclude_extensions and file_ext and file_ext in exclude_extensions:
#                print(f"[34] filter_files: {file} skipped, ext={file_ext} in exclude_extensions={exclude_extensions}")
                continue
            if exclude_extension_patterns and file_ext:
                if any(fnmatch.fnmatch(file_ext, pattern.lower()) for pattern in exclude_extension_patterns):
#                    print(f"[35] filter_files: {file} skipped, matched exclude_extension_patterns={exclude_extension_patterns}")
                    continue
            
#            print(f"[51] filter_files: {file} included")
            filtered.append(file)
        
        return filtered




    def is_within_date_range(self, file_date):
        """Check if a file's date is within the specified range."""
        if not self.date_filter_var.get() or file_date is None:
            return True

        start_date_str = self.start_date_entry.get()
        end_date_str = self.end_date_entry.get()
        if not start_date_str or not end_date_str:
            return True  # Skip filtering if dates aren’t set

        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d') + timedelta(days=1) - timedelta(seconds=1)
            return start_date <= file_date <= end_date
        except ValueError:
            return True  # Skip if invalid dates

    def get_file_date(self, file_entry):
        """Get the date of a file, handling compressed file contents via 7-Zip extraction."""
        if isinstance(file_entry, list) and len(file_entry) == 3 and file_entry[2]:
            return file_entry[2]  # Precomputed datetime from result_files

        if isinstance(file_entry, list):
            file_name = file_entry[0].split()[1].replace('\\', '/') if file_entry[0] else file_entry[1]
            full_path = file_entry[1]
        else:
            file_name = file_entry
            full_path = file_entry

        basename = os.path.basename(file_name)
        file_ext = os.path.splitext(basename)[1].lower()[1:]

        # Try parsing timestamp from filename
        timestamp_match = re.search(r'(\d{4,})', basename)
        if timestamp_match:
            ts = timestamp_match.group(1)
            if len(ts) >= 4:
                if len(ts) == 4:
                    ts = ts + '0101000000'  # Year only -> Jan 1, 00:00:00
                elif len(ts) == 8:
                    ts = ts + '000000'      # YYYYMMDD -> 00:00:00
                elif len(ts) == 12:
                    ts = ts + '00'          # YYYYMMDDHHMM -> 00
                try:
                    year = int(ts[:4])
                    month = int(ts[4:6])
                    day = int(ts[6:8])
                    hour = int(ts[8:10])
                    minute = int(ts[10:12])
                    second = int(ts[12:14])
                    return datetime(year, month, day, hour, minute, second)
                except ValueError:
                    pass

        # For compressed files, use mtime; for uncompressed, use ctime; extracted files check _log_scraper
        if file_ext in ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2']:
            try:
                mtime = os.path.getmtime(full_path)  # Archive: use mtime
                return datetime.fromtimestamp(mtime)
            except OSError:
                pass
        elif '_log_scraper' in full_path:  # Extracted file from compressed source
            try:
                mtime = os.path.getmtime(full_path)  # Use mtime for consistency with source archive
                return datetime.fromtimestamp(mtime)
            except OSError:
                pass
        else:
            # Regular uncompressed file: use ctime
            try:
                ctime = os.path.getctime(full_path)
                return datetime.fromtimestamp(ctime)
            except OSError:
                pass

        return None  # Fallback if all else fails
 
    def do_search(self, files, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search):
        start_time = time.time()

        if not self.searching or not self.search_files_queue:
            self.search_completed(has_results, error_files)
            return

        file = self.search_files_queue.pop(0)
        processed += 1
        file_ext = os.path.splitext(file)[1].lower()[1:]

        self.update_gui(processed, self.progress['maximum'], file)
        if not self.searching:
            self.search_completed(has_results, error_files)
            return

        try:
            if file_ext in ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2']:
                if not self.extract_compressed:
                    self.master.after(0, self.do_search, self.search_files_queue, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search)
                    return

                original_zip = self.path_var.get()
                is_path_var = file == original_zip
                if not is_path_var:
                    # Check if file matches include_compressed_files or include_compressed_file_patterns
                    file_name = os.path.basename(file).lower()
                    if file_name in self.include_compressed_files:
#                        print(f"[DS1] do_search: Processing target compressed file {file} from include_compressed_files={file_name}")
                        extract_dir = self.extract_archive(file, file_ext)
                        if extract_dir:
                            nested_files = []
                            for sub_root, _, sub_fs in os.walk(extract_dir):
                                nested_files.extend([os.path.join(sub_root, sub_f).replace('\\', '/') for sub_f in sub_fs])
                            filtered_nested_files = self.filter_files(nested_files, is_compressed=True, extensions=extensions, exclude_extensions=exclude_extensions, extension_patterns=self.extension_patterns, exclude_extension_patterns=self.exclude_extension_patterns, include_folders=self.include_folders, exclude_folders=self.exclude_folders, include_folder_patterns=self.include_folder_patterns, exclude_folder_patterns=self.exclude_folder_patterns, include_files=self.include_files, exclude_files=self.exclude_files, include_file_patterns=self.include_file_patterns, exclude_file_patterns=self.exclude_file_patterns, include_compressed_files=self.include_compressed_files, include_compressed_file_patterns=self.include_compressed_file_patterns, exclude_compressed_files=self.exclude_compressed_files, exclude_compressed_file_patterns=self.exclude_compressed_file_patterns, stop_folders=self.stop_folders, stop_folder_patterns=self.stop_folder_patterns, stop_compressed_files=self.stop_compressed_files, stop_compressed_file_patterns=self.stop_compressed_file_patterns)
                            filtered_nested_files = [f for f in filtered_nested_files if self.is_within_date_range(self.get_file_date(f))]
                            self.search_files_queue.extend(filtered_nested_files)
                            self.progress['maximum'] = processed + len(self.search_files_queue)
                            self.zip_cache[file] = extract_dir
                        else:
                            error_files.append(os.path.basename(file))
                        self.master.after(0, self.do_search, self.search_files_queue, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search)
                        return
                    for pattern in self.include_compressed_file_patterns:
                        if fnmatch.fnmatch(file_name, pattern.lower()):
#                            print(f"[DS2] do_search: Processing target compressed file {file} from include_compressed_file_patterns={pattern}")
                            extract_dir = self.extract_archive(file, file_ext)
                            if extract_dir:
                                nested_files = []
                                for sub_root, _, sub_fs in os.walk(extract_dir):
                                    nested_files.extend([os.path.join(sub_root, sub_f).replace('\\', '/') for sub_f in sub_fs])
                                filtered_nested_files = self.filter_files(nested_files, is_compressed=True, extensions=extensions, exclude_extensions=exclude_extensions, extension_patterns=self.extension_patterns, exclude_extension_patterns=self.exclude_extension_patterns, include_folders=self.include_folders, exclude_folders=self.exclude_folders, include_folder_patterns=self.include_folder_patterns, exclude_folder_patterns=self.exclude_folder_patterns, include_files=self.include_files, exclude_files=self.exclude_files, include_file_patterns=self.include_file_patterns, exclude_file_patterns=self.exclude_file_patterns, include_compressed_files=self.include_compressed_files, include_compressed_file_patterns=self.include_compressed_file_patterns, exclude_compressed_files=self.exclude_compressed_files, exclude_compressed_file_patterns=self.exclude_compressed_file_patterns, stop_folders=self.stop_folders, stop_folder_patterns=self.stop_folder_patterns, stop_compressed_files=self.stop_compressed_files, stop_compressed_file_patterns=self.stop_compressed_file_patterns)
                                filtered_nested_files = [f for f in filtered_nested_files if self.is_within_date_range(self.get_file_date(f))]
                                self.search_files_queue.extend(filtered_nested_files)
                                self.progress['maximum'] = processed + len(self.search_files_queue)
                                self.zip_cache[file] = extract_dir
                            else:
                                error_files.append(os.path.basename(file))
                            self.master.after(0, self.do_search, self.search_files_queue, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search)
                            return
                    # Apply filter_files to non-target nested compressed files
                    filtered = self.filter_files([file], is_compressed=True, extensions=extensions, exclude_extensions=exclude_extensions, extension_patterns=self.extension_patterns, exclude_extension_patterns=self.exclude_extension_patterns, include_folders=self.include_folders, exclude_folders=self.exclude_folders, include_folder_patterns=self.include_folder_patterns, exclude_folder_patterns=self.exclude_folder_patterns, include_files=self.include_files, exclude_files=self.exclude_files, include_file_patterns=self.include_file_patterns, exclude_file_patterns=self.exclude_file_patterns, include_compressed_files=self.include_compressed_files, include_compressed_file_patterns=self.include_compressed_file_patterns, exclude_compressed_files=self.exclude_compressed_files, exclude_compressed_file_patterns=self.exclude_compressed_file_patterns, stop_folders=self.stop_folders, stop_folder_patterns=self.stop_folder_patterns, stop_compressed_files=self.stop_compressed_files, stop_compressed_file_patterns=self.stop_compressed_file_patterns)
                    if not filtered:
#                        print(f"[DS3] do_search: Skipped {file} after filter_files")
                        self.master.after(0, self.do_search, self.search_files_queue, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search)
                        return

                path_changed = not hasattr(self, '_committed_last_path') or self._committed_last_path != original_zip
                filter_changed = self.extensions_var.get().replace(' ', '') != getattr(self, '_committed_last_filter', '')
                initial_search_done = getattr(self, 'initial_search_done', False)

                if path_changed or filter_changed or not initial_search_done:
                    self.zip_cache.clear()
                    if self.last_temp_dir and os.path.exists(self.last_temp_dir):
                        shutil.rmtree(self.last_temp_dir, ignore_errors=True)
                    self.last_temp_dir = None
                    extract_dir = self.extract_archive(file, file_ext)
                    if not extract_dir:
                        error_files.append(os.path.basename(file))
                        self.master.after(0, self.do_search, self.search_files_queue, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search)
                        return

                    nested_files = []
                    self.handle_nested_archives(extract_dir, nested_files, extensions, exclude_extensions, exclude_names, path_type)
                    if nested_files:  # Only process if nested_files were populated
                        filtered_nested_files = self.filter_files(nested_files, is_compressed=True, extensions=extensions, exclude_extensions=exclude_extensions, extension_patterns=self.extension_patterns, exclude_extension_patterns=self.exclude_extension_patterns, include_folders=self.include_folders, exclude_folders=self.exclude_folders, include_folder_patterns=self.include_folder_patterns, exclude_folder_patterns=self.exclude_folder_patterns, include_files=self.include_files, exclude_files=self.exclude_files, include_file_patterns=self.include_file_patterns, exclude_file_patterns=self.exclude_file_patterns, include_compressed_files=self.include_compressed_files, include_compressed_file_patterns=self.include_compressed_file_patterns, exclude_compressed_files=self.exclude_compressed_files, exclude_compressed_file_patterns=self.exclude_compressed_file_patterns, stop_folders=self.stop_folders, stop_folder_patterns=self.stop_folder_patterns, stop_compressed_files=self.stop_compressed_files, stop_compressed_file_patterns=self.stop_compressed_file_patterns)
                        filtered_nested_files = [f for f in filtered_nested_files if self.is_within_date_range(self.get_file_date(f))]
                        self.search_files_queue.extend(filtered_nested_files)
                        self.progress['maximum'] = processed + len(self.search_files_queue)
                    self.last_temp_dir = extract_dir
                    self._committed_last_path = original_zip
                    self._committed_last_filter = self.extensions_var.get().replace(' ', '')
                    self.initial_search_done = True
                else:
                    nested_files = []
                    if file in self.zip_cache and os.path.exists(self.zip_cache[file]):
                        self.handle_nested_archives(self.zip_cache[file], nested_files, extensions, exclude_extensions, exclude_names, path_type)
                        if nested_files:
                            filtered_nested_files = self.filter_files(nested_files, is_compressed=True, extensions=extensions, exclude_extensions=exclude_extensions, extension_patterns=self.extension_patterns, exclude_extension_patterns=self.exclude_extension_patterns, include_folders=self.include_folders, exclude_folders=self.exclude_folders, include_folder_patterns=self.include_folder_patterns, exclude_folder_patterns=self.exclude_folder_patterns, include_files=self.include_files, exclude_files=self.exclude_files, include_file_patterns=self.include_file_patterns, exclude_file_patterns=self.exclude_file_patterns, include_compressed_files=self.include_compressed_files, include_compressed_file_patterns=self.include_compressed_file_patterns, exclude_compressed_files=self.exclude_compressed_files, exclude_compressed_file_patterns=self.exclude_compressed_file_patterns, stop_folders=self.stop_folders, stop_folder_patterns=self.stop_folder_patterns, stop_compressed_files=self.stop_compressed_files, stop_compressed_file_patterns=self.stop_compressed_file_patterns)
                            filtered_nested_files = [f for f in filtered_nested_files if self.is_within_date_range(self.get_file_date(f))]
                            self.search_files_queue.extend(filtered_nested_files)
                            self.progress['maximum'] = processed + len(self.search_files_queue)
                            self.last_temp_dir = self.zip_cache[file]
            else:
                if self.is_within_date_range(self.get_file_date(file)):
                    try:
                        with open(file, 'r', encoding='utf-8', errors='ignore') as f:
                            lines = f.readlines()
                            processed_lines = []
                            for line in lines:
                                if len(line) > 1000 and line.count('\x00') > (len(line) - 100):
                                    stripped = line.replace('\x00', '').rstrip()
                                    processed_lines.append(stripped + ' ' if stripped else ' ')
                                else:
                                    processed_lines.append(line)
                            lines = processed_lines
                            if use_filter_search:
                                self.current_file = file
                                results = self.search_by_filter(lines)
                                if results:
                                    has_results = True
                                    self.result_count += len(results)
                                    self.result_files.add(file.replace('\\', '/'))
                                    formatted_results = [results]
                                    try:
                                        self.display_search_results(os.path.basename(file), formatted_results, [], 0, 0)
                                    except Exception as e:
                                        error_files.append(os.path.basename(file))
                            elif following_terms:
                                self.current_file = file
                                results = self.search_file_with_following_lines(lines, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, error_files)
                                if results:
                                    has_results = True
                                    self.result_count += len(results)
                                    self.result_files.add(file.replace('\\', '/'))
                                    self.current_file = file
                                    self.display_search_results(os.path.basename(file), results, following_terms, main_before, main_after)
                            else:
                                self.current_file = file
                                results = self.search_file_only_main_lines(lines, main_terms, exclude_main_terms, main_before, main_after, error_files)
                                if results:
                                    has_results = True
                                    self.result_count += len(results)
                                    self.result_files.add(file.replace('\\', '/'))
                                    self.current_file = file
                                    self.display_search_results(os.path.basename(file), results, following_terms, main_before, main_after)
                    except Exception as e:
                        error_files.append(os.path.basename(file))
        except Exception as e:
            error_files.append(os.path.basename(file))

        if self.searching:
            self.master.after(0, self.do_search, self.search_files_queue, path_type, extensions, exclude_extensions, exclude_names, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, has_results, error_files, processed, use_filter_search)
        else:
            self.search_completed(has_results, error_files)
            
    def get_context(self, lines, index, before, after):
        start = max(0, index - before)
        end = min(len(lines), index + after + 1)
        return [line.strip() for i, line in enumerate(lines[start:end])]

    def extract_archive(self, file_path, file_ext):
        """Extract an archive using 7-Zip with abort support, handling nested compression."""
        if not self.searching:
            return None  # Skip extraction if search is aborted

        extract_dir = f"{os.path.splitext(file_path)[0]}_log_scraper"
        try:
            os.makedirs(extract_dir, exist_ok=True)
            base_path = self.path_var.get()
            relative_file = os.path.relpath(file_path, base_path) if file_path.startswith(base_path) else os.path.basename(file_path)
            self.update_gui(text=f"Extracting {relative_file}...")
            zip_path = self.seven_zip_path
            if not os.path.exists(zip_path):
                raise FileNotFoundError(f"7-Zip not found at {zip_path}")
           
            # First extraction with explicit format for .bz2
            cmd = [zip_path, 'x', file_path, f'-o{extract_dir}', '-y']
            if file_ext == 'bz2':
                cmd.insert(2, '-tbzip2')  # Reintroduce for .bz2 to fix nested extraction
            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW  # Hide console window
            )
            
            while process.poll() is None:
                if not self.searching:
                    process.terminate()
                    return None
                self.master.update()
                time.sleep(0.1)
            
            stdout, stderr = process.communicate()
            if process.returncode != 0:
                return None

            # Check for double compression (e.g., .tar.gz -> .tar)
            extracted_files = [os.path.join(extract_dir, f) for f in os.listdir(extract_dir)]
            tar_file = None
            for f in extracted_files:
                if os.path.splitext(f)[1].lower()[1:] == 'tar':
                    tar_file = f
                    break
            
            if tar_file and self.searching:  # Only proceed if not aborted
                tar_extract_dir = os.path.join(extract_dir, "tar_contents")
                os.makedirs(tar_extract_dir, exist_ok=True)
                cmd = [zip_path, 'x', tar_file, f'-o{tar_extract_dir}', '-y']
                process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    creationflags=subprocess.CREATE_NO_WINDOW  # Hide console window
                )
                while process.poll() is None:
                    if not self.searching:
                        process.terminate()
                        return extract_dir
                    self.master.update()
                    time.sleep(0.1)
                stdout, stderr = process.communicate()
                if process.returncode != 0:
                    return extract_dir
                
                time.sleep(0.5)  # Ensure files are ready
                extracted_tar_files = os.listdir(tar_extract_dir)
                for item in extracted_tar_files:
                    src = os.path.join(tar_extract_dir, item)
                    dst = os.path.join(extract_dir, item)
                    if os.path.exists(dst):
                        shutil.rmtree(dst, ignore_errors=True) if os.path.isdir(dst) else os.remove(dst)
                    os.rename(src, dst)
                shutil.rmtree(tar_extract_dir, ignore_errors=True)
                os.remove(tar_file)
            else:
                pass

            self.zip_cache[file_path] = extract_dir
            # Store the top-level root folder only for the first extraction in a search
            if not hasattr(self, 'root_extract_dir') or not self.root_extract_dir:
                self.root_extract_dir = extract_dir
            self.prev_temp_dir = self.last_temp_dir
            self.last_temp_dir = extract_dir
            return extract_dir
        except Exception as e:
            return extract_dir  # Keep partial results

    def handle_nested_archives(self, extract_dir, files, extensions, exclude_extensions, exclude_names, path_type):
        if not extract_dir:
            return

        files.clear()  # Clear files to ensure only target contents are included
        compressed_exts = ['zip', '7z', 'cab', 'tar', 'gz', 'bz2']
        for root, _, fs in os.walk(extract_dir):
            root_lower = root.lower()
            root_parts = root.split(os.sep)

            # Stop extraction if folder matches include_folders or include_folder_patterns
            if self.include_folders:
                for folder in self.include_folders:
                    folder = folder.lower().strip(os.sep)
                    if folder in [part.lower() for part in root_parts]:
#                        print(f"[HNA1] handle_nested_archives: Stopped at folder {root} due to include_folders={folder}")
                        files.extend([os.path.join(root, f).replace('\\', '/') for f in fs if os.path.splitext(f)[1].lower()[1:] not in compressed_exts or f.lower() in self.include_compressed_files or any(fnmatch.fnmatch(f.lower(), p.lower()) for p in self.include_compressed_file_patterns)])
                        return
            if self.include_folder_patterns:
                for pattern in self.include_folder_patterns:
                    for part in root_parts:
                        if fnmatch.fnmatch(part.lower(), pattern.lower()):
#                            print(f"[HNA2] handle_nested_archives: Stopped at folder {root} due to include_folder_patterns={pattern}")
                            files.extend([os.path.join(root, f).replace('\\', '/') for f in fs if os.path.splitext(f)[1].lower()[1:] not in compressed_exts or f.lower() in self.include_compressed_files or any(fnmatch.fnmatch(f.lower(), p.lower()) for p in self.include_compressed_file_patterns)])
                            return

            # Skip excluded folders
            if self.exclude_folders:
                if any(folder.lower().strip(os.sep) in [part.lower() for part in root_parts] for folder in self.exclude_folders):
#                    print(f"[HNA3] handle_nested_archives: Skipped folder {root} due to exclude_folders={self.exclude_folders}")
                    continue
            if self.exclude_folder_patterns:
                if any(fnmatch.fnmatch(part.lower(), pattern.lower()) for pattern in self.exclude_folder_patterns for part in root_parts):
#                    print(f"[HNA4] handle_nested_archives: Skipped folder {root} due to exclude_folder_patterns={self.exclude_folder_patterns}")
                    continue

            for f in fs:
                nested_path = os.path.join(root, f).replace('\\', '/')
                nested_file_lower = f.lower()
                nested_ext = os.path.splitext(f)[1].lower()[1:]

                # Stop extraction if compressed file matches include_compressed_files or include_compressed_file_patterns
                if nested_ext in compressed_exts:
                    if nested_file_lower in self.include_compressed_files:
#                        print(f"[HNA5] handle_nested_archives: Stopped at compressed file {nested_path} due to include_compressed_files={nested_file_lower}")
                        extract_dir = self.extract_archive(nested_path, nested_ext)
                        if extract_dir:
                            files.clear()  # Clear previous files to include only target archive contents
                            for sub_root, _, sub_fs in os.walk(extract_dir):
                                files.extend([os.path.join(sub_root, sub_f).replace('\\', '/') for sub_f in sub_fs])
                            self.zip_cache[nested_path] = extract_dir
                        else:
                            files.append(nested_path)
                        return
                    for pattern in self.include_compressed_file_patterns:
                        if fnmatch.fnmatch(nested_file_lower, pattern.lower()):
#                            print(f"[HNA6] handle_nested_archives: Stopped at compressed file {nested_path} due to include_compressed_file_patterns={pattern}")
                            extract_dir = self.extract_archive(nested_path, nested_ext)
                            if extract_dir:
                                files.clear()  # Clear previous files to include only target archive contents
                                for sub_root, _, sub_fs in os.walk(extract_dir):
                                    files.extend([os.path.join(sub_root, sub_f).replace('\\', '/') for sub_f in sub_fs])
                                self.zip_cache[nested_path] = extract_dir
                            else:
                                files.append(nested_path)
                            return
                    # Extract non-target nested archive
                    if nested_path not in self.zip_cache or not os.path.exists(self.zip_cache[nested_path]):
                        nested_extract_dir = self.extract_archive(nested_path, nested_ext)
                        if nested_extract_dir:
                            self.zip_cache[nested_path] = nested_extract_dir
                            self.handle_nested_archives(nested_extract_dir, files, extensions, exclude_extensions, exclude_names, path_type)
                    else:
                        self.handle_nested_archives(self.zip_cache[nested_path], files, extensions, exclude_extensions, exclude_names, path_type)
                else:
                    files.append(nested_path)

    def update_gui(self, processed=None, total=None, current_file=None, text=None):
        """Update progress bar and status with relative path from search root."""
        base_path = self.path_var.get()
        if processed is not None and total is not None and current_file:
            self.progress['value'] = processed
            self.progress['maximum'] = total
            relative_file = os.path.relpath(current_file, base_path) if current_file.startswith(base_path) else os.path.basename(current_file)
            self.status_bar.config(text=f"Searching {processed}/{total} in {relative_file}")
        elif text:
            self.status_bar.config(text=text)
        self.master.update_idletasks()
        if self.searching:
            self.search_button.config(text="Abort Search", command=self.abort_search)
        else:
            self.search_button.config(text="Search", command=self.start_search)

    def display_search_results(self, file_name, results, following_terms, main_before, main_after):
        # Extract just the file name if nested within an archive
        file_identifier = file_name.split('/')[-1] if '/' in file_name else file_name
        self.output.insert(tk.END, f">>>> {file_identifier} <<<<\n", "file_marker")
        
        for result in results:
            if following_terms:
                main_term_index = main_before
                following_terms_found = set()
                for i, (line_number, line) in enumerate(result):
                    if i == main_term_index:
                        self.output.insert(tk.END, f"{line_number + 1}.mm {line.strip()}\n")
                    elif i < main_term_index:
                        self.output.insert(tk.END, f"{line_number + 1}.mc {line.strip()}\n")
                    elif i > main_term_index and i <= main_term_index + main_after:
                        self.output.insert(tk.END, f"{line_number + 1}.mc {line.strip()}\n")
                    else:
                        line_stripped = line.lower().strip()
                        if any(term in line_stripped for term in [term.lower().strip() for term in following_terms if term.lower().strip() not in following_terms_found]):
                            term_found = next(term for term in following_terms if term.lower().strip() in line_stripped)
                            self.output.insert(tk.END, f"{line_number + 1}.ff {line.strip()}\n")
                            following_terms_found.add(term_found.lower().strip())
                        else:
                            self.output.insert(tk.END, f"{line_number + 1}.fc {line.strip()}\n")
            else:
                for i, (line_number, line) in enumerate(result):
                    if i == main_before:  
                        self.output.insert(tk.END, f"{line_number + 1}.mm {line.strip()}\n")
                    elif i < main_before:
                        self.output.insert(tk.END, f"{line_number + 1}.mc {line.strip()}\n")
                    else:
                        self.output.insert(tk.END, f"{line_number + 1}.mc {line.strip()}\n")

            self.output.insert(tk.END, '\n')

    def abort_search(self):
        self.searching = False
        if hasattr(self, 'search_files_queue'):
            self.search_files_queue.clear()  # Clear the files queue to stop do_search immediately
        self.search_button.config(text="Search", command=self.start_search)
        self.search_completed(False, [])  # Trigger cleanup and GUI updates

    def search_completed(self, has_results, error_files):
        def sort_key(file_entry):
            if isinstance(file_entry, list):
                file_name = file_entry[0].split()[1].replace('\\', '/')
                basename = os.path.basename(file_name)
            else:
                basename = os.path.basename(file_entry)
            timestamp_match = re.search(r'(\d{4,})', basename)
            if timestamp_match:
                ts = timestamp_match.group(1)
                if len(ts) == 4:
                    ts = ts + '0101000000000'
                elif len(ts) == 8:
                    ts = ts + '000000000'
                elif len(ts) == 12:
                    ts = ts + '00000'
                elif len(ts) == 14:
                    ts = ts + '000'
                return (0, ts)
            full_path = file_entry if isinstance(file_entry, str) else next((f for f in self.result_files if os.path.basename(f) == basename), None)
            if full_path:
                try:
                    if os.path.splitext(full_path)[1].lower()[1:] in ['7z', 'zip', 'cab', 'tar', 'gz', 'bz2'] or '_log_scraper' in full_path:
                        time_val = os.path.getmtime(full_path)  # Compressed or extracted: use mtime
                    else:
                        time_val = os.path.getctime(full_path)  # Uncompressed: use ctime
                    dt = datetime.fromtimestamp(time_val).strftime('%Y%m%d%H%M%S') + '000'
                    return (1, dt)
                except OSError:
                    pass
            return (2, basename)

        def finalize_sort():
            self.status_bar.config(text="Sorting results...")
            self.master.update_idletasks()
            
            output_text = self.output.get("1.0", tk.END).rstrip().splitlines()
            file_blocks = []
            current_block = []
            
            for line in output_text:
                if line.startswith(">>>> ") and line.endswith(" <<<<"):
                    if current_block:
                        file_blocks.append(current_block)
                    current_block = [line]
                elif current_block:
                    current_block.append(line)
            if current_block:
                file_blocks.append(current_block)
            sorted_blocks = sorted(file_blocks, key=sort_key)
            # Clear and rebuild the output
            self.output.delete("1.0", tk.END)
            
            for block in sorted_blocks:
                for line in block:
                    # Same check for highlighting
                    if line.startswith(">>>> ") and line.endswith(" <<<<"):
                        self.output.insert(tk.END, f"{line}\n", "file_marker")
                    else:
                        self.output.insert(tk.END, f"{line}\n")
            self.output.see("1.0")

        if self.searching:
            finalize_sort()
            self.result_files = sorted(list(self.result_files), key=sort_key)
            total_searched = self.progress['maximum'] + len(error_files)
            valid_result_files = [f for f in self.result_files if os.path.exists(f)]
            if not has_results:
                result_msg = f"No results found in {total_searched} files"
                self.status_bar.config(text=result_msg)
                self.output.insert(tk.END, "No results found.\n")
            else:
                result_msg = f"Search Completed: Found {self.result_count} results in {len(valid_result_files)} out of {total_searched} files"
                self.status_bar.config(text=result_msg)
                self.output.insert(tk.END, "\n")

            # Scroll to last file if checkbox is checked
            if self.show_last_file_var.get():
                self.toggle_show_last_file()
            else:
                text = self.output.get("1.0", tk.END)
                lines = text.splitlines()
                file_positions = [i + 1 for i, line in enumerate(lines) if line.strip().startswith(">>>> ") and line.strip().endswith(" <<<<")]
                if file_positions:
                    first_file_line = file_positions[0]
                    self.output.tag_remove("sel_file", "1.0", tk.END)
                    self.output.tag_remove("sel_result", "1.0", tk.END)
                    self.output.tag_add("sel_file", f"{first_file_line}.0", f"{first_file_line + 1}.0")
                    self.output.mark_set(tk.INSERT, f"{first_file_line}.0")
                    self.output.yview(f"{first_file_line - 2}.0" if first_file_line > 2 else "1.0")
            self.update_output_result_label()

            self.master.update_idletasks()  # Force GUI update before cleanup
            
            cleanup_occurred = False
            if not hasattr(self, 'failed_cleanup_dirs'):
                self.failed_cleanup_dirs = []  # Initialize list for failed cleanup directories
            if hasattr(self, 'prev_root_extract_dir') and self.prev_root_extract_dir and os.path.exists(self.prev_root_extract_dir):
                if not hasattr(self, 'last_temp_dir') or not self.last_temp_dir or self.prev_root_extract_dir != self.last_temp_dir:
                    failed_deletions = self.cleanup_directory(self.prev_root_extract_dir)
                    if not os.path.exists(self.prev_root_extract_dir):
                        cleanup_occurred = True
                    elif failed_deletions:  # If cleanup failed and there are failed deletions
                        self.failed_cleanup_dirs.append(self.prev_root_extract_dir)
                delattr(self, 'prev_root_extract_dir')

            if hasattr(self, 'root_extract_dir') and self.root_extract_dir:
                self.prev_root_extract_dir = self.root_extract_dir
                delattr(self, 'root_extract_dir')
        else:
            # Clear error_files on abort to avoid false "file read errors"
            error_files = []
            self.status_bar.config(text="Search aborted.")
        
        # Restore state for next search
        self.last_path = self.path_var.get().strip()
        self.last_filter = self.extensions_var.get().replace(' ', '')
        if hasattr(self, 'current_search_path'):
            delattr(self, 'current_search_path')
        if hasattr(self, 'current_search_filter'):
            delattr(self, 'current_search_filter')

        self.searching = False
        self.search_complete = True
        self.search_button.config(text="Search", command=self.start_search)

        valid_files = [f for f in self.result_files if os.path.exists(f)]
        has_valid_files = bool(valid_files)
        # Update buttons and labels for each log viewer
        for window, state in self.log_viewer_states.items():
            if 'next_file_button' in state and window.winfo_exists():
                current_file_valid = state['current_file'] in valid_files
                if not has_valid_files or cleanup_occurred or not current_file_valid:
                    state['next_file_button'].config(state='disabled')
                    state['prev_file_button'].config(state='disabled')
                    state['cycle_files_checkbox'].config(state='disabled')
                else:
                    state['next_file_button'].config(state='normal')
                    state['prev_file_button'].config(state='normal')
                    state['cycle_files_checkbox'].config(state='normal')
                # Update label immediately using stored log_text
                if 'log_text' in state:
                    self.update_file_result_label(window, state['log_text'])

        if error_files and self.searching:  # Only show errors if not aborted
            error_msg = "The following files could not be read or processed:\n" + "\n".join(error_files)
            messagebox.showerror("File Read Errors", error_msg, parent=self.master)

        self.master.update_idletasks()
            
    def search_file_with_following_lines(self, lines, main_terms, exclude_main_terms, following_terms, exclude_following_terms, stop_after, main_before, main_after, following_before, following_after, error_files):
        results = []
        try:
            i = 0
            while i < len(lines):
                line_lower = lines[i].lower()
                
                # Check for main term
                for main_term in main_terms:
                    if main_term in line_lower and not any(exclude in line_lower for exclude in exclude_main_terms):
                        main_context = self.get_context(lines, i, main_before, main_after)
                        main_match = [(i - main_before + idx, line) for idx, line in enumerate(main_context)]
                        
                        found_following_terms = []
                        discovered_following = set()  # Track which following terms have been found

                        # Look for following terms
                        j = i + 1
                        while j < len(lines):
                            follow_line_lower = lines[j].lower()
                            
                            # Check for stop conditions
                            if any(term in follow_line_lower for term in stop_after) or '--' in follow_line_lower:
                                break

                            if any(exclude in follow_line_lower for exclude in exclude_following_terms):
                                j += 1
                                continue  # Skip to next line if an exclude term is found

                            # Check for another main term 
                            if any(check_main_term in follow_line_lower for check_main_term in main_terms) and not any(exclude in follow_line_lower for exclude in exclude_main_terms):
                                break  # Stop if a new main term is found

                            for term in following_terms:
                                if term in follow_line_lower and term not in discovered_following:
                                    following_context = self.get_context(lines, j, following_before, following_after)
                                    following_match = [(j - following_before + idx, line) for idx, line in enumerate(following_context)]
                                    found_following_terms.extend(following_match)
                                    discovered_following.add(term)  # Mark this term as found
                                    break  # Move to next line after finding one instance of the term
                            
                            if len(discovered_following) == len(following_terms):  # All unique following terms found
                                break

                            j += 1

                        # Add result if any following terms were found
                        if found_following_terms:
                            main_match.extend(found_following_terms)
                            results.append(main_match)
                        
                        # Move to the next line after processing this main term
                        i = j
                        break  # We've processed this main term, move to the next line check
                else:
                    i += 1  # If no main term matched, move to next line

            return results
        except Exception as e:
            return []

    def search_file_only_main_lines(self, lines, main_terms, exclude_main_terms, before, after, error_files):
        results = []
        try:
            for i, line in enumerate(lines):
                line_lower = line.lower()
                # Check if any of the exclude_main_terms are in the line before checking main terms
                if any(exclude in line_lower for exclude in exclude_main_terms):
                    continue  # Skip this line if it contains an excluded main term

                for main_term in main_terms:
                    if main_term in line_lower:
                        context = self.get_context(lines, i, before, after)
                        # Add line numbers to results
                        results.append([(i - before + idx, line) for idx, line in enumerate(context)])
                        break  # Move to the next line after finding a match
            return results
        except Exception as e:
            return []
########
    def search_by_filter(self, lines):
        """Return all lines from the file, as it has already been filtered by filter_files."""
        results = []
        if not lines:
            return results

        if not hasattr(self, 'current_file') or not self.current_file:
            return results

        # Return all lines with line numbers, as file has already been filtered
        for i, line in enumerate(lines):
            results.append((i, line))

        return results




#### Log_Viewer section -----------------------------------------------------------------

    def take_me_to_log(self, file_name, line_number):
        try:
            # Try exact match first
            file_path = next((f for f in self.result_files if os.path.basename(f) == file_name), None)
            # If no match, look for file_name as a prefix of basename (before notes)
            if not file_path:
                for f in self.result_files:
                    basename = os.path.basename(f)
                    if basename.startswith(file_name):
                        file_path = f
                        break
            if not file_path or not os.path.exists(file_path):
                messagebox.showerror("Error", f"File {file_name} not found or no longer exists.\nChecked: {file_path if file_path else 'None'}", parent=self.master)
                return
            
            self.clicked_file = file_path
            self.clicked_line = line_number
            self.open_log_at_line(file_path, line_number)
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while opening the file: {e}", parent=self.master)
    
    def take_me_to_log_from_event(self, event):
        try:
            # Get the exact line double-clicked
            index = self.output.index(tk.CURRENT)
            line_number = int(index.split('.')[0])

            # Check if it’s a result line (e.g., "4336.mm ...")
            line_text = self.output.get(f"{line_number}.0", f"{line_number}.end")
            match = re.match(r'(\d+)\.([mmcfl]{2})\s', line_text)
            if match:
                result_line_number = int(match.group(1))
                # Search backward for the file name
                start_index = self.output.search(">>>> ", f"{line_number}.0", "1.0", regexp=False, backwards=True)
                if start_index:
                    line_num = int(start_index.split('.')[0])
                    full_line = self.output.get(f"{line_num}.0", f"{line_num}.end").rstrip('\n')
                    if full_line.startswith(">>>> "):
                        file_name = full_line[5:-5].strip() 
                        self.take_me_to_log(file_name, result_line_number)
            # No action if not a result line (e.g., file name or empty)
        except ValueError:
            pass  # Silently handle invalid clicks

    def open_log_at_line(self, file_path, line_number):
#        start_total = time.time()
        try:
            file_path = file_path.replace('\\', '/')

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()

            processed_lines = []
            for line in lines:
                if len(line) > 1000 and line.count('\x00') > (len(line) - 100):
                    stripped = line.replace('\x00', '').rstrip()
                    processed_lines.append(stripped + ' ' if stripped else ' ')
                else:
                    processed_lines.append(line)
            lines = processed_lines

            new_window = tk.Toplevel(self.master)
            new_window.title(f"Log Viewer - {os.path.basename(file_path)}")
            geometry = self.preferences.get('log_viewer_geometry', "770x730+30+30") if not hasattr(self, 'current_geometry') else self.current_geometry
            new_window.geometry(geometry)
            new_window.resizable(True, True)
            new_window.lift()
            new_window.focus_set()

            self.log_viewer_states[new_window] = {
                'all_results': {},  # {file_path: [(line_num, tag), ...]}
                'current_highlight_index': -1,
                'current_file': file_path,
                'cycle_across_files': tk.BooleanVar(value=self.preferences.get('cycle_across_files', False)),
                'clicked_line': line_number if file_path == getattr(self, 'clicked_file', None) else None
            }
            state = self.log_viewer_states[new_window]

            main_frame = tk.Frame(new_window)
            main_frame.grid(row=0, column=0, sticky="nsew")
            new_window.grid_rowconfigure(0, weight=1)
            new_window.grid_columnconfigure(0, weight=1)

            log_frame = tk.Frame(main_frame)
            log_frame.grid(row=1, column=0, sticky="nsew")
            main_frame.grid_rowconfigure(1, weight=1)
            main_frame.grid_columnconfigure(0, weight=1)

            log_text = tk.Text(log_frame, wrap='word')
            log_text.grid(row=0, column=0, sticky="nsew")
            scrollbar = tk.Scrollbar(log_frame, orient="vertical", command=log_text.yview)
            scrollbar.grid(row=0, column=1, sticky="ns")
            log_text.configure(yscrollcommand=scrollbar.set)
            log_frame.grid_columnconfigure(0, weight=1)
            log_frame.grid_rowconfigure(0, weight=1)

            state['log_text'] = log_text

            for i, line in enumerate(lines):
                log_text.insert(tk.END, line)

            control_frame = tk.Frame(main_frame)
            control_frame.grid(row=0, column=0, sticky="ew")

            files_frame = tk.Frame(control_frame)
            files_frame.grid(row=0, column=0, padx=5)
            tk.Label(files_frame, text="Cycle Files").grid(row=0, column=0)
            state['next_file_button'] = tk.Button(files_frame, text="↑", command=lambda: self.cycle_files(-1, log_text, new_window))
            state['next_file_button'].grid(row=0, column=1)
            state['prev_file_button'] = tk.Button(files_frame, text="↓", command=lambda: self.cycle_files(1, log_text, new_window))
            state['prev_file_button'].grid(row=0, column=2)

            nav_frame = tk.Frame(control_frame)
            nav_frame.grid(row=0, column=1, padx=5)
            tk.Label(nav_frame, text="Cycle Results").grid(row=0, column=0)
            tk.Button(nav_frame, text="←", command=lambda: self.navigate_highlights(log_text, -1, new_window)).grid(row=0, column=1)
            tk.Button(nav_frame, text="→", command=lambda: self.navigate_highlights(log_text, 1, new_window)).grid(row=0, column=2)

            new_window.bind("<Up>", lambda event: self.cycle_files(-1, log_text, new_window) or "break")
            new_window.bind("<Down>", lambda event: self.cycle_files(1, log_text, new_window) or "break")
            log_text.bind("<Left>", lambda event: self.navigate_highlights(log_text, -1, new_window) or "break")
            log_text.bind("<Right>", lambda event: self.navigate_highlights(log_text, 1, new_window))

            state['cycle_files_checkbox'] = tk.Checkbutton(control_frame, text="Cycle Across Files", variable=state['cycle_across_files'])
            state['cycle_files_checkbox'].grid(row=0, column=2, padx=5)

            search_frame = tk.Frame(control_frame)
            search_frame.grid(row=0, column=3, padx=5)
            tk.Label(search_frame, text="Search:").grid(row=0, column=0)
            search_var = tk.StringVar()
            search_entry = tk.Entry(search_frame, textvariable=search_var, width=20)
            search_entry.grid(row=0, column=1)
            search_entry.bind("<Return>", lambda event: self.search_log(search_var.get(), 1, log_text))
            search_entry.bind("<Shift-Return>", lambda event: self.search_log(search_var.get(), -1, log_text))
            tk.Button(search_frame, text="↑", command=lambda: self.search_log(search_var.get(), -1, log_text)).grid(row=0, column=2)
            tk.Button(search_frame, text="↓", command=lambda: self.search_log(search_var.get(), 1, log_text)).grid(row=0, column=3)
            
            state['file_result_label'] = tk.Label(control_frame, text="")
            state['file_result_label'].grid(row=0, column=4, padx=5)

            tk.Button(control_frame, text="Export", command=lambda: self.export_log_viewer(file_path, log_text)).grid(row=0, column=5, sticky="e", padx=5)
            control_frame.grid_columnconfigure(5, weight=1)

            # Parse self.output for all files’ highlights
            output_text = self.output.get("1.0", tk.END)
            lines_in_output = output_text.splitlines()
            state['all_results'] = {}
            current_file = None
            file_highlights = []
            for output_line in lines_in_output:
                if output_line.startswith(">>>> ") and output_line.endswith(" <<<<"):    
                    if current_file is not None:
                        state['all_results'][current_file] = file_highlights
                        file_highlights = []
                    current_file = output_line.split('>>>>')[1].split('<<<<')[0].strip().replace('\\', '/')
                    for parse_file in self.result_files:
                        base_name = os.path.basename(parse_file)
                        if base_name == current_file or base_name.startswith(current_file) or current_file.startswith(base_name.split()[0]):
                            current_file = parse_file
                            break
                    else:
                        current_file = None
                elif current_file is not None:
                    match = re.match(r'(\d+)\.([mmcfl]{2})\s', output_line)
                    if match:
                        output_line_num = int(match.group(1))
                        tag = {
                            'mm': "main",
                            'ff': "following",
                            'mc': "main_context",
                            'fc': "following_context",
                            'cl': "clicked"
                        }.get(match.group(2), "unknown")
                        if tag == "clicked" and (current_file != getattr(self, 'clicked_file', None) or output_line_num != getattr(self, 'clicked_line', None)):
                            continue
                        file_highlights.append((output_line_num, tag))
            if current_file is not None:
                state['all_results'][current_file] = file_highlights

            # Apply highlights for current file
            for line_num, tag in state['all_results'].get(file_path, []):
                log_text.tag_add(tag, f"{line_num}.0", f"{line_num}.end")
                if file_path == getattr(self, 'clicked_file', None) and line_num == line_number:
                    log_text.tag_add("selected", f"{line_num}.0", f"{line_num}.end")

            log_text.tag_config("main", background="#9FC0D6")
            log_text.tag_config("following", background="#77C877")
            log_text.tag_config("main_context", background="#FFE0D0")
            log_text.tag_config("following_context", background="#FFE0D0")
            log_text.tag_config("clicked", background="yellow")
            log_text.tag_config("search", background="orange")
            log_text.tag_config("selected", foreground="red", font=("Arial", "12", "bold"))
            log_text.tag_config("sel", background="#1e90ff", foreground="white")
            log_text.tag_raise("selected")
            log_text.tag_raise("sel")

            if state['clicked_line'] and file_path == getattr(self, 'clicked_file', None):
                log_text.tag_add("clicked", f"{state['clicked_line']}.0", f"{state['clicked_line']}.end")

            top_line = max(0, line_number - 15) if line_number else 0
            log_text.yview(f"{top_line}.0")
            if line_number:
                log_text.mark_set(tk.INSERT, f"{line_number}.0")
                # Set current_highlight_index for the opened line if it's a main or following highlight
                filtered_highlights = [hl for hl in state['all_results'].get(file_path, []) if hl[1] in ('main', 'following')]
                sorted_highlights = sorted(filtered_highlights, key=lambda x: x[0])
                state['current_highlight_index'] = next((i for i, (line, _) in enumerate(sorted_highlights) if line == line_number), -1)
                log_text.focus_set()

            def show_context_menu(event):
                try:
                    log_text.focus_set()
                    sel_text = log_text.get(tk.SEL_FIRST, tk.SEL_LAST) if log_text.tag_ranges(tk.SEL) else None
                    context_menu = tk.Menu(new_window, tearoff=0)
                    if sel_text:
                        context_menu.add_command(label="Search FIXS", command=lambda: self.search_fixs(sel_text))
                    else:
                        context_menu.add_command(label="Search FIXS", state=tk.DISABLED)
                    if sel_text:
                        context_menu.add_command(label="Search Intranet", command=lambda: self.search_intranet(sel_text))
                    else:
                        context_menu.add_command(label="Search Intranet", state=tk.DISABLED)
                    if sel_text:
                        search_submenu = tk.Menu(context_menu, tearoff=0)
                        for name in self.search_location_order:
                            path = self.search_locations.get(name, '')
                            if os.path.exists(path.replace('/', '\\')):
                                search_submenu.add_command(label=f"{name}", 
                                                        command=lambda p=path: self.search_windows_location(sel_text, p))
                        search_submenu.add_command(label="Browse", command=lambda: self.browse_search_windows(sel_text, log_text))
                        context_menu.add_cascade(label="Search Windows", menu=search_submenu)
                    else:
                        context_menu.add_command(label="Search Windows", state=tk.DISABLED)
                    context_menu.add_separator()
                    context_menu.add_command(label="Copy", command=lambda: self.copy_to_clipboard(log_text))
                    context_menu.add_command(label="Select All", command=lambda: self.select_all(log_text))
                    context_menu.post(event.x_root, event.y_root)
                except tk.TclError:
                    context_menu = tk.Menu(new_window, tearoff=0)
                    context_menu.add_command(label="Search FIXS", state=tk.DISABLED)
                    context_menu.add_command(label="Search Intranet", state=tk.DISABLED)
                    context_menu.add_command(label="Search Windows", state=tk.DISABLED)
                    context_menu.add_separator()
                    context_menu.add_command(label="Copy", command=lambda: self.copy_to_clipboard(log_text))
                    context_menu.add_command(label="Select All", command=lambda: self.select_all(log_text))
                    context_menu.post(event.x_root, event.y_root)

            context_menu = tk.Menu(new_window, tearoff=0)
            log_text.bind("<Button-3>", show_context_menu)

            new_window.protocol("WM_DELETE_WINDOW", lambda: self.on_log_window_closing(new_window))

            self.update_file_result_label(new_window, log_text)
#            print(f"Total open_log_at_line took {time.time() - start_total:.2f} seconds")
#               To use the print above make sure to uncomment start_total at the top of this function.
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open log: {str(e)}", parent=self.master)

    def navigate_highlights(self, text_widget, direction, window):
        state = self.log_viewer_states.get(window)
        if not state or not state['all_results'].get(state['current_file'], []):
            return

        filtered_highlights = [(line, tag) for line, tag in state['all_results'].get(state['current_file'], []) if tag in ('main', 'following')]
        if not filtered_highlights:
            return

        sorted_highlights = sorted(filtered_highlights, key=lambda x: x[0])
        current_line = int(text_widget.index(tk.INSERT).split('.')[0])

        cycling_disabled = 'next_file_button' in state and state['next_file_button']['state'] == 'disabled'

        if direction == 1:
            new_index = next((i for i, (line, _) in enumerate(sorted_highlights) if line > current_line), None)
            if new_index is None:
                if not cycling_disabled and state['cycle_across_files'].get() and self.result_files:
                    file_paths = self.result_files
                    try:
                        current_idx = file_paths.index(state['current_file'])
                    except ValueError:
                        return
                    for i in range(1, len(file_paths)):
                        next_idx = (current_idx + i) % len(file_paths)
                        next_file = file_paths[next_idx]
                        next_highlights = [(hl[0], hl[1]) for hl in state['all_results'].get(next_file, []) if hl[1] in ('main', 'following')]
                        if next_highlights:
                            self.cycle_to_next_file(1, text_widget, window)
                            return
                else:
                    new_index = 0  # Wrap to first highlight in current file
                    popup = tk.Toplevel(text_widget.winfo_toplevel())
                    popup.transient(text_widget.winfo_toplevel())
                    popup.overrideredirect(True)
                    # Center on log viewer
                    log_window = text_widget.winfo_toplevel()
                    log_x = log_window.winfo_x()
                    log_y = log_window.winfo_y()
                    log_width = log_window.winfo_width()
                    log_height = log_window.winfo_height()
                    popup_width = 300
                    popup_height = 50
                    popup_x = log_x + (log_width - popup_width) // 2
                    popup_y = log_y + (log_height - popup_height) // 2
                    popup.geometry(f"{popup_width}x{popup_height}+{popup_x}+{popup_y}")
                    tk.Label(popup, text="Results wrapped around to the beginning.", bg="lightblue", fg="black").pack(expand=True, fill="both")
                    popup.lift()
                    popup.update_idletasks()  # Ensure display
                    popup.after(1000, popup.destroy)                    
        else:
            new_index = next((i for i, (line, _) in reversed(list(enumerate(sorted_highlights))) if line < current_line), None)
            if new_index is None:
                if not cycling_disabled and state['cycle_across_files'].get() and self.result_files:
                    file_paths = self.result_files
                    try:
                        current_idx = file_paths.index(state['current_file'])
                    except ValueError:
                        return
                    for i in range(1, len(file_paths)):
                        prev_idx = (current_idx - i) % len(file_paths)
                        prev_file = file_paths[prev_idx]
                        prev_highlights = [(hl[0], hl[1]) for hl in state['all_results'].get(prev_file, []) if hl[1] in ('main', 'following')]
                        if prev_highlights:
                            self.cycle_to_next_file(-1, text_widget, window)
                            return
                else:
                    new_index = len(sorted_highlights) - 1  # Wrap to last highlight in current file
                    popup = tk.Toplevel(text_widget.winfo_toplevel())
                    popup.transient(text_widget.winfo_toplevel())
                    popup.overrideredirect(True)
                    # Center on log viewer
                    log_window = text_widget.winfo_toplevel()
                    log_x = log_window.winfo_x()
                    log_y = log_window.winfo_y()
                    log_width = log_window.winfo_width()
                    log_height = log_window.winfo_height()
                    popup_width = 300
                    popup_height = 50
                    popup_x = log_x + (log_width - popup_width) // 2
                    popup_y = log_y + (log_height - popup_height) // 2
                    popup.geometry(f"{popup_width}x{popup_height}+{popup_x}+{popup_y}")
                    tk.Label(popup, text="Results wrapped around to the end.", bg="lightblue", fg="black").pack(expand=True, fill="both")
                    popup.lift()
                    popup.update_idletasks()  # Ensure display
                    popup.after(1000, popup.destroy)

        if new_index is not None:
            new_line = sorted_highlights[new_index][0]
            text_widget.tag_remove("selected", "1.0", tk.END)
            text_widget.tag_add("selected", f"{new_line}.0", f"{new_line}.end")
            text_widget.mark_set(tk.INSERT, f"{new_line}.0")
            state['current_highlight_index'] = new_index
            
            top_line = max(0, new_line - 15)
            text_widget.yview(f"{top_line}.0")
            text_widget.update_idletasks()
            
            self.update_file_result_label(window, text_widget)

    def update_file_result_label(self, window, text_widget):
        state = self.log_viewer_states.get(window)
        if not state or 'file_result_label' not in state:
            return  # Exit if state or label is missing
        
        try:
            file_paths = self.result_files
            cycling_disabled = 'next_file_button' in state and state['next_file_button']['state'] == 'disabled'
            
            if cycling_disabled or not file_paths:
                # Single-file mode when cycling is disabled or no files
                filtered_highlights = [hl for hl in state['all_results'].get(state['current_file'], []) if hl[1] in ('main', 'following')]
                sorted_highlights = sorted(filtered_highlights, key=lambda x: x[0])
                result_idx = state['current_highlight_index'] + 1 if state['current_highlight_index'] >= 0 else 0
                state['file_result_label'].config(text=f"File 1 of 1, Result {result_idx} of {len(sorted_highlights)}")
            else:
                # Normal multi-file mode
                current_file_idx = file_paths.index(state['current_file']) + 1  # 1-based
                total_files = len(file_paths)
                filtered_highlights = [hl for hl in state['all_results'].get(state['current_file'], []) if hl[1] in ('main', 'following')]
                sorted_highlights = sorted(filtered_highlights, key=lambda x: x[0])
                result_idx = state['current_highlight_index'] + 1 if state['current_highlight_index'] >= 0 else 0
                state['file_result_label'].config(text=f"File {current_file_idx} of {total_files}, Result {result_idx} of {len(sorted_highlights)}")
            
            text_widget.update_idletasks()  # Force UI refresh
        except Exception:
            pass  # Silently handle errors

    def position_popup_bottom_right(self, parent_window, filename):
        popup = tk.Toplevel(parent_window)
        text = f"Moving to new file: {os.path.basename(filename)}"
        font = ("Arial", 14)  # Bigger font
        label = tk.Label(popup, text=text, font=font)
        label.pack(pady=10, padx=10)

        # Calculate width based on text length and font
        width = len(text) * 10 + 20  # ~10px per char at size 12 + padding
        height = 60  # Fixed height, adjusted for bigger font

        parent_x = parent_window.winfo_x()
        parent_y = parent_window.winfo_y()
        parent_width = parent_window.winfo_width()
        parent_height = parent_window.winfo_height()

        popup_x = parent_x + parent_width - width
        popup_y = parent_y + parent_height - height

        popup.geometry(f"{width}x{height}+{popup_x}+{popup_y}")
        popup.title("File Change")
        popup.transient(parent_window)
        popup.after(1000, popup.destroy)  # Auto-close after 1 second
    
    def cycle_files(self, direction, log_text, window):
        state = self.log_viewer_states.get(window)
        if not state or not self.result_files:
            return

        file_paths = self.result_files
        try:
            current_idx = file_paths.index(state['current_file'])
        except ValueError:
            return

        new_idx = (current_idx + direction) % len(file_paths)
        new_file = file_paths[new_idx]
        if not os.path.exists(new_file):
            messagebox.showerror("Error", f"File not found: {new_file}", parent=window)
            return

        self.position_popup_bottom_right(window, new_file)

        log_text.delete('1.0', tk.END)
        with open(new_file, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            processed_lines = []
            for line in lines:
                if len(line) > 1000 and line.count('\x00') > (len(line) - 100):
                    stripped = line.replace('\x00', '').rstrip()
                    processed_lines.append(stripped + ' ' if stripped else ' ')
                else:
                    processed_lines.append(line)
            lines = processed_lines
            log_text.insert('1.0', ''.join(lines))

        state['current_file'] = new_file
        filtered_highlights = [hl for hl in state['all_results'].get(new_file, []) if hl[1] in ('main', 'following')]
        sorted_highlights = sorted(filtered_highlights, key=lambda x: x[0])
        state['current_highlight_index'] = 0 if sorted_highlights else -1
        for line_num, tag in state['all_results'].get(new_file, []):
            log_text.tag_add(tag, f"{line_num}.0", f"{line_num}.end")
        if new_file == getattr(self, 'clicked_file', None) and getattr(self, 'clicked_line', None):
            log_text.tag_add("clicked", f"{self.clicked_line}.0", f"{self.clicked_line}.end")

        target_line = sorted_highlights[0][0] if sorted_highlights else 1
        log_text.tag_remove("selected", "1.0", tk.END)
        log_text.tag_add("selected", f"{target_line}.0", f"{target_line}.end")
        log_text.mark_set(tk.INSERT, f"{target_line}.0")
        top_line = max(0, target_line - 15)
        log_text.yview(f"{top_line}.0")
        window.title(f"Log Viewer - {os.path.basename(new_file)}")
        self.update_file_result_label(window, log_text)

    def cycle_to_next_file(self, direction, log_text, window):
        state = self.log_viewer_states.get(window)
        if not state or not self.result_files:
            return
        
        file_paths = self.result_files
        valid_files = [f for f in file_paths if os.path.exists(f)]
        if not valid_files:
            if 'next_file_button' in state:
                state['next_file_button'].config(state='disabled')
                state['prev_file_button'].config(state='disabled')
                state['cycle_files_checkbox'].config(state='disabled')
            return
        
        try:
            current_idx = file_paths.index(state['current_file'])
        except ValueError:
            return
        
        new_idx = (current_idx + direction) % len(file_paths)
        new_file = file_paths[new_idx]
        if not os.path.exists(new_file):
            return
        
        self.position_popup_bottom_right(window, new_file)

        log_text.delete('1.0', tk.END)
        with open(new_file, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            processed_lines = []
            for line in lines:
                if len(line) > 1000 and line.count('\x00') > (len(line) - 100):
                    stripped = line.replace('\x00', '').rstrip()
                    processed_lines.append(stripped + ' ' if stripped else ' ')
                else:
                    processed_lines.append(line)
            lines = processed_lines
            log_text.insert('1.0', ''.join(lines))

        state['current_file'] = new_file
        filtered_highlights = [hl for hl in state['all_results'].get(new_file, []) if hl[1] in ('main', 'following')]
        sorted_highlights = sorted(filtered_highlights, key=lambda x: x[0])
        state['current_highlight_index'] = 0 if direction == 1 and sorted_highlights else len(sorted_highlights) - 1 if direction == -1 and sorted_highlights else -1
        for line_num, tag in state['all_results'].get(new_file, []):
            log_text.tag_add(tag, f"{line_num}.0", f"{line_num}.end")
        if new_file == getattr(self, 'clicked_file', None) and getattr(self, 'clicked_line', None):
            log_text.tag_add("clicked", f"{self.clicked_line}.0", f"{self.clicked_line}.end")

        target_line = sorted_highlights[0][0] if direction == 1 and sorted_highlights else sorted_highlights[-1][0] if direction == -1 and sorted_highlights else 1
        log_text.tag_add("selected", f"{target_line}.0", f"{target_line}.end")
        log_text.mark_set(tk.INSERT, f"{target_line}.0")
        top_line = max(0, target_line - 15)
        log_text.yview(f"{top_line}.0")
        window.title(f"Log Viewer - {os.path.basename(new_file)}")
        self.update_file_result_label(window, log_text)

    def search_log(self, term, direction, text_widget):
        if not term:
            return
        
        # Get current cursor position
        current_index = text_widget.index(tk.INSERT)
        current_line, current_col = map(int, current_index.split('.'))
        lines = text_widget.get("1.0", tk.END).split('\n')
        
        # Adjust starting position based on direction
        if direction == 1:  # Down
            start_line = current_line
            start_col = current_col + 1 if current_col > 0 else 0  # Move past current position for down search
            end_line = len(lines) + 1
            step = 1
        else:  # Up
            start_line = current_line
            start_col = current_col - 1  # Move before current position for up search
            end_line = 0
            step = -1

        found = False
        wrapped = False

        # First search from current position to end/beginning
        for i in range(start_line, end_line, step):
            line_text = lines[i - 1].lower()
            term_lower = term.lower()
            
            if direction == 1:
                start_pos = line_text.find(term_lower, start_col if i == start_line else 0)
            else:
                start_pos = line_text.rfind(term_lower, 0, start_col + 1 if i == start_line else None)

            if start_pos != -1:
                # Move cursor and highlight
                text_widget.mark_set(tk.INSERT, f"{i}.{start_pos}")
                text_widget.see(tk.INSERT)
                
                # Remove previous highlights
                text_widget.tag_remove("search", "1.0", tk.END)
                
                # Highlight the match
                text_widget.tag_add("search", f"{i}.{start_pos}", f"{i}.{start_pos + len(term)}")
                found = True
                break
            start_col = 0  # Reset column for subsequent lines

        # If not found, wrap around and search from the opposite end
        if not found:
            wrapped = True
            if direction == 1:
                start_line = 1
                end_line = current_line
                step = 1
                start_col = 0
            else:
                start_line = len(lines)
                end_line = current_line
                step = -1
                start_col = float('inf')

            for i in range(start_line, end_line, step):
                line_text = lines[i - 1].lower()
                term_lower = term.lower()
                
                if direction == 1:
                    start_pos = line_text.find(term_lower)
                else:
                    start_pos = line_text.rfind(term_lower)

                if start_pos != -1:
                    text_widget.mark_set(tk.INSERT, f"{i}.{start_pos}")
                    text_widget.see(tk.INSERT)
                    
                    text_widget.tag_remove("search", "1.0", tk.END)
                    text_widget.tag_add("search", f"{i}.{start_pos}", f"{i}.{start_pos + len(term)}")
                    found = True
                    break

        if not found:
            popup = tk.Toplevel(text_widget.winfo_toplevel())
            popup.transient(text_widget.winfo_toplevel())
            popup.overrideredirect(True)
            # Center on log viewer
            log_window = text_widget.winfo_toplevel()
            log_x = log_window.winfo_x()
            log_y = log_window.winfo_y()
            log_width = log_window.winfo_width()
            log_height = log_window.winfo_height()
            popup_width = 300
            popup_height = 50
            popup_x = log_x + (log_width - popup_width) // 2
            popup_y = log_y + (log_height - popup_height) // 2
            popup.geometry(f"{popup_width}x{popup_height}+{popup_x}+{popup_y}")
            tk.Label(popup, text="No matches found.", bg="lightblue", fg="black").pack(expand=True, fill="both")
            popup.lift()
            popup.update_idletasks()  # Ensure display
            popup.after(2000, popup.destroy)
        elif wrapped:
            popup = tk.Toplevel(text_widget.winfo_toplevel())
            popup.transient(text_widget.winfo_toplevel())
            popup.overrideredirect(True)
            # Center on log viewer
            log_window = text_widget.winfo_toplevel()
            log_x = log_window.winfo_x()
            log_y = log_window.winfo_y()
            log_width = log_window.winfo_width()
            log_height = log_window.winfo_height()
            popup_width = 300
            popup_height = 50
            popup_x = log_x + (log_width - popup_width) // 2
            popup_y = log_y + (log_height - popup_height) // 2
            popup.geometry(f"{popup_width}x{popup_height}+{popup_x}+{popup_y}")
            tk.Label(popup, text="Search wrapped around to the beginning/end.", bg="lightblue", fg="black").pack(expand=True, fill="both")
            popup.lift()
            popup.update_idletasks()  # Ensure display
            popup.after(2000, popup.destroy)
    
    def export_log_viewer(self, file_path, log_text):
        try:
            now = datetime.now().strftime("%Y%m%d%H%M%S%f")[:-3]
            initial_filename = f"{os.path.basename(file_path)}_{now}_highlighted"
            # Find the log viewer window for the given file_path
            log_viewer = None
            for window, state in self.log_viewer_states.items():
                if state['current_file'] == file_path and window.winfo_exists():
                    log_viewer = window
                    break

            # Prepare log viewer geometry (or fallback to main window)
            if log_viewer:
                log_geometry = log_viewer.winfo_geometry()  # Format: "widthxheight+x+y"
                log_width, log_height, log_x, log_y = map(int, re.split('[x+]', log_geometry))
            else:
                self.master.update_idletasks()
                log_x = self.master.winfo_x()
                log_y = self.master.winfo_y()
                log_width = self.master.winfo_width()
                log_height = self.master.winfo_height()

            # Set initial directory based on self.path_var
            initial_path = self.path_var.get().strip()
            if ": " in initial_path:
                initial_path = initial_path.split(": ", 1)[1].strip()  # Strip "Recent 1: " or "(Current): "
            if os.path.isfile(initial_path):  # If it's a file, use parent directory
                initial_path = os.path.dirname(initial_path)
            if not initial_path or not os.path.exists(initial_path):
                initial_path = "C:/"

            # File dialog centered on log viewer via parent
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=initial_filename,
                initialdir=initial_path,  # Added
                filetypes=[("DOCX files", "*.docx"), ("Text files", "*.txt"), ("All files", "*.*")],
                title="Save Log With Highlights",
                parent=log_viewer if log_viewer else self.master
            )
            if not save_path:
                return

            # Progress window centered on log viewer
            progress_window = tk.Toplevel(self.master)
            progress_window.title("Exporting...")
            progress_width = 300
            progress_height = 120
            center_x = log_x + (log_width - progress_width) // 2
            center_y = log_y + (log_height - progress_height) // 2
            center_x = max(0, min(center_x, self.master.winfo_screenwidth() - progress_width))
            center_y = max(0, min(center_y, self.master.winfo_screenheight() - progress_height))
            progress_window.geometry(f"{progress_width}x{progress_height}+{center_x}+{center_y}")
            progress_window.transient(self.master)
            progress_window.grab_set()
            tk.Label(progress_window, text="Exporting document...").pack(pady=10)
            progress = ttk.Progressbar(progress_window, mode='determinate', maximum=100)
            progress.pack(pady=10, padx=20, fill=tk.X)
            cancel_var = tk.BooleanVar(value=False)
            tk.Button(progress_window, text="Cancel", command=lambda: cancel_var.set(True)).pack(pady=5)
            progress_window.update()

            if save_path.endswith('.docx'):
                doc = Document()
                sections = doc.sections
                for section in sections:
                    section.top_margin = Pt(0)
                    section.bottom_margin = Pt(0)
                    section.left_margin = Pt(0)
                    section.right_margin = Pt(0)
                    section.page_height = Inches(1000)
                
                end_index = log_text.index(tk.END)
                current_index = "1.0"
                total_lines = int(end_index.split('.')[0]) - 1
                processed_lines = 0
                
                while log_text.compare(current_index, "<", end_index):
                    if cancel_var.get():
                        progress_window.destroy()
                        return

                    line_end = log_text.index(f"{current_index} lineend")
                    line_text = log_text.get(current_index, line_end)
                    tags = log_text.tag_names(current_index)

                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    run = p.add_run(line_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    if "main" in tags:
                        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    elif "following" in tags:
                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    elif "main_context" in tags or "following_context" in tags:
                        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                    elif "clicked" in tags:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif "search" in tags:
                        run.font.highlight_color = WD_COLOR_INDEX.DARK_YELLOW

                    current_index = log_text.index(f"{line_end} + 1 char")
                    processed_lines += 1
                    progress['value'] = (processed_lines / total_lines) * 100
                    progress_window.update()

                if not cancel_var.get():
                    doc.save(save_path)
            else:  # .txt export
                with open(save_path, 'w', encoding='utf-8') as f:
                    end_index = log_text.index(tk.END)
                    current_index = "1.0"
                    total_lines = int(end_index.split('.')[0]) - 1
                    processed_lines = 0

                    while log_text.compare(current_index, "<", end_index):
                        if cancel_var.get():
                            progress_window.destroy()
                            return
                        line_end = log_text.index(f"{current_index} lineend")
                        line_text = log_text.get(current_index, line_end)
                        f.write(line_text + '\n')
                        current_index = log_text.index(f"{line_end} + 1 char")
                        processed_lines += 1
                        progress['value'] = (processed_lines / total_lines) * 100
                        progress_window.update()

            progress_window.destroy()
            if not cancel_var.get():
                # Success popup centered on log viewer
                popup = tk.Toplevel(self.master)
                popup.title("Export")
                popup_width = 300
                popup_height = 120
                center_x = log_x + (log_width - popup_width) // 2
                center_y = log_y + (log_height - popup_height) // 2
                center_x = max(0, min(center_x, self.master.winfo_screenwidth() - popup_width))
                center_y = max(0, min(center_y, self.master.winfo_screenheight() - popup_height))
                popup.geometry(f"{popup_width}x{popup_height}+{center_x}+{center_y}")
                popup.transient(self.master)
                popup.grab_set()
                tk.Label(popup, text="Export finished successfully!\nOpen file?", font=("Arial", "12")).pack(pady=10)
                btn_frame = tk.Frame(popup)
                btn_frame.pack(pady=10)
                tk.Button(btn_frame, text="Yes", command=lambda: [os.startfile(save_path), popup.destroy()]).pack(side=tk.LEFT, padx=5)
                tk.Button(btn_frame, text="No", command=popup.destroy).pack(side=tk.LEFT, padx=5)
        except Exception as e:
            progress_window.destroy()
            messagebox.showerror("Error", f"Failed to export log: {str(e)}", parent=self.master)

    def export_to_docx(self, output_text, file_path, progress, cancel_var, progress_window):
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Pt(0)
                section.bottom_margin = Pt(0)
                section.left_margin = Pt(0)
                section.right_margin = Pt(0)
                section.page_height = Inches(1000)

            lines = output_text.splitlines()
            total_lines = len(lines)
            text_widget = self.output

            for i, line in enumerate(lines):
                if cancel_var.get():
                    return False
                if not line.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    continue
                start_idx = text_widget.search(line, "1.0", tk.END)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(line)
                run.font.name = "Courier New"
                if start_idx:
                    tags = text_widget.tag_names(start_idx)
                    run.font.size = Pt(14) if "file_marker" in tags else Pt(11)
                else:
                    run.font.size = Pt(11)
                progress['value'] = ((i + 1) / total_lines) * 100
                progress_window.update()

            doc.save(file_path)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export to DOCX: {str(e)}", parent=self.master)
            return False

    def on_log_window_closing(self, window): 
    # Saves to json when log_viewer closes
        self.current_geometry = window.winfo_geometry()
        if hasattr(self, 'save_prefs_var') and self.save_prefs_var.get():
            self.preferences['log_viewer_geometry'] = self.current_geometry
            state = self.log_viewer_states.get(window)
            if state:
                self.preferences['cycle_across_files'] = state['cycle_across_files'].get()
            try:
                with open(self.config_path, 'w') as f:
                    json.dump(self.preferences, f, indent=4)
            except Exception as e:
                pass
        if window in self.log_viewer_states:
            del self.log_viewer_states[window]
        window.destroy()


#### Cleanup and Closing Application. -----------------------------------------------------------------------

    def cleanup_directory(self, path):
        """Attempt to clean up a directory and its contents with retry logic, tracking failures."""
        failed_deletions = []
        for root, dirs, files in os.walk(path, topdown=False):
            for name in files:
                file_path = os.path.join(root, name)
                retries = 3
                deleted = False
                # Method 1: Standard file deletion with retries using os.remove
                for attempt in range(retries):
                    try:
                        os.remove(file_path)
                        deleted = True
                        break
                    except PermissionError as e:
                        if attempt < retries - 1:
                            time.sleep(1)
                    except FileNotFoundError as e:
                        deleted = True
                        break
                    except Exception as e:
                        break
                # Method 2: Force deletion with os.system (Windows)
                if not deleted and os.name == 'nt':
                    try:
                        os.system(f'del /F /Q "{file_path}" >nul 2>&1')  # Silent output
                        if not os.path.exists(file_path):
                            deleted = True
                    except Exception as e:
                        pass
                if not deleted:
                    failed_deletions.append(file_path)
            for name in dirs:
                dir_path = os.path.join(root, name)
                deleted = False
                # Method 3: Standard directory deletion with os.rmdir
                try:
                    os.rmdir(dir_path)
                    deleted = True
                except OSError as e:
                    # Method 4: Force directory deletion with shutil.rmtree
                    try:
                        shutil.rmtree(dir_path, ignore_errors=False)
                        deleted = True
                    except Exception as e2:
                        # Method 5: Force directory deletion with os.system (Windows)
                        if os.name == 'nt':
                            try:
                                os.system(f'rmdir /S /Q "{dir_path}" >nul 2>&1')  # Silent output
                                if not os.path.exists(dir_path):
                                    deleted = True
                            except Exception as e3:
                                pass
                if not deleted:
                    failed_deletions.append(dir_path)
        # Method 6: Standard root directory deletion with os.rmdir
        deleted = False
        try:
            os.rmdir(path)
            deleted = True
        except OSError as e:
            # Method 7: Force root directory deletion with shutil.rmtree
            try:
                shutil.rmtree(path, ignore_errors=False)
                deleted = True
            except Exception as e2:
                # Method 8: Force root directory deletion with os.system (Windows)
                if os.name == 'nt':
                    try:
                        os.system(f'rmdir /S /Q "{path}" >nul 2>&1')  # Silent output
                        if not os.path.exists(path):
                            deleted = True
                    except Exception as e3:
                        pass
        if not deleted:
            failed_deletions.append(path)
        return failed_deletions
 
    def on_closing(self):
        if self.save_prefs_var.get():
            self.save_preferences()
        
        self.status_bar.config(text="Cleaning up extracted files...")
        self.master.update_idletasks()

        # Temporary popup for cleanup
        cleanup_popup = tk.Toplevel(self.master)
        cleanup_popup.overrideredirect(True)
        tk.Label(cleanup_popup, text="Cleaning up extracted files...", bg="yellow", fg="black", font=("Arial", 14)).pack(padx=10, pady=5)
        cleanup_popup.update_idletasks()
        popup_width = cleanup_popup.winfo_width()
        popup_height = cleanup_popup.winfo_height()
        x = self.master.winfo_x() + (self.master.winfo_width() - popup_width) // 2
        y = self.master.winfo_y() + (self.master.winfo_height() - popup_height) // 2
        cleanup_popup.geometry(f"{popup_width}x{popup_height}+{x}+{y}")
        cleanup_popup.lift()
        cleanup_popup.update()  # Ensure it’s rendered

        failed_deletions = []
        base_extract_dirs = set()
        for temp_dir in list(self.zip_cache.values()):
            if os.path.exists(temp_dir):
                base_extract_dirs.add(temp_dir)
                failed_deletions.extend(self.cleanup_directory(temp_dir))
        
        if self.last_temp_dir and self.last_temp_dir not in self.zip_cache.values() and os.path.exists(self.last_temp_dir):
            base_extract_dirs.add(self.last_temp_dir)
            failed_deletions.extend(self.cleanup_directory(self.last_temp_dir))
        
        if hasattr(self, 'root_extract_dir') and self.root_extract_dir and os.path.exists(self.root_extract_dir):
            failed_deletions.extend(self.cleanup_directory(self.root_extract_dir))
        
        # Attempt cleanup of directories failed in search_completed
        if hasattr(self, 'failed_cleanup_dirs'):
            for failed_dir in self.failed_cleanup_dirs:
                if os.path.exists(failed_dir):
                    failed_deletions.extend(self.cleanup_directory(failed_dir))
        
        failed_deletions = [path for path in failed_deletions if os.path.exists(path)]
        if failed_deletions:
            # Handle failed deletions (unchanged logic)
            tree_msg = ["Folder(s) to delete:"]
            # Determine the parent folder to open and display
            if failed_dirs := sorted(set(failed_path for failed_path in failed_deletions if os.path.isdir(failed_path))):
                explorer_dir = os.path.dirname(failed_dirs[0])  # Parent of first failed folder
            elif hasattr(self, 'root_extract_dir') and self.root_extract_dir:
                explorer_dir = os.path.dirname(self.root_extract_dir)
            elif self.last_temp_dir:
                explorer_dir = os.path.dirname(self.last_temp_dir)
            else:
                explorer_dir = self.path_var.get() if self.path_var.get() else ""
            parent_dir_name = os.path.basename(explorer_dir) if explorer_dir else ""
            
            # List failed folders with their immediate parent
            for failed_dir in failed_dirs:
                dir_name = os.path.basename(failed_dir)
                tree_msg.append(f"  /{parent_dir_name}/{dir_name}/")
            
            tree_msg.append("")  # Blank line
            tree_msg.append("File(s) causing deletion failure:")
            # List only files that still exist
            failed_files = sorted(set(failed_path for failed_path in failed_deletions if os.path.isfile(failed_path)))
            for failed_file in failed_files:
                tree_msg.append(f"  {os.path.basename(failed_file)}")
            tree_msg.append("")  # Extra blank line
            tree_msg.append("")  # Second extra blank line
            
            full_msg = "\n".join(tree_msg)
            popup_msg = f"Please delete these folder(s) manually.\nWould you like to open the parent folder?"
            
            popup = tk.Toplevel(self.master)
            popup.title("Cleanup Failed")
            popup.transient(self.master)
            popup.grab_set()

            msg_label = tk.Label(popup, text=popup_msg, justify=tk.LEFT, wraplength=380)
            msg_label.pack(pady=5)

            button_frame = tk.Frame(popup)
            button_frame.pack(pady=5)
            explorer_dir = explorer_dir.replace('/', '\\')  # Ensure backslashes
            explorer_cmd = f'explorer "{explorer_dir}"'
            yes_button = tk.Button(button_frame, text="Yes", command=lambda: [subprocess.Popen(explorer_cmd, shell=True), popup.destroy()])
            yes_button.pack(side=tk.LEFT, padx=5)
            no_button = tk.Button(button_frame, text="No", command=popup.destroy)
            no_button.pack(side=tk.LEFT, padx=5)

            details_label = tk.Label(popup, text=full_msg, justify=tk.LEFT, wraplength=380)
            details_label.pack(pady=5)

            popup.update_idletasks()
            width = max(msg_label.winfo_reqwidth(), button_frame.winfo_reqwidth(), details_label.winfo_reqwidth()) + 10
            height = msg_label.winfo_reqheight() + button_frame.winfo_reqheight() + details_label.winfo_reqheight() + 20
            popup.geometry(f"{width}x{height}")
            
            x = (self.master.winfo_screenwidth() - width) // 2
            y = (self.master.winfo_screenheight() - height) // 2
            popup.geometry(f"+{x}+{y}")
            popup.wait_window()
        
        self.master.destroy()
  
     
# Main script execution
if __name__ == "__main__":
    root = tk.Tk()
    app = TextFileSearchGUI(root)
    root.mainloop()